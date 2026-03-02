"""
Generate Word Document: SHAP Explainability Notebook 04 - 4 Layers x 3 Perspectives
Agent Fraud Sentinel (BAFS) Project

Produces: docs/shap_04_analysis_matrix.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent.parent
FIGURES_DIR = BASE_DIR / "figures" / "shap"
OUTPUT_PATH = BASE_DIR / "docs" / "shap_04_analysis_matrix.docx"

# ── Colors ─────────────────────────────────────────────────────────────
C_DARK_BLUE = "1B4F72"
C_MED_BLUE = "2980B9"
C_LIGHT_BLUE = "D6EAF8"
C_FRAUD_RED = "E74C3C"
C_LEGIT_GREEN = "2ECC71"
C_LIGHT_GRAY = "F2F3F4"
C_WHITE = "FFFFFF"
C_BLACK = "000000"
C_DARK_GRAY = "2C3E50"
C_INSIGHT_BG = "D6EAF8"
C_INSIGHT_BORDER = "2980B9"
C_BUSINESS_BG = "D5F5E3"
C_BUSINESS_BORDER = "27AE60"
C_CAUTION_BG = "FADBD8"
C_CAUTION_BORDER = "E74C3C"

LAYERS = ["WHAT did I do?", "WHY did I do it?", "HOW does it work?", "WHAT does the bank gain?"]

# ══════════════════════════════════════════════════════════════════════
#  CONTENT: ALL CHAPTER MATRICES
# ══════════════════════════════════════════════════════════════════════

CH0_MATRIX = {
    (0, 0): (
        "Applied SHAP (SHapley Additive exPlanations) to decompose the XGBoost "
        "fraud model's predictions into feature-level contributions. Computed exact "
        "Shapley values via TreeExplainer for 118,108 test transactions (7 features). "
        "Produced global explainability (summary plots, feature importance, dependence "
        "plots), local explainability (6 case studies with waterfall plots and plain-"
        "English explanations), business insights (fraud vs legitimate comparison, "
        "risk tier decomposition), and regulatory compliance documentation (SR 11-7, "
        "fair lending review, right-to-explanation, audit trail)."
    ),
    (0, 1): (
        "Made the fraud detection model transparent and explainable. For every "
        "transaction, the system can now answer 'Why was this flagged?' with "
        "specific feature contributions. Documented 6 representative case studies "
        "covering true positives, false negatives, false positives, and borderline "
        "cases. Prepared regulatory documentation for SR 11-7 model risk management, "
        "fair lending compliance, and customer dispute resolution."
    ),
    (0, 2): (
        "We opened the 'black box' of the fraud detection system and made it explain "
        "itself. For any transaction, it can now say: 'I flagged this because of 5 "
        "purchases in one hour, an unusual amount, and a 3 AM timing.' We also "
        "prepared paperwork for bank regulators and created a process for when "
        "honest customers are accidentally blocked."
    ),
    (1, 0): (
        "A production fraud model requires explainability for three audiences: "
        "analysts (understand the review queue), regulators (verify model governance "
        "per SR 11-7), and customers (right to explanation under GDPR Art. 22). "
        "SHAP provides mathematically exact, game-theoretic feature attributions "
        "that satisfy all three requirements. The 6 case studies demonstrate the "
        "model's strengths and weaknesses across decision boundaries."
    ),
    (1, 1): (
        "A fraud model that says 'blocked' without explanation creates customer "
        "complaints, regulatory risk, and analyst frustration. Banks must prove "
        "their models are fair, documented, and explainable. SHAP provides the "
        "evidence trail that connects each decision to specific transaction "
        "characteristics, enabling dispute resolution, audit compliance, and "
        "continuous model improvement."
    ),
    (1, 2): (
        "Imagine a security guard who stops someone but can't explain why. That "
        "makes customers angry and gets the bank in trouble with inspectors. "
        "We taught the system to explain every decision: 'I stopped this because "
        "of X, Y, and Z.' Now customers, regulators, and the fraud team all "
        "understand why each decision was made."
    ),
    (2, 0): (
        "SHAP TreeExplainer computes exact Shapley values for tree-based models "
        "in O(TLD\u00b2) time (T=trees, L=leaves, D=depth). Base value = 0.0178 "
        "(average model output). Each feature's SHAP value represents its marginal "
        "contribution to the prediction. Global analysis: mean |SHAP| ranking, "
        "beeswarm plots, dependence plots. Local analysis: waterfall decomposition "
        "per transaction. Full test set (118,108 rows) computed in seconds."
    ),
    (2, 1): (
        "For each transaction, SHAP assigns a contribution score to each of the "
        "7 features. Positive contributions push toward fraud; negative push toward "
        "legitimate. The base value (1.78% \u2014 the average fraud rate the model "
        "expects before seeing any features) is adjusted up or down by each feature. "
        "This creates a complete, auditable explanation for every decision."
    ),
    (2, 2): (
        "Think of a jury trial: each piece of evidence (feature) either points "
        "toward 'guilty' (fraud) or 'innocent' (legitimate). SHAP is like a "
        "judge who writes down exactly how much each piece of evidence influenced "
        "the verdict. The starting point is 'probably innocent' (only 1.78% are "
        "fraud), and each feature adjusts this up or down."
    ),
    (3, 0): (
        "Complete explainability stack: global feature importance ranking "
        "(TransactionAmt > txn_count_24hr > amount_deviation > hour_of_day), "
        "local per-transaction explanations, fraud-vs-legitimate SHAP comparison, "
        "risk tier decomposition. SR 11-7 documentation with 8/12 governance "
        "items completed. 6 SHAP visualization artifacts saved for dashboard "
        "integration (Phase 5)."
    ),
    (3, 1): (
        "The bank gets a fully explainable fraud system. Every blocked transaction "
        "has a documented reason. Regulators see complete model governance. "
        "Customer disputes can be resolved with clear, specific explanations. "
        "The fraud team understands which patterns the model detects and where "
        "it has blind spots, enabling targeted improvements."
    ),
    (3, 2): (
        "The bank can now explain every decision to anyone who asks. Regulators "
        "get the paperwork they need. Customers get clear answers when their "
        "purchase is blocked. And the fraud team knows exactly what the system "
        "is good at and where it needs improvement."
    ),
}

CH1_MATRIX = {
    (0, 0): (
        "Loaded XGBoost model (xgboost_final.pkl), StandardScaler (scaler.pkl), and "
        "threshold configuration (threshold_config.pkl) from Phase 3. Loaded test set "
        "(118,108 transactions, 3.44% fraud rate, 4,064 frauds). Extracted production "
        "thresholds: auto-block \u22650.90, manual review \u22650.41. Generated fraud "
        "probability scores for all test transactions. Initialized SHAP TreeExplainer."
    ),
    (0, 1): (
        "Loaded the trained fraud detection model and the test dataset (most recent "
        "transactions, never seen during training). Set up the explanation engine "
        "(SHAP) that will decompose every prediction into feature-level contributions. "
        "Confirmed production thresholds: auto-block at 90%+ confidence, manual review "
        "at 41%+ confidence."
    ),
    (0, 2): (
        "We loaded the trained fraud detector and a set of 118,000 transactions it "
        "has never seen before. We also loaded the 'explanation engine' that will "
        "tell us why the detector makes each decision. The system has two alarm "
        "levels: 'definitely stop this' (90%+ suspicion) and 'check this' (41%+ "
        "suspicion)."
    ),
    (1, 0): (
        "The test set provides unbiased performance estimates (temporal split ensures "
        "no data leakage). TreeExplainer is chosen over KernelExplainer because it "
        "computes exact (not approximate) Shapley values for tree models. The "
        "threshold_config.pkl ensures the same thresholds used in model selection "
        "are applied consistently in the explainability analysis."
    ),
    (1, 1): (
        "Using the test set (data the model never saw during training) ensures honest "
        "evaluation. The SHAP engine is exact for tree models \u2014 no approximation "
        "or sampling errors. Loading the same thresholds from Phase 3 ensures "
        "consistency: the explanations match the actual production decisions."
    ),
    (1, 2): (
        "We use fresh data the system has never seen to test honestly. The explanation "
        "engine gives exact answers (no guessing). And we use the same alarm settings "
        "from the training phase so everything is consistent."
    ),
    (2, 0): (
        "joblib.load() deserializes model artifacts. model.predict_proba(X)[:,1] "
        "generates fraud probabilities. X_test cleaned: inf\u219210, -inf\u2192-10, "
        "NaN\u21920 (same as notebook 03). shap.TreeExplainer(model) initializes "
        "the exact explainer. explainer.expected_value = 0.0178 (base prediction "
        "before any features). 7 features: txn_count_1hr, txn_count_24hr, "
        "amount_deviation, is_first_transaction, hour_of_day, is_weekend, TransactionAmt."
    ),
    (2, 1): (
        "The saved model files are loaded, and fraud scores are generated for every "
        "test transaction. The SHAP explainer is initialized with a base value of "
        "1.78% \u2014 the model's starting prediction before looking at any features. "
        "Each feature then adjusts this base value up or down."
    ),
    (2, 2): (
        "We opened the saved model files and asked it to score all 118,000 test "
        "transactions. The explanation engine starts with a baseline of 'about "
        "1.78% chance of fraud' and then adjusts based on each transaction's "
        "characteristics."
    ),
    (3, 0): (
        "Reproducible pipeline: model + scaler + thresholds loaded from serialized "
        "artifacts. Test set fraud score distribution: mean 0.4030, median 0.3913. "
        "23 transactions exceed auto-block (0.90), 54,249 exceed manual review (0.41). "
        "TreeExplainer ready for both global (sample-based) and local (full-set) analysis."
    ),
    (3, 1): (
        "The bank has a reproducible, auditable pipeline: same model, same thresholds, "
        "same data cleaning. Of 118,108 test transactions, 23 would be auto-blocked and "
        "54,249 would go to manual review. The explanation system is ready to decompose "
        "every single decision."
    ),
    (3, 2): (
        "Everything is set up consistently and reproducibly. Out of 118,000 transactions, "
        "23 would be automatically stopped and about 54,000 would need a human checker. "
        "The explanation system is ready to tell us why."
    ),
}

CH2_MATRIX = {
    (0, 0): (
        "Computed SHAP values for a 2,000-transaction sample (global analysis). "
        "Generated beeswarm summary plot showing each feature's value-to-impact "
        "relationship. Created mean |SHAP| bar chart ranking feature importance: "
        "TransactionAmt (highest), txn_count_24hr, amount_deviation, hour_of_day, "
        "txn_count_1hr, is_first_transaction, is_weekend (lowest). Used banking-"
        "friendly labels (e.g., 'Transaction Velocity (1 hour)')."
    ),
    (0, 1): (
        "Identified which features matter most for fraud detection across all "
        "transactions. Transaction amount and 24-hour velocity are the top drivers. "
        "The beeswarm plot shows exactly how each feature's value pushes the fraud "
        "score up or down. Feature labels were translated into business-friendly "
        "language for stakeholder presentations."
    ),
    (0, 2): (
        "We found out which clues matter most overall. The transaction amount and "
        "how many purchases someone makes in 24 hours are the biggest factors. "
        "We created two charts: one showing how each clue's value affects the "
        "fraud score (like a scatter plot), and one ranking them from most to "
        "least important."
    ),
    (1, 0): (
        "Global feature importance answers 'What does the model pay attention to?' "
        "A 2,000-transaction sample is sufficient for stable SHAP estimates with "
        "7 features (standard error < 0.01). The beeswarm plot reveals non-linear "
        "relationships (e.g., TransactionAmt has bimodal SHAP: both very high and "
        "very low amounts increase fraud risk). Mean |SHAP| provides a single "
        "importance number per feature for executive reporting."
    ),
    (1, 1): (
        "Knowing which features drive the model's decisions helps the bank "
        "understand its fraud detection strategy. If the model relies heavily on "
        "velocity, the bank knows that rapid-fire transactions are the primary "
        "detection mechanism. This guides where to invest in additional data "
        "sources (e.g., merchant categories to catch velocity-evasion fraud)."
    ),
    (1, 2): (
        "It's like asking a detective: 'What clues do you look at most?' "
        "If the detective says 'speed of purchases,' the bank knows that's the "
        "main defense. If fraudsters learn to shop slowly, the bank needs new "
        "clues \u2014 and now they know where to look."
    ),
    (2, 0): (
        "shap.TreeExplainer(model).shap_values(X_sample) returns an (n, 7) array. "
        "np.abs(shap_values).mean(axis=0) computes mean absolute SHAP per feature. "
        "shap.summary_plot() generates beeswarm: x=SHAP value, color=feature value, "
        "y=feature rank. Banking-friendly labels via FEATURE_LABELS dict mapping "
        "technical names to business names (e.g., txn_count_1hr \u2192 'Transaction "
        "Velocity (1 hour)')."
    ),
    (2, 1): (
        "For 2,000 transactions, each feature gets a contribution score. The absolute "
        "values are averaged to create a single importance ranking. The beeswarm plot "
        "uses color (red=high value, blue=low value) to show how feature values relate "
        "to their impact on the fraud score. Technical feature names are translated "
        "into banking language."
    ),
    (2, 2): (
        "We scored 2,000 transactions, measured how much each clue mattered, and "
        "averaged the results. The beeswarm chart is like a weather map: each dot "
        "is one transaction, its color shows the clue's value, and its position "
        "shows whether that clue pointed toward fraud (right) or safety (left)."
    ),
    (3, 0): (
        "Feature importance ranking confirmed: TransactionAmt and velocity features "
        "dominate. This differs from LR coefficients (where txn_count_1hr was #1), "
        "confirming non-linear interactions captured by XGBoost. The beeswarm plot "
        "reveals that both very high AND very low TransactionAmt increase fraud "
        "risk \u2014 a pattern invisible to linear models. Visualizations saved for "
        "Phase 5 dashboard integration."
    ),
    (3, 1): (
        "The bank learns that transaction amount and velocity are the primary fraud "
        "signals. The discovery that both very high and very low amounts increase "
        "risk guides fraud analyst training: watch for micro-transactions (card "
        "testing) AND large purchases (account takeover). The ranking is documented "
        "for regulatory audit trails."
    ),
    (3, 2): (
        "The bank now knows its top fraud clues: how much money and how fast. "
        "Interestingly, both very big AND very small purchases are suspicious \u2014 "
        "thieves test stolen cards with tiny purchases, then make big ones. "
        "The fraud team can use this knowledge to spot fraud patterns."
    ),
}

CH3_MATRIX = {
    (0, 0): (
        "Generated SHAP dependence plots for top 4 features: TransactionAmt, "
        "txn_count_24hr, amount_deviation, hour_of_day. Each plot shows feature "
        "value (x-axis) vs SHAP contribution (y-axis) with automatic interaction "
        "coloring. Produced global insights summary: HIGH values of TransactionAmt, "
        "txn_count_24hr, amount_deviation, and hour_of_day all INCREASE fraud risk. "
        "Documented operational takeaways for fraud analysts."
    ),
    (0, 1): (
        "Created detailed charts showing exactly how each feature's value affects "
        "fraud risk. For example: at what transaction count per hour does the model "
        "get concerned? How much spending deviation triggers a flag? Which hours are "
        "riskier? Translated these patterns into actionable intelligence for fraud "
        "analysts: look for rapid bursts, unusual amounts, and overnight activity."
    ),
    (0, 2): (
        "We made charts that show exactly when each clue raises suspicion. For "
        "example: more than a few purchases per hour makes the system nervous. "
        "Spending way more than usual is a red flag. Late-night shopping is "
        "suspicious. We wrote this up as instructions for the fraud team."
    ),
    (1, 0): (
        "Dependence plots reveal the functional form of each feature's effect on "
        "the prediction. Unlike feature importance (which gives a single number), "
        "dependence plots show the full relationship: thresholds, non-linearities, "
        "and interaction effects. The automatic SHAP interaction coloring reveals "
        "which features amplify each other (e.g., high velocity + early morning)."
    ),
    (1, 1): (
        "Feature importance says 'velocity matters,' but dependence plots say "
        "'velocity matters A LOT after 5 transactions per hour, and even more "
        "when combined with unusual amounts.' This level of detail helps analysts "
        "understand exactly when the model becomes concerned and why."
    ),
    (1, 2): (
        "Instead of just knowing which clues matter, we now know exactly how much "
        "of each clue is needed to raise the alarm. It's like knowing that one "
        "cookie from the jar is fine, but five cookies in an hour is suspicious."
    ),
    (2, 0): (
        "shap.dependence_plot(feat_idx, shap_values, X_sample) creates scatter "
        "plots with automatic interaction feature coloring. For each feature, "
        "computed median split: avg SHAP for high vs low values determines "
        "directionality. FEATURE_INSIGHTS dict maps each feature to a banking-"
        "friendly pattern description. 2x2 subplot grid for the top 4 features."
    ),
    (2, 1): (
        "For each top feature, a scatter chart plots the feature's actual value "
        "against its impact on the fraud score. Points are colored by the strongest "
        "interacting feature. The median split analysis determines whether high or "
        "low values increase fraud risk \u2014 confirmed: all top features show "
        "'high values = higher risk.'"
    ),
    (2, 2): (
        "We plotted each clue's value against its suspicion impact. The dots "
        "are colored by a related clue to show how they work together. We also "
        "calculated which direction matters: for all top clues, 'more is more "
        "suspicious.'"
    ),
    (3, 0): (
        "Non-linear feature effects documented: TransactionAmt shows bimodal risk "
        "(high AND low amounts), txn_count_24hr shows a step function above ~5 "
        "transactions, amount_deviation is approximately linear, hour_of_day shows "
        "elevated risk at hours 6-9 (early morning). These functional forms guide "
        "future feature engineering and rule-based fallback systems."
    ),
    (3, 1): (
        "The bank gets a detailed playbook: fraud risk increases sharply when "
        "24-hour velocity exceeds ~5 transactions, when amounts deviate "
        "significantly from history, and during early morning hours (6-9 AM). "
        "Analysts can use these thresholds as mental shortcuts when reviewing "
        "flagged transactions."
    ),
    (3, 2): (
        "The bank now has specific 'danger zones': more than 5 purchases a day, "
        "spending much more than usual, and shopping between 6 and 9 AM. "
        "These are the situations that make the fraud detector most nervous."
    ),
}

CH4_MATRIX = {
    (0, 0): (
        "Computed full SHAP values for all 118,108 test transactions. Selected 6 "
        "representative case studies: (1) True Positive \u2014 clear fraud (score 0.9094), "
        "(2) True Positive \u2014 velocity-driven (score 0.7332), (3) False Negative \u2014 "
        "missed fraud (score 0.0853), (4) False Positive \u2014 legitimate flagged "
        "(score 0.9342), (5) Auto-Block candidate (score 0.9342), (6) Borderline "
        "case (score 0.3648). Generated plain-English explanations and waterfall "
        "plots for each case."
    ),
    (0, 1): (
        "Demonstrated the system's ability to explain individual transactions \u2014 "
        "critical for fraud analysts, customer disputes, and regulatory audits. "
        "Six cases cover every important scenario: correctly caught fraud, missed "
        "fraud (model limitation), false alarms (customer impact), and borderline "
        "decisions (analyst judgment needed). Each explanation uses plain language "
        "with risk-level descriptions."
    ),
    (0, 2): (
        "We picked 6 example transactions and asked the system to explain each "
        "decision in plain English. We covered the good (fraud caught), the bad "
        "(fraud missed), the annoying (honest customer blocked), and the tricky "
        "(transactions right on the border). Each explanation says exactly which "
        "clues mattered and how much."
    ),
    (1, 0): (
        "Local explainability is required by GDPR Art. 22 (right to explanation "
        "for automated decisions) and supports SR 11-7 model governance. The 6 case "
        "studies are strategically chosen to cover the confusion matrix quadrants "
        "(TP, FP, FN) plus operational scenarios (auto-block, borderline). Waterfall "
        "plots show the cumulative feature contribution from base value (0.0178) to "
        "the final fraud score."
    ),
    (1, 1): (
        "When a customer calls to ask why their purchase was blocked, the bank "
        "needs a clear answer. When a regulator asks how individual decisions are "
        "made, the bank needs documentation. These 6 cases prove the system can "
        "explain any decision, covering every possible outcome \u2014 from clear "
        "fraud to honest mistakes."
    ),
    (1, 2): (
        "When a customer calls the bank saying 'Why was my purchase blocked?', "
        "the system can now explain: 'Your purchase looked suspicious because "
        "you made 5 purchases in one hour, spent more than usual, and it was "
        "3 AM.' This makes customers feel heard and helps the bank comply "
        "with regulations."
    ),
    (2, 0): (
        "Full SHAP values: explainer.shap_values(X_test) returns (118108, 7) array. "
        "Case selection: score-based filtering (e.g., TP: isFraud=1 AND score>=0.80, "
        "FN: isFraud=1 AND score<0.41). explain_transaction() helper translates "
        "SHAP values into risk language: |SHAP|<0.05='Minimal', 0.05-0.2='Moderate', "
        "0.2-0.5='Strong', >0.5='VERY STRONG'. shap.Explanation() objects feed "
        "waterfall plots. 3x2 subplot grid for all 6 cases."
    ),
    (2, 1): (
        "For each case, the system ranks all 7 features by their impact, translates "
        "technical SHAP values into risk language ('STRONG \u2014 INCREASES fraud risk'), "
        "and formats feature values in human-readable form ($17.52, 6:00 AM, etc.). "
        "Waterfall plots show the step-by-step build-up from baseline (1.78%) to the "
        "final fraud score."
    ),
    (2, 2): (
        "For each example, we listed the clues from most important to least. Each "
        "clue gets a 'strength rating' (from 'minimal impact' to 'VERY STRONG'). "
        "The waterfall chart shows how each clue adds to or subtracts from the "
        "suspicion score, like a running tally."
    ),
    (3, 0): (
        "Case studies reveal model behavior across decision boundaries. Key findings: "
        "FN cases lack velocity signal (zero txn counts), suggesting velocity is "
        "necessary for detection. FP cases accumulate multiple moderate signals. "
        "Auto-block cases show extreme multi-feature contributions. The plain-English "
        "explanation template is production-ready for customer communication."
    ),
    (3, 1): (
        "The bank gets a proven explanation system covering every scenario. Missed "
        "frauds teach the team where to improve (add merchant data). False alarms "
        "have documented reasons for customer communication. The explanation template "
        "is ready for production: analysts can generate plain-English explanations "
        "for any flagged transaction."
    ),
    (3, 2): (
        "The bank can now explain every decision clearly. When fraud is missed, "
        "the team knows why (the thief was too careful). When honest customers "
        "are blocked, the team can apologize and explain what triggered the alarm. "
        "This builds trust with customers and regulators."
    ),
}

CH5_MATRIX = {
    (0, 0): (
        "Computed mean SHAP per feature for fraud vs legitimate transactions across "
        "the full test set. Generated grouped bar chart showing feature impact "
        "differences. Created risk tier decomposition: mean |SHAP| per feature for "
        "auto-approve, manual review, and auto-block tiers. Produced 5-point "
        "actionable insights report for fraud operations (queue prioritization, "
        "FP reduction, FN patterns, auto-block validation, threshold monitoring)."
    ),
    (0, 1): (
        "Analyzed how the model distinguishes fraud from legitimate transactions "
        "at the feature level. Created a comparison showing which features push "
        "toward fraud vs toward legitimate. Decomposed feature contributions by "
        "risk tier (auto-approve, manual review, auto-block). Produced actionable "
        "recommendations for fraud operations: how to prioritize the review queue, "
        "reduce false alarms, and address missed fraud patterns."
    ),
    (0, 2): (
        "We compared how the system sees fraud vs honest transactions. The biggest "
        "differences are in how fast someone shops and how much they spend compared "
        "to their history. We also analyzed what makes the system say 'stop,' "
        "'check,' or 'go,' and wrote recommendations for the fraud team."
    ),
    (1, 0): (
        "The fraud-vs-legitimate SHAP comparison reveals the model's discrimination "
        "strategy. If features have similar SHAP distributions for both classes, "
        "the model struggles to separate them \u2014 this identifies improvement "
        "targets. The risk tier decomposition shows whether the three-tier strategy "
        "creates qualitatively different feature profiles, validating the multi-"
        "threshold design from Phase 3."
    ),
    (1, 1): (
        "Understanding HOW the model separates fraud from legitimate transactions "
        "informs operational strategy. If the key distinction is velocity, the bank "
        "knows fraudsters who pace their transactions are a blind spot. The tier "
        "analysis validates that auto-block, review, and approve are qualitatively "
        "different \u2014 not just arbitrary score cutoffs."
    ),
    (1, 2): (
        "We wanted to know: what makes the system think something is fraud vs honest? "
        "The answer: speed and unusual spending. We also checked that the three "
        "alarm levels (stop, check, go) actually look at different patterns, not "
        "just arbitrary lines."
    ),
    (2, 0): (
        "Fraud mask: y_test == 1, legit mask: y_test == 0. Mean SHAP computed per "
        "class: shap_values_full[mask].mean(axis=0). Grouped bar chart with #e74c3c "
        "(fraud) and #2ecc71 (legitimate). Risk tiers: pd.cut(fraud_scores, bins) "
        "with labels. Per-tier mean |SHAP|: np.abs(shap_values_full[tier_mask])."
        "mean(axis=0). Feature contribution stacked bar chart by tier."
    ),
    (2, 1): (
        "For fraud and legitimate groups separately, the average SHAP contribution "
        "of each feature is computed. The grouped bar chart places them side by side "
        "for easy comparison. The risk tier analysis groups all transactions into "
        "three buckets (auto-approve, review, auto-block) and computes which "
        "features contribute most in each bucket."
    ),
    (2, 2): (
        "We calculated the average impact of each clue for fraud transactions and "
        "for honest transactions, then compared them side by side. We also split "
        "all transactions into three groups (go, check, stop) and looked at which "
        "clues matter most in each group."
    ),
    (3, 0): (
        "Key discrimination features: TransactionAmt and velocity show the largest "
        "SHAP gap between fraud and legitimate. Auto-block tier shows 3-5x higher "
        "mean |SHAP| than auto-approve, confirming the tiers represent qualitatively "
        "different risk profiles. Actionable insights document 5 operational "
        "recommendations: queue prioritization, FP reduction via merchant checks, "
        "FN mitigation via additional features, monthly auto-block audits, quarterly "
        "threshold recalibration."
    ),
    (3, 1): (
        "The bank gets a clear picture of its detection strategy and its limits. "
        "The 5 actionable insights provide immediate operational improvements: "
        "prioritize multi-signal alerts, reduce false alarms with merchant checks, "
        "address missed fraud by adding new data sources. The tier analysis confirms "
        "the three-tier strategy is well-designed."
    ),
    (3, 2): (
        "The bank learns its strengths (catching fast spenders) and weaknesses "
        "(missing slow, careful thieves). The five recommendations tell the fraud "
        "team exactly what to do: focus on the most suspicious alerts first, "
        "check with stores to reduce false alarms, and add new clues to catch "
        "the sneaky thieves."
    ),
}

CH6_MATRIX = {
    (0, 0): (
        "Produced comprehensive regulatory documentation: (1) SR 11-7 model "
        "documentation (identification, inputs, outputs, assumptions, limitations), "
        "(2) Fair lending feature review (5 features assessed, 2 medium-risk: "
        "is_first_transaction, hour_of_day/is_weekend), (3) Right-to-explanation "
        "capability (feature attribution, quantified contribution, baseline comparison), "
        "(4) Audit trail specification (fields, retention, dispute workflow), "
        "(5) Governance checklist (8/12 items complete, 4 pending). "
        "Documented 6 model limitations with compensating controls."
    ),
    (0, 1): (
        "Created the complete regulatory package for the fraud model. Documented "
        "everything regulators need: what the model is, how it works, where it's "
        "limited, and how it's governed. Reviewed all features for potential "
        "discrimination against protected groups. Established a dispute resolution "
        "workflow and audit trail. Identified 4 remaining governance items that "
        "require production data to complete."
    ),
    (0, 2): (
        "We prepared all the paperwork the bank inspectors need. We wrote down "
        "what the system does, how it works, and what it can't do. We checked "
        "that the system doesn't discriminate against any group of people. We "
        "created a process for when customers complain about blocked purchases. "
        "And we made a checklist to keep the system well-maintained."
    ),
    (1, 0): (
        "SR 11-7 and OCC 2011-12 require comprehensive model documentation before "
        "production deployment. ECOA/Fair Lending requires proxy discrimination "
        "analysis even when no protected attributes are direct inputs. GDPR Art. 22 "
        "mandates right to explanation for automated decisions affecting customers. "
        "Proactive documentation reduces regulatory risk and accelerates the "
        "approval process."
    ),
    (1, 1): (
        "Banks cannot deploy fraud models without regulatory approval. The "
        "documentation package addresses the three main regulatory requirements: "
        "(1) model risk management (SR 11-7), (2) fair lending (no discrimination), "
        "and (3) customer rights (explanation on request). Completing this "
        "documentation proactively demonstrates good governance practices."
    ),
    (1, 2): (
        "Bank inspectors need to see that the system is fair, documented, and "
        "well-maintained. We created all the paperwork before they even ask \u2014 "
        "like doing your homework before the teacher checks. This shows the bank "
        "takes fraud detection seriously and responsibly."
    ),
    (2, 0): (
        "SR 11-7 documentation: model identification (name, version, type, purpose, "
        "owner, date), inputs (7 features with descriptions), outputs (probability "
        "score + 3-tier decision), assumptions (5 key assumptions), performance "
        "metrics (PR-AUC 0.1098, recall ~74-76%, precision ~5-6%). Fair lending: "
        "per-feature risk assessment (LOW to MEDIUM). Audit trail: transaction_id, "
        "fraud_score, threshold, decision, SHAP values, model_version, timestamp. "
        "Retention: 7 years."
    ),
    (2, 1): (
        "The regulatory package follows standard banking templates. Model card "
        "documents what goes in and what comes out. Fair lending review assesses "
        "each feature for potential proxy discrimination (2 features flagged as "
        "MEDIUM risk: first-time transactions and time-of-day). Audit trail "
        "specification ensures every decision can be reconstructed and explained "
        "years later."
    ),
    (2, 2): (
        "We filled out all the standard forms: what the system is called, what "
        "it does, what goes in, what comes out, and what it can't do. We checked "
        "each clue for fairness. We specified exactly what to record for each "
        "decision so the bank can look it up years later."
    ),
    (3, 0): (
        "Regulatory readiness: 8/12 governance items completed (model documentation, "
        "performance metrics, global + local explainability, limitations, fair lending "
        "review, right-to-explanation, audit trail spec). 4 items pending production "
        "data: disparate impact testing, champion/challenger, monitoring dashboard, "
        "revalidation schedule. 6 limitations documented with compensating controls. "
        "Monitoring schedule: daily/weekly/monthly/quarterly/annual cadence."
    ),
    (3, 1): (
        "The bank is 67% through its regulatory checklist (8 of 12 items). The 4 "
        "remaining items require production data or operational setup that can't be "
        "completed during development. The documentation package is ready for initial "
        "regulatory review and production deployment with appropriate monitoring. "
        "A 5-tier monitoring schedule ensures ongoing compliance."
    ),
    (3, 2): (
        "The bank has completed most of its regulatory homework (8 out of 12 items). "
        "The remaining 4 items need real-world data that isn't available yet. "
        "The system is ready to go live with a monitoring plan to keep checking "
        "that everything works properly: daily, weekly, monthly, quarterly, and "
        "yearly reviews."
    ),
}

ALL_CHAPTERS = [
    {
        "number": 0,
        "title": "Executive Overview",
        "subtitle": "Full-Project Summary",
        "narrative": (
            "This notebook makes the XGBoost fraud detection model from Phase 3 fully "
            "transparent and explainable using SHAP (SHapley Additive exPlanations). "
            "SHAP decomposes every prediction into exact feature-level contributions, "
            "answering the critical question: 'Why did the model flag or approve this "
            "transaction?' The analysis covers global explainability (which features "
            "matter most overall), local explainability (why specific transactions were "
            "flagged), business insights (fraud vs legitimate patterns, risk tier analysis), "
            "and regulatory compliance (SR 11-7, fair lending, right-to-explanation). "
            "Six SHAP visualizations and 6 case studies provide the evidence base for "
            "Phase 5 dashboard deployment."
        ),
        "matrix": CH0_MATRIX,
        "figures": [],
        "callouts": [
            ("insight",
             "SHAP provides mathematically exact, game-theoretic explanations for every "
             "prediction. Unlike approximation-based methods, TreeExplainer computes the "
             "true Shapley values for tree models \u2014 no sampling error, no approximation."),
        ],
    },
    {
        "number": 1,
        "title": "Setup & Model Loading",
        "subtitle": "Notebook Section 1: Libraries, Artifacts, Test Data, SHAP Initialization",
        "narrative": (
            "The pipeline begins by loading the serialized model artifacts from Phase 3: "
            "the tuned XGBoost model, StandardScaler, and production threshold configuration "
            "(auto-block \u22650.90, manual review \u22650.41). The test set (118,108 "
            "transactions, 3.44% fraud rate) is loaded with the same data cleaning applied "
            "in Phase 3 (inf\u219210, NaN\u21920). Fraud probability scores are generated "
            "for all test transactions, and the SHAP TreeExplainer is initialized with "
            "a base value of 0.0178."
        ),
        "matrix": CH1_MATRIX,
        "figures": [],
        "callouts": [
            ("business",
             "The base value (0.0178 = 1.78%) represents the model's prediction before "
             "seeing any features \u2014 essentially the background fraud rate. Each feature "
             "then adjusts this baseline up or down. This is the starting point for every "
             "explanation."),
        ],
    },
    {
        "number": 2,
        "title": "Global Explainability \u2014 Feature Importance",
        "subtitle": "Notebook Section 2 (Parts 2.1-2.2): SHAP Summary Plot, Feature Importance Ranking",
        "narrative": (
            "Global explainability answers: 'Across all transactions, which features "
            "matter most?' A 2,000-transaction sample is analyzed with the beeswarm "
            "summary plot (showing value-to-impact relationships) and the mean |SHAP| "
            "bar chart (ranking overall importance). TransactionAmt emerges as the top "
            "feature, followed by txn_count_24hr and amount_deviation. A key discovery: "
            "TransactionAmt has bimodal risk \u2014 both very high AND very low amounts "
            "increase fraud risk (card testing vs account takeover)."
        ),
        "matrix": CH2_MATRIX,
        "figures": [
            ("shap_summary_beeswarm.png",
             "Figure 1: SHAP Beeswarm \u2014 Each dot is a transaction; color = feature value; "
             "position = impact on fraud score."),
            ("shap_feature_importance_bar.png",
             "Figure 2: Mean |SHAP| Feature Importance \u2014 Overall ranking of the 7 model features."),
        ],
        "callouts": [
            ("insight",
             "TransactionAmt shows bimodal risk: both very high and very low amounts "
             "increase fraud risk. This pattern is invisible to linear models (like "
             "Logistic Regression) and explains part of XGBoost's 33.8% improvement."),
        ],
    },
    {
        "number": 3,
        "title": "Global Explainability \u2014 Feature Dependence",
        "subtitle": "Notebook Section 2 (Parts 2.3-2.4): Dependence Plots, Analyst Insights",
        "narrative": (
            "Dependence plots reveal the functional relationship between each feature's "
            "value and its SHAP contribution, going beyond simple importance rankings. "
            "Key patterns: txn_count_24hr shows a step function above ~5 transactions, "
            "amount_deviation is approximately linear, and hour_of_day shows elevated "
            "risk at 6-9 AM. Global insights are translated into operational takeaways "
            "for fraud analysts, identifying the specific patterns that align with known "
            "fraud typologies (card testing, account takeover, new account fraud)."
        ),
        "matrix": CH3_MATRIX,
        "figures": [],
        "callouts": [
            ("business",
             "Operational takeaway for fraud analysts: prioritize transactions with "
             "rapid bursts (>5 per day), amounts deviating significantly from client "
             "history, and overnight timing (especially 6-9 AM). These patterns align "
             "with card testing, account takeover, and new account fraud."),
        ],
    },
    {
        "number": 4,
        "title": "Local Explainability \u2014 Case Studies",
        "subtitle": "Notebook Section 3: 6 Case Studies, Plain-English Explanations, Waterfall Plots",
        "narrative": (
            "Local explainability demonstrates the system's ability to explain individual "
            "transactions. SHAP values are computed for all 118,108 test transactions, "
            "and 6 representative cases are selected: (1) clear true positive (score 0.91), "
            "(2) velocity-driven true positive (score 0.73), (3) missed fraud / false "
            "negative (score 0.09), (4) false positive (score 0.93), (5) auto-block "
            "candidate (score 0.93), (6) borderline case (score 0.36). Each case receives "
            "a plain-English explanation with risk-level descriptions (Minimal/Moderate/"
            "Strong/VERY STRONG) and a waterfall plot showing cumulative feature contributions."
        ),
        "matrix": CH4_MATRIX,
        "figures": [
            ("shap_waterfall_cases.png",
             "Figure 3: SHAP Waterfall Plots \u2014 Step-by-step decomposition of 6 case study decisions."),
        ],
        "callouts": [
            ("caution",
             "Case 3 (missed fraud, score 0.09) reveals a model limitation: when all "
             "velocity features are zero and the amount is normal, the model has no "
             "signal to detect fraud. Sophisticated fraudsters who pace their transactions "
             "can evade velocity-based detection. Future improvement: add merchant-category "
             "and device-fingerprinting features."),
        ],
    },
    {
        "number": 5,
        "title": "Business Insights & Operational Intelligence",
        "subtitle": "Notebook Section 4: Fraud vs Legitimate Analysis, Risk Tier Decomposition, Actionable Insights",
        "narrative": (
            "Business-focused analysis translates SHAP findings into operational "
            "intelligence. The fraud-vs-legitimate comparison reveals that TransactionAmt "
            "and velocity features show the largest SHAP gap between the two groups. "
            "Risk tier decomposition confirms that auto-block, manual review, and auto-approve "
            "represent qualitatively different risk profiles (auto-block shows 3-5x higher "
            "mean |SHAP|). Five actionable insights guide fraud operations: queue "
            "prioritization, false positive reduction, false negative patterns, "
            "auto-block validation, and threshold monitoring."
        ),
        "matrix": CH5_MATRIX,
        "figures": [
            ("shap_fraud_vs_legit.png",
             "Figure 4: Fraud vs Legitimate SHAP Comparison \u2014 How the model distinguishes "
             "the two groups at the feature level."),
            ("shap_risk_tiers.png",
             "Figure 5: Feature Contribution by Risk Tier \u2014 What drives each decision level "
             "(auto-approve, manual review, auto-block)."),
        ],
        "callouts": [
            ("business",
             "Five actionable insights for fraud operations: (1) Prioritize multi-signal "
             "alerts in the review queue, (2) Reduce false positives with merchant consistency "
             "checks, (3) Address missed fraud by adding merchant and geolocation features, "
             "(4) Audit auto-block decisions monthly, (5) Recalibrate thresholds quarterly."),
        ],
    },
    {
        "number": 6,
        "title": "Regulatory Compliance & Model Governance",
        "subtitle": "Notebook Section 5: SR 11-7 Documentation, Fair Lending, Right-to-Explanation, Audit Trail",
        "narrative": (
            "The regulatory section addresses three critical compliance frameworks: "
            "SR 11-7 / OCC 2011-12 (model risk management), ECOA / Fair Lending "
            "(anti-discrimination), and GDPR Art. 22 (right to explanation). Model "
            "documentation covers identification, inputs, outputs, assumptions, "
            "performance metrics, and 6 documented limitations with compensating "
            "controls. Fair lending review assesses all 7 features (2 flagged as "
            "MEDIUM risk). Right-to-explanation capability is demonstrated with "
            "feature attribution, quantified contribution, and baseline comparison. "
            "A governance checklist shows 8/12 items complete with 4 pending "
            "production deployment."
        ),
        "matrix": CH6_MATRIX,
        "figures": [],
        "callouts": [
            ("caution",
             "Two features carry MEDIUM fair lending risk: is_first_transaction "
             "(new customers disproportionately flagged) and hour_of_day/is_weekend "
             "(shift workers and different time zones affected). Disparate impact "
             "testing should be conducted when demographic data becomes available."),
            ("business",
             "The model governance checklist is 67% complete (8/12 items). The 4 "
             "remaining items (disparate impact testing, champion/challenger framework, "
             "monitoring dashboard, revalidation schedule) require production data or "
             "operational infrastructure that cannot be completed during development."),
        ],
    },
]

SUMMARY_TABLE_DATA = [
    ("Explainability Method", "SHAP (SHapley Additive exPlanations)"),
    ("SHAP Algorithm", "TreeExplainer (exact, not approximate)"),
    ("Model Explained", "XGBoost (tuned, 7 features)"),
    ("Test Set", "118,108 transactions (3.44% fraud)"),
    ("Base Value", "0.0178 (model output before any features)"),
    ("Top Feature (Global)", "TransactionAmt (highest mean |SHAP|)"),
    ("Second Feature", "txn_count_24hr (24-hour velocity)"),
    ("Third Feature", "amount_deviation (spending anomaly)"),
    ("Case Studies", "6 (TP clear, TP velocity, FN, FP, auto-block, borderline)"),
    ("Visualizations Produced", "5 (beeswarm, bar, waterfall, fraud-vs-legit, tiers)"),
    ("Fair Lending Features Reviewed", "7 (2 flagged MEDIUM risk)"),
    ("SR 11-7 Items Complete", "8 of 12 (67%)"),
    ("Governance Items Pending", "4 (require production data)"),
    ("Documented Limitations", "6 (with compensating controls)"),
    ("Regulatory Frameworks", "SR 11-7, OCC 2011-12, ECOA, GDPR Art. 22"),
    ("Monitoring Schedule", "Daily / Weekly / Monthly / Quarterly / Annual"),
    ("Audit Trail Retention", "Minimum 7 years"),
    ("Model Risk Tier", "Tier 2 (material financial impact)"),
]

GLOSSARY = [
    ("Base Value", "The model's average prediction before seeing any features (0.0178 = 1.78% fraud rate). Every explanation starts here."),
    ("Beeswarm Plot", "A SHAP visualization where each dot is a transaction; position shows impact on fraud score; color shows the feature's actual value."),
    ("Compensating Controls", "Additional safeguards that mitigate known model limitations (e.g., manual review layer compensates for model uncertainty)."),
    ("Dependence Plot", "A chart showing how a feature's value relates to its SHAP contribution, revealing non-linear relationships and interaction effects."),
    ("Disparate Impact", "When a seemingly neutral practice (like a fraud model) disproportionately affects a protected group (race, gender, age)."),
    ("ECOA", "Equal Credit Opportunity Act \u2014 prohibits discrimination in credit decisions based on protected attributes."),
    ("Fair Lending", "Regulatory requirement ensuring financial models do not discriminate against protected classes, even unintentionally."),
    ("Feature Attribution", "Assigning a specific contribution value to each input feature for a given prediction. SHAP provides exact attributions."),
    ("GDPR Art. 22", "European regulation granting customers the right to an explanation when automated decisions significantly affect them."),
    ("Global Explainability", "Understanding which features matter most across ALL transactions (the model's overall strategy)."),
    ("Local Explainability", "Understanding why the model made a specific decision for ONE particular transaction."),
    ("Mean |SHAP|", "The average of absolute SHAP values for a feature across all transactions \u2014 a single importance number."),
    ("Proxy Discrimination", "When a non-protected feature (like time of day) correlates with a protected attribute (like geography or ethnicity)."),
    ("SHAP", "SHapley Additive exPlanations \u2014 a game-theoretic method that assigns each feature a contribution to each prediction."),
    ("Shapley Value", "From cooperative game theory: the fair allocation of a 'payout' (prediction) among 'players' (features)."),
    ("SR 11-7", "Federal Reserve guidance on Model Risk Management. Requires model documentation, validation, and ongoing monitoring."),
    ("TreeExplainer", "SHAP's exact algorithm for tree-based models. Computes true Shapley values in polynomial time (no approximation)."),
    ("Waterfall Plot", "A SHAP visualization showing step-by-step how each feature moves the prediction from the base value to the final score."),
]


# ══════════════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_paragraph_spacing(paragraph, before=0, after=0, line=240):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_formatted_text(cell, text, font_name="Calibri", font_size=10,
                       bold=False, color_hex=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=20, after=20, line=240)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def add_callout_box(doc, text, box_type="insight"):
    colors = {
        "insight": (C_INSIGHT_BG, C_INSIGHT_BORDER, "Key Insight"),
        "business": (C_BUSINESS_BG, C_BUSINESS_BORDER, "Business Impact"),
        "caution": (C_CAUTION_BG, C_CAUTION_BORDER, "Caution"),
    }
    bg_color, border_color, label = colors.get(box_type, colors["insight"])

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=120, line=264)

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), bg_color)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)

    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), border_color)
    left.set(qn("w:space"), "4")
    pBdr.append(left)
    pPr.append(pBdr)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:right"), "360")
    pPr.append(ind)

    label_run = p.add_run(f"{label}: ")
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(10)
    label_run.font.bold = True
    label_run.font.color.rgb = RGBColor.from_string(border_color)

    content_run = p.add_run(text)
    content_run.font.name = "Calibri"
    content_run.font.size = Pt(10)
    content_run.font.italic = True
    content_run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)


def add_figure(doc, image_filename, caption_text):
    img_path = FIGURES_DIR / image_filename
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cap, before=40, after=200, line=240)
        cap_run = cap.add_run(caption_text)
        cap_run.font.name = "Calibri"
        cap_run.font.size = Pt(9)
        cap_run.font.italic = True
        cap_run.font.color.rgb = RGBColor.from_string("666666")
    else:
        p = doc.add_paragraph(
            f"[Figure not available: {image_filename} \u2014 run notebook to generate]"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor.from_string("999999")


def add_matrix_table(doc, matrix_data):
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9360")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    headers = [
        "Layer", "Technical\n(BDS Colleague)",
        "Business\n(Manager / Regulator)", "Simple\n(Grandmother)"
    ]
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, layer in enumerate(LAYERS):
        row_idx = i + 1
        layer_cell = table.cell(row_idx, 0)
        set_cell_shading(layer_cell, C_MED_BLUE)
        set_cell_margins(layer_cell, top=60, bottom=60, left=100, right=100)
        add_formatted_text(layer_cell, layer, font_size=9, bold=True, color_hex=C_WHITE)
        layer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for j in range(3):
            cell = table.cell(row_idx, j + 1)
            bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            text = matrix_data.get((i, j), "")
            add_formatted_text(cell, text, font_size=9)

    doc.add_paragraph("")
    return table


def add_section_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)
    return heading


def add_page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
#  UNIQUE ELEMENTS
# ══════════════════════════════════════════════════════════════════════

def add_case_studies_table(doc):
    """Add the 6 case studies summary table unique to this notebook."""
    add_section_heading(doc, "Case Studies Overview", level=2)

    data = [
        ("Case 1", "True Positive \u2014 Clear Fraud", "0.9094", "FRAUD", "AUTO-BLOCK"),
        ("Case 2", "True Positive \u2014 Velocity-Driven", "0.7332", "FRAUD", "MANUAL REVIEW"),
        ("Case 3", "False Negative \u2014 Missed Fraud", "0.0853", "FRAUD", "AUTO-APPROVE"),
        ("Case 4", "False Positive \u2014 Legitimate Flagged", "0.9342", "LEGIT", "AUTO-BLOCK"),
        ("Case 5", "Auto-Block \u2014 High Confidence", "0.9342", "LEGIT", "AUTO-BLOCK"),
        ("Case 6", "Borderline \u2014 Near Threshold", "0.3648", "LEGIT", "AUTO-APPROVE"),
    ]

    table = doc.add_table(rows=len(data) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Case", "Type", "Score", "Actual", "Decision"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (case, case_type, score, actual, decision) in enumerate(data):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        for j, text in enumerate([case, case_type, score, actual, decision]):
            cell = table.cell(i + 1, j)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=80, right=80)
            add_formatted_text(cell, text, font_size=10, bold=(j == 0))
            if j > 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")


def add_governance_checklist_table(doc):
    """Add the SR 11-7 governance checklist table."""
    add_section_heading(doc, "Model Governance Checklist (SR 11-7)", level=2)

    data = [
        ("\u2713", "Model documentation (purpose, inputs, outputs, assumptions)", "Complete"),
        ("\u2713", "Performance metrics on held-out test data", "Complete"),
        ("\u2713", "Global explainability (feature importance, SHAP summary)", "Complete"),
        ("\u2713", "Local explainability (individual transaction SHAP)", "Complete"),
        ("\u2713", "Limitations and known risks documented", "Complete"),
        ("\u2713", "Fair lending feature review conducted", "Complete"),
        ("\u2713", "Right-to-explanation capability demonstrated", "Complete"),
        ("\u2713", "Audit trail requirements specified", "Complete"),
        ("\u2717", "Disparate impact testing (requires demographic data)", "Pending"),
        ("\u2717", "Champion/challenger framework", "Pending"),
        ("\u2717", "Ongoing monitoring dashboard (drift detection)", "Pending"),
        ("\u2717", "Quarterly model revalidation schedule", "Pending"),
    ]

    table = doc.add_table(rows=len(data) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Status", "Item", "State"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (status, item, state) in enumerate(data):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        cell_s = table.cell(i + 1, 0)
        set_cell_shading(cell_s, bg)
        set_cell_margins(cell_s, top=30, bottom=30, left=80, right=80)
        color = C_LEGIT_GREEN if state == "Complete" else C_FRAUD_RED
        add_formatted_text(cell_s, status, font_size=12, bold=True, color_hex=color)
        cell_s.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_i = table.cell(i + 1, 1)
        set_cell_shading(cell_i, bg)
        set_cell_margins(cell_i, top=30, bottom=30, left=80, right=80)
        add_formatted_text(cell_i, item, font_size=10)

        cell_st = table.cell(i + 1, 2)
        set_cell_shading(cell_st, bg)
        set_cell_margins(cell_st, top=30, bottom=30, left=80, right=80)
        add_formatted_text(cell_st, state, font_size=10, bold=True,
                           color_hex=(C_LEGIT_GREEN if state == "Complete" else C_FRAUD_RED))
        cell_st.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")


# ══════════════════════════════════════════════════════════════════════
#  DOCUMENT CONSTRUCTION
# ══════════════════════════════════════════════════════════════════════

def add_cover_page(doc):
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Agent Fraud Sentinel")
    run.font.name = "Calibri"
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("SHAP Explainability Deep Dive: 4 Layers \u00d7 3 Perspectives")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("\u2500" * 50)
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)
    run.font.size = Pt(12)

    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ref.add_run("Notebook: 04_shap_explainability.ipynb")
    run.font.name = "Consolas"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        "IEEE-CIS Fraud Detection Dataset\n"
        "SHAP TreeExplainer | Global & Local Explainability\n"
        "6 Case Studies | Regulatory Compliance (SR 11-7)"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)

    for _ in range(3):
        doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("BAFS Project \u2014 February 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("888888")

    add_page_break(doc)


def add_toc_placeholder(doc):
    add_section_heading(doc, "Table of Contents", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)
    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar_separate)

    placeholder_run = p.add_run(
        "(Right-click here and select 'Update Field' to generate Table of Contents)"
    )
    placeholder_run.font.name = "Calibri"
    placeholder_run.font.size = Pt(10)
    placeholder_run.font.italic = True
    placeholder_run.font.color.rgb = RGBColor.from_string("999999")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    placeholder_run._r.append(fldChar_end)

    add_page_break(doc)


def add_framework_explanation(doc):
    add_section_heading(doc, "The 4-Layer \u00d7 3-Perspective Framework", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "This document examines each section of the SHAP Explainability notebook through "
        "two dimensions: four analytical layers and three audience perspectives. This "
        "framework ensures complete understanding \u2014 from raw technical detail to "
        "business impact \u2014 accessible to any reader regardless of their background."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    add_section_heading(doc, "The Four Layers", level=3)
    for title_text, desc in [
        ("Layer 1 \u2014 WHAT did I do?",
         "Describes the concrete actions: SHAP values computed, plots generated, cases analyzed."),
        ("Layer 2 \u2014 WHY did I do it?",
         "Explains the motivation: why explainability, why these cases, why regulatory compliance."),
        ("Layer 3 \u2014 HOW does it work?",
         "Details the mechanics: TreeExplainer, Shapley values, dependence analysis, waterfall decomposition."),
        ("Layer 4 \u2014 WHAT does the bank gain?",
         "Translates results into value: regulatory readiness, analyst tools, customer trust, model improvement."),
    ]:
        p = doc.add_paragraph()
        bold_run = p.add_run(title_text + "  ")
        bold_run.font.name = "Calibri"
        bold_run.font.size = Pt(11)
        bold_run.font.bold = True
        bold_run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)
        desc_run = p.add_run(desc)
        desc_run.font.name = "Calibri"
        desc_run.font.size = Pt(11)

    add_section_heading(doc, "The Three Perspectives", level=3)
    for title_text, desc in [
        ("Technical (BDS Colleague)",
         "Uses ML/XAI terminology, references Shapley values, TreeExplainer, dependence functions."),
        ("Business (Manager / Regulator)",
         "Focuses on compliance, audit trails, operational insights, and customer impact."),
        ("Simple (Grandmother)",
         "Uses everyday analogies and plain language. No technical background assumed."),
    ]:
        p = doc.add_paragraph()
        bold_run = p.add_run(title_text + "  ")
        bold_run.font.name = "Calibri"
        bold_run.font.size = Pt(11)
        bold_run.font.bold = True
        bold_run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)
        desc_run = p.add_run(desc)
        desc_run.font.name = "Calibri"
        desc_run.font.size = Pt(11)


def add_chapter(doc, chapter_data):
    num = chapter_data["number"]
    title = chapter_data["title"]

    add_page_break(doc)
    add_section_heading(doc, f"Chapter {num}: {title}", level=1)

    if chapter_data.get("subtitle"):
        p = doc.add_paragraph()
        run = p.add_run(chapter_data["subtitle"])
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=200, line=276)
    run = p.add_run(chapter_data["narrative"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    add_section_heading(doc, "Analysis Matrix", level=2)
    add_matrix_table(doc, chapter_data["matrix"])

    for box_type, text in chapter_data.get("callouts", []):
        add_callout_box(doc, text, box_type)

    for fig_filename, caption in chapter_data.get("figures", []):
        add_figure(doc, fig_filename, caption)


def add_summary_statistics_table(doc):
    add_section_heading(doc, "SHAP Explainability Summary", level=2)

    table = doc.add_table(rows=len(SUMMARY_TABLE_DATA) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Metric", "Value"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (metric, value) in enumerate(SUMMARY_TABLE_DATA):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        for j, text in enumerate([metric, value]):
            cell = table.cell(i + 1, j)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=100, right=100)
            add_formatted_text(cell, text, font_size=10, bold=(j == 0))
            if j == 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_appendix_gallery(doc):
    add_page_break(doc)
    add_section_heading(doc, "Appendix A: SHAP Visualization Gallery", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "All six SHAP visualizations produced in this notebook, presented with "
        "captions explaining their purpose and how to interpret them."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    set_paragraph_spacing(p, after=200)

    gallery = [
        ("shap_summary_beeswarm.png",
         "Beeswarm Summary Plot: Each dot represents one transaction. The horizontal "
         "position shows how much the feature pushed the fraud score up (right) or "
         "down (left). Color indicates the feature's actual value (red = high, blue = low). "
         "Features are sorted by overall importance."),
        ("shap_feature_importance_bar.png",
         "Feature Importance Bar Chart: The mean absolute SHAP value for each feature, "
         "providing a single importance ranking. TransactionAmt is the most important "
         "feature overall, followed by 24-hour velocity and spending anomaly."),
        ("shap_dependence_top4.png",
         "Dependence Plots (Top 4 Features): Each plot shows the relationship between "
         "a feature's value (x-axis) and its SHAP contribution (y-axis). Color indicates "
         "the strongest interacting feature. Key insights: velocity shows step-function "
         "risk above ~5 transactions; hour_of_day shows elevated risk at 6-9 AM."),
        ("shap_waterfall_cases.png",
         "Waterfall Plots (6 Case Studies): Step-by-step decomposition of each prediction, "
         "from the base value (0.0178) to the final fraud score. Red bars push toward "
         "fraud; blue bars push toward legitimate. Cases cover TP, FN, FP, auto-block, "
         "and borderline scenarios."),
        ("shap_fraud_vs_legit.png",
         "Fraud vs Legitimate Comparison: Grouped bar chart showing the average SHAP "
         "contribution per feature for fraud and legitimate transactions separately. "
         "The gap between bars indicates the feature's discrimination power."),
        ("shap_risk_tiers.png",
         "Feature Contribution by Risk Tier: How feature importance changes across "
         "auto-approve, manual review, and auto-block tiers. Auto-block transactions "
         "show 3-5x higher mean |SHAP| than auto-approve, confirming qualitatively "
         "different risk profiles."),
    ]

    for filename, caption in gallery:
        add_figure(doc, filename, caption)


def add_appendix_glossary(doc):
    add_page_break(doc)
    add_section_heading(doc, "Appendix B: Glossary", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "Key terms used throughout this document, defined for non-technical readers."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    set_paragraph_spacing(p, after=200)

    table = doc.add_table(rows=len(GLOSSARY) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Term", "Definition"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)

    for i, (term, definition) in enumerate(GLOSSARY):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        cell_term = table.cell(i + 1, 0)
        set_cell_shading(cell_term, bg)
        set_cell_margins(cell_term, top=30, bottom=30, left=100, right=100)
        add_formatted_text(cell_term, term, font_size=10, bold=True)

        cell_def = table.cell(i + 1, 1)
        set_cell_shading(cell_def, bg)
        set_cell_margins(cell_def, top=30, bottom=30, left=100, right=100)
        add_formatted_text(cell_def, definition, font_size=10)


def setup_header_footer(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        h_para.text = ""
        run_left = h_para.add_run("Agent Fraud Sentinel \u2014 SHAP Explainability")
        run_left.font.name = "Calibri"
        run_left.font.size = Pt(8)
        run_left.font.color.rgb = RGBColor.from_string("999999")
        h_para.add_run("\t\t")
        run_right = h_para.add_run("BAFS")
        run_right.font.name = "Calibri"
        run_right.font.size = Pt(8)
        run_right.font.bold = True
        run_right.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

        footer = section.footer
        footer.is_linked_to_previous = False
        f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        f_para.text = ""
        f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_page = f_para.add_run("Page ")
        run_page.font.name = "Calibri"
        run_page.font.size = Pt(8)
        run_page.font.color.rgb = RGBColor.from_string("999999")
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run_page._r.append(fldChar1)
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run_page._r.append(instrText)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        run_page._r.append(fldChar2)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run_page._r.append(fldChar3)


# ══════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════

def main():
    print("Generating SHAP Explainability Analysis Matrix document...")

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i in range(1, 4):
        h_style = doc.styles[f"Heading {i}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

    # ── Build Document ─────────────────────────────────────────────
    add_cover_page(doc)
    add_toc_placeholder(doc)

    # Chapter 0
    ch0 = ALL_CHAPTERS[0]
    add_section_heading(doc, f"Chapter 0: {ch0['title']}", level=1)
    p = doc.add_paragraph()
    run = p.add_run(ch0["subtitle"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=200, line=276)
    run = p.add_run(ch0["narrative"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    add_framework_explanation(doc)
    add_case_studies_table(doc)
    add_governance_checklist_table(doc)

    add_section_heading(doc, "Full-Project Analysis Matrix", level=2)
    add_matrix_table(doc, ch0["matrix"])

    for box_type, text in ch0.get("callouts", []):
        add_callout_box(doc, text, box_type)

    # Chapters 1-6
    for chapter_data in ALL_CHAPTERS[1:]:
        add_chapter(doc, chapter_data)

    # Summary
    add_summary_statistics_table(doc)

    # Appendices
    add_appendix_gallery(doc)
    add_appendix_glossary(doc)

    # Header/Footer
    setup_header_footer(doc)

    # Save
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"File size: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
