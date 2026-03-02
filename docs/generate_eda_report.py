"""
Generate Word Document: EDA Notebook 01 - 4 Layers x 3 Perspectives
Agent Fraud Sentinel (BAFS) Project

Produces: docs/eda_01_analysis_matrix.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import os

# ── Paths ──────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent.parent
FIGURES_DIR = BASE_DIR / "data" / "processed"
OUTPUT_PATH = BASE_DIR / "docs" / "eda_01_analysis_matrix.docx"

# ── Colors (hex without #) ─────────────────────────────────────────────
C_DARK_BLUE = "1B4F72"
C_MED_BLUE = "2980B9"
C_LIGHT_BLUE = "D6EAF8"
C_FRAUD_RED = "E74C3C"
C_LEGIT_GREEN = "2ECC71"
C_LIGHT_GRAY = "F2F3F4"
C_WHITE = "FFFFFF"
C_BLACK = "000000"
C_DARK_GRAY = "2C3E50"
C_INSIGHT_BG = "D6EAF8"
C_INSIGHT_BORDER = "2980B9"
C_BUSINESS_BG = "D5F5E3"
C_BUSINESS_BORDER = "27AE60"

LAYERS = ["WHAT did I do?", "WHY did I do it?", "HOW does it work?", "WHAT does the bank gain?"]
PERSPECTIVES = ["Technical", "Business", "Simple"]

# ── Content: All Chapter Matrices ──────────────────────────────────────

CH0_MATRIX = {
    (0, 0): "Performed exploratory data analysis on 590,540 IEEE-CIS transactions (434 features) using pandas, matplotlib, and seaborn. Computed correlations, missingness profiles, distributional statistics, and temporal patterns across 32 notebook cells.",
    (0, 1): "Conducted a comprehensive diagnostic of the bank's transaction data to understand where fraud hides, how much data is usable, and what patterns distinguish fraudulent from legitimate transactions.",
    (0, 2): "We looked at about 590,000 bank transactions to find clues about which ones might be fraud. We checked what the data looks like, what's missing, and when fraud tends to happen.",
    (1, 0): "To identify discriminative features, quantify class imbalance (3.50%, 1:27 ratio), assess data quality (214/434 features >50% missing), and establish cost assumptions (FN=$227, FP=$10) before feature engineering and modeling.",
    (1, 1): "Because building a fraud detection system without understanding the data is like building a house without surveying the land. The bank needs to know the fraud rate, data gaps, and cost of errors before investing in a model.",
    (1, 2): "Because you can't catch thieves if you don't first understand how they behave. We needed to learn the patterns before building the alarm system.",
    (2, 0): "Left-joined transaction (590K x 394) and identity (144K x 41) tables on TransactionID. Computed absolute Pearson correlations with isFraud, profiled missing values per column, plotted log-transformed amount distributions, and extracted hour/day-of-week from TransactionDT (seconds since epoch, base date 2017-12-01).",
    (2, 1): "Merged two data files into a single view, ranked all 434 features by their connection to fraud, mapped data quality gaps, compared dollar amounts between fraud and legitimate transactions, and analyzed what times and days fraud peaks.",
    (2, 2): "We combined two spreadsheets into one big table, then sorted through all 434 columns to find which ones are most connected to fraud. We also checked which hours and days have the most fraud.",
    (3, 0): "Quantified inputs for scale_pos_weight (28.56), identified that top correlated features (V-series) have 76-78% missing data (ruling them out as direct model inputs), and established the 22.7:1 cost ratio that drives threshold optimization in Phase 3.",
    (3, 1): "The bank now knows: (1) fraud is 3.5% of transactions (manageable but costly), (2) half the data columns are unreliable, (3) early mornings and weekends are high-risk, and (4) missing one fraud costs 22.7x more than a false alarm.",
    (3, 2): "The bank learns when to be extra careful (early mornings, weekends), that a missed fraud costs $227 but checking a good transaction only costs $10, and that the data needs cleaning before it can be useful.",
}

CH1_MATRIX = {
    (0, 0): "Loaded train_transaction.csv (590,540 x 394) and train_identity.csv (144,233 x 41). Performed left join on TransactionID producing a merged DataFrame of 590,540 x 434. Computed isFraud value counts: 20,663 positive (3.50%), 569,877 negative (96.50%), imbalance ratio 1:27. Generated bar chart and pie chart of class distribution.",
    (0, 1): "Combined the bank's transaction records with customer identity data into a single unified dataset. Measured the fraud rate: out of every 100 transactions, roughly 3.5 are fraudulent. Visualized the extreme imbalance between fraud and legitimate transactions.",
    (0, 2): "We took two big lists \u2014 one with transaction details and one with customer info \u2014 and combined them. Then we counted: out of about 590,000 transactions, only about 20,600 were fraud. That means for every 1 fraud, there are 27 normal transactions.",
    (1, 0): "Identity features (device info, browser data) are available for only 24.4% of transactions (144K/590K). Left join preserves all transactions while adding identity where available. The 1:27 imbalance necessitates cost-sensitive learning or resampling; standard accuracy would yield 96.5% by predicting all-legitimate.",
    (1, 1): "The bank stores transaction data and identity data in separate systems. Combining them gives more clues per transaction. The extreme rarity of fraud (3.5%) means a naive system that says 'everything is fine' would be 'right' 96.5% of the time while catching zero fraud.",
    (1, 2): "We joined the two lists because having more information about each transaction helps us spot fraud. And we noticed fraud is very rare \u2014 like finding a needle in a haystack \u2014 so we need a smart approach, not a simple one.",
    (2, 0): "pd.merge(train_transaction, train_identity, on='TransactionID', how='left') preserves all 590K rows. df['isFraud'].value_counts() yields class counts. Imbalance ratio computed as (total - fraud) / fraud. Visualization uses matplotlib subplots: bar chart with count annotations and pie chart with exploded fraud slice.",
    (2, 1): "Data engineers merge the tables using a common transaction identifier. The fraud count is divided by total transactions to get the percentage. Two charts are produced: a bar chart showing the raw counts and a pie chart showing the proportions.",
    (2, 2): "We matched transactions to identities using a shared ID number. Then we counted frauds and made two pictures: one showing the actual numbers and one showing the percentages, like a slice of pie.",
    (3, 0): "The 1:27 ratio directly determines scale_pos_weight=28.56 in XGBoost and justifies using PR-AUC over ROC-AUC as the primary evaluation metric (avoids inflated scores from true negatives in imbalanced data).",
    (3, 1): "The bank understands that its fraud problem, while small in percentage, represents 20,663 incidents. At a median loss of $75 per fraud, that is approximately $1.55M in potential losses in this dataset alone. This justifies the investment in a detection system.",
    (3, 2): "The bank learns that even though fraud seems rare (only 3.5%), it adds up to over 20,000 cases and potentially over a million dollars in losses. That's why building this system matters.",
}

CH2_MATRIX = {
    (0, 0): "Computed absolute Pearson correlation of all 432 numeric features with isFraud. Top 10: V257 (0.383), V246 (0.367), V244 (0.364), V242 (0.361), V201 (0.328), V200 (0.319), V189 (0.308), V188 (0.304), V258 (0.297), V45 (0.282). Profiled missing data: 214/434 features have >50% missing, 208 have >75% missing.",
    (0, 1): "Ranked all 434 data columns by how strongly they relate to fraud, and assessed data quality. The most fraud-predictive features are anonymized 'V' columns, but nearly half of all features have more than 50% of their values missing.",
    (0, 2): "We checked which of the 434 columns are most connected to fraud, and found the top ones are mysterious 'V' columns. We also discovered that about half the columns are mostly empty \u2014 like a form where people skip most questions.",
    (1, 0): "To identify candidate features for modeling and assess data quality constraints. High correlation with the target suggests discriminative power, but high missingness limits direct usability. The paradox (best predictors are mostly missing) motivates engineering new features from low-missing columns in Phase 2.",
    (1, 1): "The bank needs to know which signals are strongest so it can focus resources. Discovering that the best signals come from unreliable data (76-78% missing) is a critical strategic finding \u2014 it means the bank cannot rely on these features alone and must create new, more reliable indicators.",
    (1, 2): "We wanted to know which clues are best for catching fraud. The strongest clues turned out to be in columns where most of the information is missing. So we can't rely on them alone and need to create better clues from the data we do have.",
    (2, 0): "df[numeric_cols].corrwith(df['isFraud']).abs().sort_values(ascending=False) computes pairwise Pearson r. Missing analysis: (df.isnull().sum() / len(df) * 100). Horizontal bar chart with RdYlGn_r colormap for correlations; dual-panel figure (histogram of missing percentages + top-15 missing features bar chart).",
    (2, 1): "Each feature's values are statistically compared against the fraud label to produce a correlation score (0 = no relationship, 1 = perfect relationship). Separately, each column is checked for the percentage of blank entries. Results are displayed as ranked bar charts.",
    (2, 2): "We measured how much each column 'moves together' with fraud \u2014 like checking if one thing goes up when the other goes up. We also counted how many blanks each column has. Then we drew charts to show the results.",
    (3, 0): "Established that V-features are not viable as direct model inputs due to missingness, justifying the Phase 2 decision to engineer 7 new leakage-free features from low-missing columns (TransactionAmt, TransactionDT, card1, addr1, P_emaildomain). Output top_features.csv provides a traceable artifact for model governance.",
    (3, 1): "The bank avoids a costly mistake: blindly building a model on the 'best' features that are 78% empty. Instead, Phase 2 will engineer reliable indicators from data the bank actually collects consistently. This ensures the model works on real incoming transactions.",
    (3, 2): "The bank learns it can't use its most fraud-connected columns because they're mostly empty. Instead, the team will build new, reliable clues from the data that is actually filled in \u2014 like time of day and transaction amount.",
}

CH3_MATRIX = {
    (0, 0): "Compared TransactionAmt distributions by fraud status. Fraud: mean=$149.24, median=$75.00, std=$232.21, max=$5,191. Legitimate: mean=$134.51, median=$68.50, std=$239.40, max=$31,937. Generated histogram ($0-$500 range), log-scale density overlay, and box plot. ~95% of both classes fall below $500.",
    (0, 1): "Analyzed how much money moves in fraudulent vs. legitimate transactions. Fraud transactions are slightly higher on average ($149 vs $135) and cluster around $75. The vast majority of both types are under $500, making it hard to distinguish fraud by amount alone.",
    (0, 2): "We looked at how much money each transaction involved. Fraud transactions tend to be a bit higher \u2014 around $75 in the middle \u2014 compared to about $68.50 for normal ones. But both types are mostly under $500, so you can't just flag big purchases.",
    (1, 0): "To assess whether TransactionAmt alone has discriminative power. The moderate difference in medians ($75 vs $68.50) and overlapping distributions confirm that amount is a weak standalone signal but valuable as a component of behavioral features (amount_deviation Z-score in Phase 2).",
    (1, 1): "The bank needs to know if 'big transactions = fraud' is a valid rule. The data shows it's not: fraudsters intentionally stay within normal ranges to avoid detection. The system needs smarter signals than just transaction size.",
    (1, 2): "We checked if fraudsters always steal big amounts. They don't \u2014 they're clever and keep amounts similar to normal purchases. So the alarm system needs to look at more than just the dollar amount.",
    (2, 0): "df.groupby('isFraud')['TransactionAmt'].describe() for summary statistics. Dollar-denominated histogram with np.arange(0, 525, 25) bins, side-by-side subplots. Log transform via np.log1p(data) for density overlay. Box plot via seaborn boxplot with y-axis capped at $1,000.",
    (2, 1): "Summary statistics (mean, median, standard deviation) are calculated separately for fraud and legitimate groups. Three chart types visualize the comparison: side-by-side histograms in dollar amounts, a log-scale overlay for shape comparison, and a box-and-whisker plot for spread.",
    (2, 2): "We calculated the average, middle value, and spread for fraud and normal transactions separately. Then we drew three types of pictures: bar charts in real dollar amounts, a special view that stretches out small differences, and a box diagram showing the range.",
    (3, 0): "Justified including TransactionAmt as one of the 7 model features and motivated the amount_deviation engineered feature (Z-score vs client history), which captures whether a transaction is unusual for that client rather than unusual in absolute terms. Median fraud amount ($75) becomes the FN cost parameter.",
    (3, 1): "The bank gains a nuanced understanding: it's not the absolute amount that matters, but whether the amount is unusual for that specific customer. A $200 transaction is normal for a business traveler but suspicious for a retiree who usually spends $30.",
    (3, 2): "The bank learns that the dollar amount alone doesn't tell you much. What matters is whether the amount is unusual for that person. The system will compare each transaction to what that customer normally spends.",
}

CH4_MATRIX = {
    (0, 0): "Derived datetime from TransactionDT (seconds since epoch, base 2017-12-01). Computed fraud rate by hour (0-23) and day of week (Mon-Sun). Peak fraud hours: 7, 8, 9 (above 3.5% baseline). Lowest fraud hours: 13, 14, 15. Peak fraud days: Fri, Sun, Sat. Lowest: Tue, Wed, Thu.",
    (0, 1): "Mapped when fraud is most likely to occur. Fraud peaks in the early morning (7-9 AM) and on weekends (Friday through Sunday). Midweek afternoons are the safest periods.",
    (0, 2): "We checked what time and what day fraud happens most. Early mornings (7-9 AM) and weekends (especially Friday, Saturday, Sunday) have the most fraud. Weekday afternoons are the safest.",
    (1, 0): "Temporal features capture cyclical fraud patterns: (a) fraudsters exploit off-hours when monitoring is lighter, (b) stolen card details are tested in early morning hours, (c) weekend shopping volume provides cover. These signals justify hour_of_day and is_weekend as engineered features.",
    (1, 1): "The bank needs to know when to be most vigilant. If fraud peaks at 7-9 AM, the fraud operations team can allocate more analysts during those hours. Weekend patterns may justify different staffing models on Friday through Sunday.",
    (1, 2): "Fraudsters prefer early mornings and weekends because fewer people are watching. Knowing this helps the bank put more guards on duty at the right times.",
    (2, 0): "pd.to_datetime('2017-12-01') + pd.to_timedelta(df['TransactionDT'], unit='s') converts epoch seconds. .dt.hour and .dt.dayofweek extract temporal features. Fraud rate: df.groupby('hour')['isFraud'].mean() * 100. Dual-panel bar charts with horizontal red dashed line at 3.5% baseline.",
    (2, 1): "The raw timestamp (seconds since a reference point) is converted to clock time and calendar day. The fraud percentage is calculated for each hour and day separately. Charts show which hours and days are above or below the average fraud rate of 3.5%.",
    (2, 2): "We converted the raw time numbers into real clock times and calendar days. Then we calculated the fraud percentage for each hour and each day. The charts show which times are worse than average.",
    (3, 0): "Two temporal features (hour_of_day, is_weekend) added to the final 7-feature model, contributing to fraud detection without data leakage (both computable at transaction time). These features interact with velocity features in XGBoost to detect patterns like 'multiple early-morning transactions from a new device.'",
    (3, 1): "The bank can implement time-based alert escalation rules: tighter thresholds during peak fraud hours (7-9 AM) and weekends. This is an immediate operational improvement that can be deployed as a simple business rule alongside the model.",
    (3, 2): "The bank can put extra people on watch during early mornings and weekends, right away, even before the full computer system is ready. It's like knowing that most burglaries happen at night \u2014 you add more patrols then.",
}

CH5_MATRIX = {
    (0, 0): "Established cost parameters: FN=$227.00 (full economic cost: transaction loss $75 + chargeback $27 + ops $50 + reputational $75), FP=$10.00 (industry benchmark for manual review), ratio 22.7:1. The median fraud TransactionAmt ($75.00) is the starting point; total economic impact reaches $227. Created summary statistics table (10 key metrics). Saved top_features.csv for Phase 2 traceability.",
    (0, 1): "Defined the financial consequences of each type of error. Missing a fraud costs the bank $227 in total economic impact (transaction loss, chargeback fees, investigation, and reputational damage). Flagging a legitimate transaction for review costs $10 in analyst time. These numbers drive every decision about how aggressive the fraud system should be.",
    (0, 2): "We figured out the 'price' of mistakes. Missing a fraud costs about $227 (the money lost plus bank costs). Checking a good transaction by mistake costs about $10 (the employee's time). So missing a fraud is about 22.7 times worse than a false alarm.",
    (1, 0): "Standard ML metrics (accuracy, F1) assign equal weight to FP and FN errors, which misrepresents business reality. The 22.7:1 asymmetry informs: (1) scale_pos_weight in XGBoost, (2) cost-weighted threshold optimization in Phase 3, and (3) the business case for a high-recall operating point.",
    (1, 1): "Because not all mistakes are equal. A bank that misses fraud loses money AND customer trust. A bank that over-flags transactions annoys customers but at a much lower cost. The 22.7:1 ratio quantifies this imbalance so every subsequent decision is grounded in financial reality.",
    (1, 2): "Because missing a thief is much worse than bothering an honest customer. If we treat both mistakes the same, the system would be lazy about catching thieves. The 22.7:1 ratio tells the system: 'catching thieves matters 22.7 times more.'",
    (2, 0): "Median via df[df['isFraud']==1]['TransactionAmt'].median() = $75.00. Mean rejected due to right-skew (mean/median ratio = 1.99). FP cost from industry benchmarks (analyst time + customer friction). Cost function: total_cost = FN_count * 227 + FP_count * 10. Top features saved via pd.DataFrame.to_csv().",
    (2, 1): "The median was chosen over the average because a few very large frauds ($5,000+) would inflate the average and make the system overly cautious. The $10 review cost reflects the real cost of an analyst spending time on a false lead. Together, these form the scoring system for the model.",
    (2, 2): "We used the 'middle' fraud amount ($75) instead of the average ($149) because a few giant frauds would throw off the average. The $10 review cost is what it costs the bank in employee time to check a flagged transaction.",
    (3, 0): "The cost framework enables dollar-denominated model evaluation: Phase 3 reports total cost at each threshold, allowing direct comparison between a 'catch-everything' strategy ($598K at 76% recall) and a 'minimize-cost' strategy ($328K at 14% recall). This makes the precision-recall tradeoff tangible.",
    (3, 1): "The bank gains a decision framework in dollars, not percentages. When the CEO asks 'how much does this model save us?', the team can answer in monetary terms. The 22.7:1 ratio also provides audit-ready documentation for regulators requiring cost-benefit justification.",
    (3, 2): "The bank gets a simple rule for judging the system: 'What is the total cost of mistakes?' This lets everyone \u2014 from the boss to the regulator \u2014 understand whether the system is working, using real dollar amounts instead of confusing percentages.",
}

ALL_CHAPTERS = [
    {
        "number": 0,
        "title": "Executive Overview",
        "subtitle": "Full-Project Summary",
        "narrative": (
            "This notebook represents the foundation of the Agent Fraud Sentinel project. "
            "Over 32 cells, it systematically investigates the IEEE-CIS fraud detection dataset "
            "(590,540 transactions, 434 features, 3.50% fraud rate) to identify fraud signals, "
            "assess data quality, and establish cost assumptions. Every downstream decision \u2014 "
            "from feature engineering to threshold optimization to regulatory documentation \u2014 "
            "traces back to insights discovered here."
        ),
        "matrix": CH0_MATRIX,
        "figures": [],
        "callouts": [
            ("insight", "The EDA is not just exploratory analysis \u2014 it is the strategic blueprint for a production-grade fraud detection system that balances detection accuracy, operational costs, and regulatory compliance."),
        ],
    },
    {
        "number": 1,
        "title": "Data Acquisition & Fraud Rate",
        "subtitle": "Notebook Sections 1\u20133: Loading, Merging, Class Distribution",
        "narrative": (
            "The analysis begins by loading and merging two data sources: the transaction table "
            "(590,540 rows, 394 columns) and the identity table (144,233 rows, 41 columns). "
            "A left join on TransactionID produces a unified dataset of 434 features. "
            "The critical first finding: only 3.50% of transactions are fraudulent (20,663 out "
            "of 590,540), creating a 1:27 class imbalance that will shape every modeling decision."
        ),
        "matrix": CH1_MATRIX,
        "figures": [
            ("class_distribution.png", "Figure 1: Class distribution \u2014 bar chart (left) and pie chart (right) showing the 96.50% vs 3.50% split between legitimate and fraudulent transactions."),
        ],
        "callouts": [
            ("business", "At a median loss of $75 per fraud, the 20,663 fraudulent transactions represent approximately $1.55 million in potential losses in this dataset alone."),
        ],
    },
    {
        "number": 2,
        "title": "Feature Landscape & Data Quality",
        "subtitle": "Notebook Sections 4\u20135: Top Correlations, Missing Data",
        "narrative": (
            "This chapter reveals a central paradox: the features most correlated with fraud "
            "(V-series features, with correlations up to 0.383) are also the most incomplete, "
            "with 76\u201378% missing values. Meanwhile, 214 out of 434 features (49.3%) have more "
            "than 50% of their values missing. This finding fundamentally shapes the modeling "
            "strategy: rather than relying on sparse V-features, Phase 2 will engineer 7 new "
            "features from reliably-collected data."
        ),
        "matrix": CH2_MATRIX,
        "figures": [
            ("top_correlations.png", "Figure 2: Top 10 features by absolute correlation with isFraud. All are anonymized V-features."),
            ("missing_data.png", "Figure 3: Missing data analysis \u2014 histogram of missing percentages (left) and top 15 most-missing features (right)."),
        ],
        "callouts": [
            ("insight", "The Paradox: The most fraud-correlated features (V257 at 0.383) have 77.9% missing data. The only top-10 feature with reasonable completeness is V45 at 28.6% missing."),
        ],
    },
    {
        "number": 3,
        "title": "Transaction Amount Patterns",
        "subtitle": "Notebook Section 6: Amount Distribution by Fraud Status",
        "narrative": (
            "Transaction amount analysis reveals subtle but important differences between fraud "
            "and legitimate transactions. Fraudulent transactions have a slightly higher median "
            "($75.00 vs $68.50) and a broader distribution, but ~95% of both classes fall below "
            "$500. The key insight: amount alone is a weak discriminator, but deviations from a "
            "client's typical spending pattern are highly informative."
        ),
        "matrix": CH3_MATRIX,
        "figures": [
            ("amount_distribution_dollars.png", "Figure 4: Transaction amount histograms ($0\u2013$500 range) for legitimate (green) and fraud (red) transactions."),
            ("amount_distribution.png", "Figure 5: Log-scale density overlay (left) and box plot (right) comparing transaction amounts by fraud status."),
        ],
        "callouts": [
            ("business", "Fraudsters intentionally keep amounts within normal ranges to avoid detection. A simple 'flag large transactions' rule would miss most fraud."),
        ],
    },
    {
        "number": 4,
        "title": "Temporal Fraud Patterns",
        "subtitle": "Notebook Section 7: Fraud Rate by Hour and Day of Week",
        "narrative": (
            "Temporal analysis uncovers clear cyclical patterns. Fraud rates peak during "
            "early morning hours (7\u20139 AM) and on weekends (Friday, Saturday, Sunday), while "
            "midweek afternoons (Tuesday\u2013Thursday, 1\u20133 PM) show the lowest rates. These "
            "patterns suggest that fraudsters exploit periods of lower monitoring and higher "
            "shopping volume."
        ),
        "matrix": CH4_MATRIX,
        "figures": [
            ("temporal_patterns.png", "Figure 6: Fraud rate by hour of day (left) and day of week (right), with red dashed line indicating the overall 3.50% baseline."),
        ],
        "callouts": [
            ("insight", "Immediate operational value: the bank can tighten alert thresholds during 7\u20139 AM and weekends even before deploying the ML model."),
        ],
    },
    {
        "number": 5,
        "title": "Cost Modeling & Strategic Outputs",
        "subtitle": "Notebook Sections 8\u201310: Cost Assumptions, Key Findings, Summary Statistics",
        "narrative": (
            "The EDA concludes by translating analytical findings into a business decision "
            "framework. Establishing the False Negative cost at $227 (full economic impact: $75 transaction "
            "loss + $27 chargeback + $50 investigation + $75 reputational damage) versus $10 for "
            "False Positive (manual review), the analysis establishes a 22.7:1 cost asymmetry. This ratio becomes the cornerstone of Phase 3 "
            "threshold optimization, ensuring the model prioritizes catching fraud over "
            "minimizing false alarms."
        ),
        "matrix": CH5_MATRIX,
        "figures": [],
        "callouts": [
            ("business", "The 22.7:1 cost ratio means the bank should tolerate up to 22.7 false alarms for every fraud it catches. This drives the model to a 76% recall operating point \u2014 catching 3 out of 4 frauds."),
        ],
    },
]

SUMMARY_TABLE_DATA = [
    ("Total Transactions", "590,540"),
    ("Fraud Transactions", "20,663"),
    ("Fraud Rate (%)", "3.50"),
    ("Class Imbalance Ratio", "1:27"),
    ("Total Features", "434"),
    ("Features >50% Missing", "214 (49.3%)"),
    ("Avg Fraud Amount ($)", "149.24"),
    ("Median Fraud Amount ($)", "75.00"),
    ("Avg Legitimate Amount ($)", "134.51"),
    ("Peak Fraud Hour", "7:00 AM"),
    ("Peak Fraud Day", "Friday"),
    ("FN Cost (missed fraud)", "$227.00"),
    ("FP Cost (false alarm)", "$10.00"),
    ("Cost Ratio (FN:FP)", "22.7 : 1"),
]

GLOSSARY = [
    ("Class Imbalance", "When one category (e.g., fraud) is much rarer than the other (e.g., legitimate). Here, fraud is only 3.5% of all transactions."),
    ("Correlation", "A statistical measure (0 to 1) of how strongly two variables move together. Higher means more connected."),
    ("EDA", "Exploratory Data Analysis \u2014 the process of examining data to discover patterns, anomalies, and insights before building models."),
    ("False Negative (FN)", "A fraud transaction that the system missed \u2014 it said 'legitimate' but it was actually fraud."),
    ("False Positive (FP)", "A legitimate transaction that the system flagged \u2014 it said 'fraud' but it was actually fine."),
    ("Feature", "A column or variable in the dataset used to make predictions. Examples: transaction amount, time of day."),
    ("Feature Engineering", "Creating new, more informative variables from raw data. Example: computing 'transactions per hour' from timestamps."),
    ("IEEE-CIS Dataset", "A large public dataset of anonymized e-commerce transactions released by the IEEE Computational Intelligence Society for fraud detection research."),
    ("Left Join", "A database operation that combines two tables, keeping all rows from the left table and matching rows from the right table where available."),
    ("Missing Data", "Blank or null values in the dataset. Features with high missingness (>50%) are unreliable for modeling."),
    ("PR-AUC", "Precision-Recall Area Under Curve \u2014 a metric suited for imbalanced datasets that measures how well the model separates fraud from legitimate transactions."),
    ("Recall", "The percentage of actual frauds that the model correctly identifies. 76% recall means catching 3 out of 4 frauds."),
    ("Scale_pos_weight", "An XGBoost parameter that tells the model how much more important fraud cases are than legitimate ones, compensating for class imbalance."),
    ("Threshold", "The probability cutoff above which the model flags a transaction as fraud. Lower thresholds catch more fraud but create more false alarms."),
    ("V-Features", "Anonymized features in the IEEE-CIS dataset (V1\u2013V339). Their real-world meaning is not disclosed for privacy reasons."),
    ("XGBoost", "Extreme Gradient Boosting \u2014 a powerful machine learning algorithm that builds many small decision trees and combines them for accurate predictions."),
]


# ── Helper Functions ───────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background color to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    """Set cell margins in twips (1/20 of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_paragraph_spacing(paragraph, before=0, after=0, line=240):
    """Set paragraph spacing in twips."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_formatted_text(cell, text, font_name="Calibri", font_size=10,
                       bold=False, color_hex=None):
    """Add formatted text to a table cell, clearing existing content."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=20, after=20, line=240)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def add_callout_box(doc, text, box_type="insight"):
    """Add a colored callout box paragraph."""
    colors = {
        "insight": (C_INSIGHT_BG, C_INSIGHT_BORDER, "Key Insight"),
        "business": (C_BUSINESS_BG, C_BUSINESS_BORDER, "Business Impact"),
    }
    bg_color, border_color, label = colors.get(box_type, colors["insight"])

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=120, line=264)

    # Shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), bg_color)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)

    # Left border
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), border_color)
    left.set(qn("w:space"), "4")
    pBdr.append(left)
    pPr.append(pBdr)

    # Indentation
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:right"), "360")
    pPr.append(ind)

    # Label run
    label_run = p.add_run(f"{label}: ")
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(10)
    label_run.font.bold = True
    label_run.font.color.rgb = RGBColor.from_string(border_color)

    # Content run
    content_run = p.add_run(text)
    content_run.font.name = "Calibri"
    content_run.font.size = Pt(10)
    content_run.font.italic = True
    content_run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)


def add_figure(doc, image_filename, caption_text):
    """Add an image with caption if the file exists."""
    img_path = FIGURES_DIR / image_filename
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cap, before=40, after=200, line=240)
        cap_run = cap.add_run(caption_text)
        cap_run.font.name = "Calibri"
        cap_run.font.size = Pt(9)
        cap_run.font.italic = True
        cap_run.font.color.rgb = RGBColor.from_string("666666")
    else:
        p = doc.add_paragraph(f"[Figure not available: {image_filename} \u2014 run notebook to generate]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor.from_string("999999")


def add_matrix_table(doc, matrix_data):
    """Add a 4-layer x 3-perspective matrix table."""
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set preferred table width to full page width
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9360")  # ~6.5 inches in twips
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Set table borders
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    # Header row
    headers = ["Layer", "Technical\n(BDS Colleague)", "Business\n(Manager / Regulator)", "Simple\n(Grandmother)"]
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, layer in enumerate(LAYERS):
        row_idx = i + 1
        # Layer label cell
        layer_cell = table.cell(row_idx, 0)
        set_cell_shading(layer_cell, C_MED_BLUE)
        set_cell_margins(layer_cell, top=60, bottom=60, left=100, right=100)
        add_formatted_text(layer_cell, layer, font_size=9, bold=True, color_hex=C_WHITE)
        layer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Perspective cells
        for j in range(3):
            cell = table.cell(row_idx, j + 1)
            bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            text = matrix_data.get((i, j), "")
            add_formatted_text(cell, text, font_size=9)

    # Add spacing after table
    doc.add_paragraph("")

    return table


def add_section_heading(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)
    return heading


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


# ── Document Construction ──────────────────────────────────────────────

def add_cover_page(doc):
    """Create the cover page."""
    # Add empty paragraphs for vertical centering
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Agent Fraud Sentinel")
    run.font.name = "Calibri"
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("EDA Deep Dive: 4 Layers \u00d7 3 Perspectives")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    # Separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("\u2500" * 50)
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)
    run.font.size = Pt(12)

    # Notebook reference
    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ref.add_run("Notebook: 01_eda_fraud_patterns.ipynb")
    run.font.name = "Consolas"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)

    # Dataset info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("IEEE-CIS Fraud Detection Dataset\n590,540 Transactions | 434 Features | 3.50% Fraud Rate")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)

    # Date and project
    for _ in range(3):
        doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("BAFS Project \u2014 February 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("888888")

    add_page_break(doc)


def add_toc_placeholder(doc):
    """Add a Table of Contents placeholder field."""
    add_section_heading(doc, "Table of Contents", level=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    # Insert TOC field code
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar_separate)

    # Placeholder text
    placeholder_run = p.add_run("(Right-click here and select 'Update Field' to generate Table of Contents)")
    placeholder_run.font.name = "Calibri"
    placeholder_run.font.size = Pt(10)
    placeholder_run.font.italic = True
    placeholder_run.font.color.rgb = RGBColor.from_string("999999")

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    placeholder_run._r.append(fldChar_end)

    add_page_break(doc)


def add_framework_explanation(doc):
    """Add the 4x3 framework explanation section."""
    add_section_heading(doc, "The 4-Layer \u00d7 3-Perspective Framework", level=2)

    p = doc.add_paragraph()
    run = p.add_run(
        "This document examines each section of the EDA notebook through two dimensions: "
        "four analytical layers and three audience perspectives. This framework ensures "
        "complete understanding \u2014 from raw technical detail to business impact \u2014 "
        "accessible to any reader regardless of their background."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    # Layers explanation
    add_section_heading(doc, "The Four Layers", level=3)
    layers_desc = [
        ("Layer 1 \u2014 WHAT did I do?", "Describes the concrete actions taken: what was computed, plotted, or measured."),
        ("Layer 2 \u2014 WHY did I do it?", "Explains the motivation and reasoning behind each analytical step."),
        ("Layer 3 \u2014 HOW does it work?", "Details the technical mechanics: methods, formulas, and code logic."),
        ("Layer 4 \u2014 WHAT does the bank gain?", "Translates findings into business value: cost savings, risk reduction, and strategic advantage."),
    ]
    for title, desc in layers_desc:
        p = doc.add_paragraph()
        bold_run = p.add_run(title + "  ")
        bold_run.font.name = "Calibri"
        bold_run.font.size = Pt(11)
        bold_run.font.bold = True
        bold_run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)
        desc_run = p.add_run(desc)
        desc_run.font.name = "Calibri"
        desc_run.font.size = Pt(11)

    # Perspectives explanation
    add_section_heading(doc, "The Three Perspectives", level=3)
    persp_desc = [
        ("Technical (BDS Colleague)", "Uses data science terminology, references specific functions and parameters, assumes ML knowledge."),
        ("Business (Manager / Regulator)", "Focuses on process, strategy, compliance, and dollar impact. Assumes no coding knowledge."),
        ("Simple (Grandmother)", "Uses everyday language, analogies, and metaphors. Assumes no technical or financial background."),
    ]
    for title, desc in persp_desc:
        p = doc.add_paragraph()
        bold_run = p.add_run(title + "  ")
        bold_run.font.name = "Calibri"
        bold_run.font.size = Pt(11)
        bold_run.font.bold = True
        bold_run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)
        desc_run = p.add_run(desc)
        desc_run.font.name = "Calibri"
        desc_run.font.size = Pt(11)


def add_chapter(doc, chapter_data):
    """Add a complete chapter with narrative, matrix, figures, and callouts."""
    num = chapter_data["number"]
    title = chapter_data["title"]

    add_page_break(doc)
    add_section_heading(doc, f"Chapter {num}: {title}", level=1)

    # Subtitle
    if chapter_data.get("subtitle"):
        p = doc.add_paragraph()
        run = p.add_run(chapter_data["subtitle"])
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    # Narrative
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=200, line=276)
    run = p.add_run(chapter_data["narrative"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    # Matrix table
    add_section_heading(doc, "Analysis Matrix", level=2)
    add_matrix_table(doc, chapter_data["matrix"])

    # Callout boxes
    for box_type, text in chapter_data.get("callouts", []):
        add_callout_box(doc, text, box_type)

    # Figures
    for fig_filename, caption in chapter_data.get("figures", []):
        add_figure(doc, fig_filename, caption)


def add_summary_statistics_table(doc):
    """Add the EDA summary statistics table in Chapter 5."""
    add_section_heading(doc, "EDA Summary Statistics", level=2)

    table = doc.add_table(rows=len(SUMMARY_TABLE_DATA) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    # Header
    for j, header in enumerate(["Metric", "Value"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, (metric, value) in enumerate(SUMMARY_TABLE_DATA):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        for j, text in enumerate([metric, value]):
            cell = table.cell(i + 1, j)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=100, right=100)
            add_formatted_text(cell, text, font_size=10, bold=(j == 0))
            if j == 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_appendix_gallery(doc):
    """Add Appendix A: Visualization Gallery."""
    add_page_break(doc)
    add_section_heading(doc, "Appendix A: Visualization Gallery", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "All visualizations generated by the EDA notebook, presented in sequence. "
        "These figures are saved in data/processed/ and can be regenerated by running the notebook."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    figures = [
        ("class_distribution.png", "Class Distribution: Bar chart and pie chart showing the 96.50% legitimate vs 3.50% fraud split across 590,540 transactions."),
        ("top_correlations.png", "Top 10 Feature Correlations: Horizontal bar chart ranking features by absolute Pearson correlation with isFraud. V257 leads at 0.383."),
        ("missing_data.png", "Missing Data Analysis: Histogram of missing percentages across all 434 features (left) and top 15 most-missing features (right)."),
        ("amount_distribution_dollars.png", "Transaction Amount ($0\u2013$500): Side-by-side histograms for legitimate (green) and fraud (red) transactions with median lines."),
        ("amount_distribution.png", "Transaction Amount (Advanced): Log-scale density overlay (left) and box plot (right) comparing distributions by fraud status."),
        ("temporal_patterns.png", "Temporal Patterns: Fraud rate by hour of day (left) and day of week (right) with 3.50% baseline reference line."),
    ]

    for fig_file, caption in figures:
        add_figure(doc, fig_file, caption)


def add_appendix_glossary(doc):
    """Add Appendix B: Glossary."""
    add_page_break(doc)
    add_section_heading(doc, "Appendix B: Glossary", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "Key terms used throughout this document, defined for non-technical readers."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    set_paragraph_spacing(p, after=200)

    table = doc.add_table(rows=len(GLOSSARY) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    # Header
    for j, header in enumerate(["Term", "Definition"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)

    # Rows
    for i, (term, definition) in enumerate(GLOSSARY):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        cell_term = table.cell(i + 1, 0)
        set_cell_shading(cell_term, bg)
        set_cell_margins(cell_term, top=30, bottom=30, left=100, right=100)
        add_formatted_text(cell_term, term, font_size=10, bold=True)

        cell_def = table.cell(i + 1, 1)
        set_cell_shading(cell_def, bg)
        set_cell_margins(cell_def, top=30, bottom=30, left=100, right=100)
        add_formatted_text(cell_def, definition, font_size=10)


def setup_header_footer(doc):
    """Add header and footer to all sections."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        h_para.text = ""
        run_left = h_para.add_run("Agent Fraud Sentinel \u2014 EDA Analysis")
        run_left.font.name = "Calibri"
        run_left.font.size = Pt(8)
        run_left.font.color.rgb = RGBColor.from_string("999999")
        h_para.add_run("\t\t")
        run_right = h_para.add_run("BAFS")
        run_right.font.name = "Calibri"
        run_right.font.size = Pt(8)
        run_right.font.bold = True
        run_right.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        f_para.text = ""
        f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_page = f_para.add_run("Page ")
        run_page.font.name = "Calibri"
        run_page.font.size = Pt(8)
        run_page.font.color.rgb = RGBColor.from_string("999999")

        # Page number field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run_page._r.append(fldChar1)

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run_page._r.append(instrText)

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        run_page._r.append(fldChar2)

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run_page._r.append(fldChar3)


def main():
    """Generate the complete Word document."""
    print("Generating EDA Analysis Matrix document...")

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Heading styles
    for i in range(1, 4):
        h_style = doc.styles[f"Heading {i}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

    # ── Build Document ─────────────────────────────────────────────
    add_cover_page(doc)
    add_toc_placeholder(doc)

    # Chapter 0: Executive Overview (with framework explanation)
    ch0 = ALL_CHAPTERS[0]
    add_section_heading(doc, f"Chapter 0: {ch0['title']}", level=1)

    # Subtitle
    p = doc.add_paragraph()
    run = p.add_run(ch0["subtitle"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    # Narrative
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=200, line=276)
    run = p.add_run(ch0["narrative"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    # Framework explanation (unique to Chapter 0)
    add_framework_explanation(doc)

    # Chapter 0 matrix
    add_section_heading(doc, "Full-Project Analysis Matrix", level=2)
    add_matrix_table(doc, ch0["matrix"])

    for box_type, text in ch0.get("callouts", []):
        add_callout_box(doc, text, box_type)

    # Chapters 1-5
    for chapter_data in ALL_CHAPTERS[1:]:
        add_chapter(doc, chapter_data)

    # Summary statistics table at end of Chapter 5
    add_summary_statistics_table(doc)

    # Appendices
    add_appendix_gallery(doc)
    add_appendix_glossary(doc)

    # Header/Footer
    setup_header_footer(doc)

    # Save
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"File size: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
