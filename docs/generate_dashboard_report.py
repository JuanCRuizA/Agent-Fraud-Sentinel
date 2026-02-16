"""
Generate Word Document: Streamlit Dashboard Notebook 05 - 4 Layers x 3 Perspectives
Agent Fraud Sentinel (BAFS) Project

Produces: docs/dashboard_05_analysis_matrix.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent.parent
FIGURES_DIR = BASE_DIR / "figures" / "shap"
OUTPUT_PATH = BASE_DIR / "docs" / "dashboard_05_analysis_matrix.docx"

# ── Colors ─────────────────────────────────────────────────────────────
C_DARK_BLUE = "1B4F72"
C_MED_BLUE = "2980B9"
C_LIGHT_BLUE = "D6EAF8"
C_FRAUD_RED = "E74C3C"
C_LEGIT_GREEN = "2ECC71"
C_LIGHT_GRAY = "F2F3F4"
C_WHITE = "FFFFFF"
C_BLACK = "000000"
C_DARK_GRAY = "2C3E50"
C_INSIGHT_BG = "D6EAF8"
C_INSIGHT_BORDER = "2980B9"
C_BUSINESS_BG = "D5F5E3"
C_BUSINESS_BORDER = "27AE60"
C_CAUTION_BG = "FADBD8"
C_CAUTION_BORDER = "E74C3C"

LAYERS = ["WHAT did I do?", "WHY did I do it?", "HOW does it work?", "WHAT does the bank gain?"]

# ══════════════════════════════════════════════════════════════════════
#  CONTENT: ALL CHAPTER MATRICES
# ══════════════════════════════════════════════════════════════════════

CH0_MATRIX = {
    (0, 0): (
        "Built a professional, interactive Streamlit dashboard (dashboard_app.py, "
        "~470 lines) that consolidates all BAFS project outputs into 4 tabs: "
        "Executive Summary (KPIs, cost analysis), Model Performance (confusion matrix, "
        "ROC/PR curves, cost-benefit table), Case Study Explorer (6 SHAP case studies "
        "with plain-English explanations), and Regulatory Compliance (SR 11-7 checklist, "
        "fair lending review, model governance, audit trail). Interactive sidebar "
        "with threshold slider and sample size selector updates all metrics dynamically."
    ),
    (0, 1): (
        "Created a control center that makes the fraud detection system accessible to "
        "three audiences: executives see KPIs and cost trade-offs, data scientists see "
        "detailed model performance, fraud analysts explore individual case studies, "
        "and compliance officers review regulatory documentation. All metrics update "
        "in real-time when the user adjusts the risk threshold."
    ),
    (0, 2): (
        "We built a 'command center' for the fraud detection system. It has four screens: "
        "one for the boss (how much fraud are we catching?), one for the tech team (how "
        "accurate is the model?), one for the investigators (why was this purchase flagged?), "
        "and one for the inspectors (are we following the rules?). Anyone can adjust the "
        "sensitivity dial and see how it changes everything."
    ),
    (1, 0): (
        "A fraud detection model is only as valuable as its ability to communicate results "
        "to stakeholders. Notebooks are not suitable for non-technical stakeholders. "
        "Streamlit provides a low-code framework for deploying data applications with "
        "built-in interactivity (sliders, dropdowns, dynamic charts). The 4-tab design "
        "matches the four stakeholder groups: executives, data scientists, fraud analysts, "
        "and compliance officers."
    ),
    (1, 1): (
        "The best fraud model in the world is useless if nobody can understand or use it. "
        "Jupyter notebooks are for data scientists, not for bank managers. The dashboard "
        "translates months of analysis into an interactive tool that any stakeholder can "
        "use. The threshold slider lets stakeholders explore 'what if we were more/less "
        "strict?' scenarios in real-time."
    ),
    (1, 2): (
        "Imagine doing all the detective work but writing the report in a language nobody "
        "speaks. The dashboard is the translation: it takes everything we've built and "
        "presents it in a way that anyone at the bank can understand and use. The "
        "sensitivity dial lets them experiment and see what happens."
    ),
    (2, 0): (
        "Streamlit with st.set_page_config(layout='wide'). @st.cache_resource for model "
        "loading (singleton), @st.cache_data for data/predictions (hash-based). Custom "
        "CSS for banking aesthetic (#1a365d headers, #f8f9fa background). st.sidebar for "
        "navigation + filters (st.radio, st.slider, st.selectbox). st.columns for "
        "responsive layouts. st.metric for KPI cards. matplotlib/seaborn for charts. "
        "PIL for loading SHAP figures. sklearn.metrics for dynamic computation."
    ),
    (2, 1): (
        "A web application built with Streamlit (Python framework for data dashboards). "
        "The model and data are loaded once and cached for performance. The sidebar "
        "contains navigation (4 tabs) and global controls (threshold slider, sample "
        "size). Charts are generated dynamically with matplotlib. SHAP figures from "
        "Phase 4 are loaded as images. A consistent footer appears on every page."
    ),
    (2, 2): (
        "We built a website using a tool called Streamlit that turns Python code into "
        "an interactive web page. The model is loaded once (so it's fast), and when "
        "you move the slider, all the numbers and charts update instantly. It has a "
        "sidebar for navigating between the four screens."
    ),
    (3, 0): (
        "Production-ready dashboard deployed via Streamlit, deployable to Streamlit "
        "Cloud with one-click. All visualizations are dynamic (threshold-responsive). "
        "File artifacts: dashboard_app.py (~470 lines), requirements.txt. Portfolio "
        "piece demonstrating full-stack ML deployment: data \u2192 model \u2192 "
        "explainability \u2192 interactive application."
    ),
    (3, 1): (
        "The bank gets an operational tool, not just a report. Executives can explore "
        "cost trade-offs by adjusting the threshold. Analysts can study case studies. "
        "Compliance officers can review governance documentation. And as a portfolio "
        "piece, it demonstrates end-to-end data science capabilities for banking roles."
    ),
    (3, 2): (
        "The bank gets a real tool that anyone can use, not just a static report. "
        "The boss can play with the sensitivity dial and see costs change. "
        "Investigators can study example cases. And inspectors can check that "
        "everything is properly documented. It's also a showcase of the entire project."
    ),
}

CH1_MATRIX = {
    (0, 0): (
        "Verified environment dependencies: streamlit, pandas, numpy, matplotlib, "
        "seaborn, joblib, sklearn, xgboost, PIL (Pillow). Confirmed model artifacts "
        "exist: xgboost_final.pkl, scaler.pkl, threshold_config.pkl. Confirmed test "
        "data: test.csv (118,108 rows). Confirmed 6 SHAP figures in figures/shap/. "
        "Designed 4-tab architecture with sidebar navigation and global filters. "
        "Documented file dependencies and design principles."
    ),
    (0, 1): (
        "Checked that all required software and data files are present before building "
        "the dashboard. Designed the layout: a sidebar for navigation and controls, "
        "and a main area with 4 tabs for different audiences. Documented the design "
        "principles: professional banking look, interactive filters, cost-sensitive "
        "framing, and regulatory awareness."
    ),
    (0, 2): (
        "We made sure all the ingredients were ready before cooking: the software, "
        "the model, the data, and the charts from the previous phase. Then we drew "
        "a blueprint of what the dashboard would look like: a sidebar with controls "
        "and four main screens."
    ),
    (1, 0): (
        "Dependency verification prevents runtime failures in production. The "
        "architecture design follows the 4-audience pattern: executive (KPIs), "
        "technical (model metrics), operational (case studies), and compliance "
        "(governance). The sidebar-plus-tabs pattern is standard for data dashboards. "
        "Design principles ensure banking-appropriate aesthetics (no emojis in the UI, "
        "blue/gray palette, professional typography)."
    ),
    (1, 1): (
        "Checking dependencies upfront prevents embarrassing crashes during a demo. "
        "The 4-tab design ensures each audience finds what they need without wading "
        "through irrelevant content. The professional banking aesthetic builds "
        "credibility \u2014 a colorful, casual dashboard would undermine trust in "
        "a financial application."
    ),
    (1, 2): (
        "Before building a house, you check that all the materials are ready. Before "
        "building the dashboard, we checked that all the software and data files "
        "exist. Then we designed each room (tab) for a specific person: one for "
        "the boss, one for the tech team, one for investigators, one for inspectors."
    ),
    (2, 0): (
        "importlib.import_module() for dependency checking. Path.exists() for "
        "artifact verification. Architecture diagram: sidebar (st.radio for tabs, "
        "st.slider for threshold 0.0-1.0, st.selectbox for sample size) + main area "
        "(4 conditional pages). Design constants: FEATURE_LABELS dict, FOOTER_HTML "
        "template, color scheme via custom CSS. %%writefile magic writes dashboard_app.py "
        "and requirements.txt from the notebook."
    ),
    (2, 1): (
        "Each software package is checked to ensure it's installed. Each data file "
        "is verified to exist. The architecture uses a sidebar for navigation (4 tabs) "
        "and controls (threshold slider, sample size dropdown). The notebook writes "
        "the complete application code to a standalone file (dashboard_app.py) that "
        "can be run independently."
    ),
    (2, 2): (
        "We checked each ingredient one by one, verified the data files exist, "
        "and drew the blueprint. The notebook writes the complete dashboard code "
        "to a separate file that can be launched on its own \u2014 like a recipe "
        "card that produces a complete meal."
    ),
    (3, 0): (
        "Clean dependency management ensures reproducible deployment. The 4-tab "
        "architecture scales naturally: additional audiences or analyses can be added "
        "as new tabs. The %%writefile pattern enables notebook-as-documentation: the "
        "notebook explains the design while producing the deployable artifact. "
        "requirements.txt enables one-command deployment on Streamlit Cloud."
    ),
    (3, 1): (
        "The bank gets a verified, reproducible setup: every dependency checked, "
        "every file confirmed. The architecture is designed to grow \u2014 new tabs "
        "can be added for future needs (e.g., a monitoring tab for model drift). "
        "The deployment file (requirements.txt) enables cloud hosting with minimal "
        "effort."
    ),
    (3, 2): (
        "Everything is verified and ready to go. The design can grow easily \u2014 "
        "like adding new rooms to a house. The setup file means the dashboard can "
        "be published online with just a few clicks."
    ),
}

CH2_MATRIX = {
    (0, 0): (
        "Built Executive Summary tab with: 4 KPI metric cards (Fraud Detected with "
        "recall, False Positive Rate, Fraud Prevented in dollars, Total Operational "
        "Cost), performance table (recall, precision, F1, FPR, threshold), risk score "
        "distribution histogram (overlapping fraud/legitimate with threshold line), and "
        "cost analysis row (Missed Fraud Cost, False Alarm Cost, Savings vs No Model). "
        "All metrics update dynamically when the sidebar threshold changes."
    ),
    (0, 1): (
        "Created a one-page executive overview showing the four numbers that matter "
        "most: how much fraud we catch, how many false alarms we generate, how much "
        "fraud we prevent in dollars, and the total operational cost. A histogram "
        "shows the distribution of fraud scores, and a cost breakdown shows where "
        "the money goes. Moving the threshold slider instantly updates everything."
    ),
    (0, 2): (
        "We built a 'dashboard for the boss' with four big numbers at the top: "
        "fraud caught, false alarms, money saved, and total cost. Below that, a "
        "chart shows how the system scores each transaction, and a cost breakdown "
        "shows where the money goes. Moving the slider changes all the numbers "
        "instantly."
    ),
    (1, 0): (
        "Executive stakeholders need high-level KPIs, not technical metrics. The "
        "4 KPI cards follow the dashboard best practice of showing key metrics at a "
        "glance. The risk score histogram provides intuition about class separability. "
        "The cost analysis ties everything to dollar values, which is the language "
        "of executive decision-making. Dynamic updating enables 'what-if' scenario "
        "analysis during budget discussions."
    ),
    (1, 1): (
        "Executives don't need precision-recall curves. They need: 'How much fraud "
        "are we catching?' and 'What does it cost?' The threshold slider lets them "
        "explore the fundamental trade-off: catching more fraud means more false "
        "alarms. The cost section translates this into dollars, enabling informed "
        "budget decisions."
    ),
    (1, 2): (
        "The boss wants to know: 'Is it working? How much does it cost?' The four "
        "big numbers answer both questions. The slider lets them ask 'What if we "
        "were more strict?' or 'What if we were more lenient?' and see the answer "
        "in dollars immediately."
    ),
    (2, 0): (
        "st.metric() for KPI cards with delta indicators. st.columns(4) for side-by-side "
        "layout. Confusion matrix computed from (y_filt, y_pred_filt) where y_pred = "
        "(scores >= threshold).astype(int). Histogram: plt.hist() with density=True, "
        "alpha=0.6 for overlapping fraud/legitimate distributions. Cost computation: "
        "missed_fraud = FN * $75, review_cost = FP * $10, savings = (total_fraud * $75) "
        "- missed_fraud. All recomputed on every threshold/sample change."
    ),
    (2, 1): (
        "Four KPI cards are displayed side by side using Streamlit's column layout. "
        "The histogram overlaps fraud scores (red) and legitimate scores (blue) to "
        "show how well the model separates them. A vertical line marks the current "
        "threshold. Cost calculations use the established ratios: $75 per missed fraud, "
        "$10 per false alarm. Everything recalculates when the slider moves."
    ),
    (2, 2): (
        "Four number boxes sit side by side at the top. Below them, a chart shows "
        "how fraud (red) and honest (blue) transactions are distributed, with a "
        "line showing where the alarm is set. The cost section multiplies missed "
        "frauds by $75 and false alarms by $10 to get the total cost."
    ),
    (3, 0): (
        "Interactive executive dashboard that enables data-driven threshold selection. "
        "The 'Savings vs No Model' metric quantifies the model's value proposition in "
        "dollars: how much the system saves compared to having no fraud detection. "
        "Dynamic updates support real-time scenario analysis during stakeholder meetings."
    ),
    (3, 1): (
        "Executives can make informed decisions about fraud detection trade-offs using "
        "real numbers. The 'Savings vs No Model' metric proves the system's ROI. "
        "The interactive threshold slider transforms a static report into a decision-"
        "support tool that can be used during budget and strategy meetings."
    ),
    (3, 2): (
        "The boss can see exactly how much money the fraud system saves compared "
        "to having no system at all. And by moving the slider, they can decide "
        "how strict they want the system to be, seeing the cost impact immediately."
    ),
}

CH3_MATRIX = {
    (0, 0): (
        "Built Model Performance tab with: confusion matrix heatmap with cost overlay "
        "(TN=$0, FP=count\u00d7$10, FN=count\u00d7$75, TP='Prevented'), ROC curve with "
        "operating point marker and AUC, Precision-Recall curve with baseline and "
        "operating point, SHAP feature importance bar chart (loaded from Phase 4 PNG), "
        "and cost-benefit table across 8 threshold values (0.20 to 0.90) showing recall, "
        "precision, TP, FP, missed frauds, and total cost."
    ),
    (0, 1): (
        "Created a deep technical evaluation page for model validators. The confusion "
        "matrix shows both counts AND dollar costs. ROC and PR curves display the "
        "full trade-off landscape with the current operating point marked. Feature "
        "importance comes directly from the SHAP analysis. The cost-benefit table "
        "lets stakeholders compare 8 different threshold options side by side."
    ),
    (0, 2): (
        "We built a page for the tech team showing: a table of right vs wrong answers "
        "(with dollar costs), two curves showing the trade-off between catching fraud "
        "and making false alarms, a chart of which clues matter most, and a comparison "
        "table showing what happens at 8 different alarm settings."
    ),
    (1, 0): (
        "Model validators require standardized evaluation charts: confusion matrix, "
        "ROC, and PR curves. The cost overlay on the confusion matrix bridges the gap "
        "between statistical metrics and business impact. PR-AUC is more informative "
        "than ROC-AUC for imbalanced data (3.44% fraud). The cost-benefit table "
        "enables systematic threshold selection by comparing all options side by side."
    ),
    (1, 1): (
        "Data science teams need detailed performance analysis to validate the model. "
        "The confusion matrix with dollar overlays helps them communicate results to "
        "business stakeholders. ROC and PR curves are standard tools for model "
        "comparison. The cost-benefit table is the bridge between model performance "
        "and operational decision-making."
    ),
    (1, 2): (
        "The tech team needs detailed charts to make sure the model is working "
        "correctly. The dollar amounts on the confusion matrix help them explain "
        "results to non-technical people. The comparison table lets them see what "
        "happens if they turn the alarm up or down."
    ),
    (2, 0): (
        "sklearn confusion_matrix() \u2192 sns.heatmap() with custom annot labels "
        "including costs. roc_curve() + auc() for ROC; precision_recall_curve() + "
        "auc() for PR. Operating point: recall_score/precision_score at current "
        "threshold, plotted as red scatter. SHAP image: PIL.Image.open() with "
        "st.image(use_container_width=True). Fallback: model.feature_importances_ "
        "if SHAP PNG missing. Cost-benefit: loop over 8 thresholds computing "
        "TP/FP/FN/recall/precision/cost."
    ),
    (2, 1): (
        "The confusion matrix is rendered as a heatmap with count AND cost annotations. "
        "ROC and PR curves are computed from the model's probability scores and plotted "
        "with the current operating point highlighted. The SHAP feature importance chart "
        "is loaded from the Phase 4 output file. If unavailable, a fallback using the "
        "model's built-in importance is displayed. The cost-benefit table evaluates "
        "8 threshold values."
    ),
    (2, 2): (
        "The confusion matrix is drawn as a color-coded table showing right and wrong "
        "answers, with dollar costs added. Two curves show the detection trade-off. "
        "The feature importance chart is borrowed from the previous phase. The "
        "comparison table tests 8 different alarm levels and shows the results."
    ),
    (3, 0): (
        "Complete model validation dashboard with cost-aware evaluation. The operating "
        "point visualization on ROC/PR curves enables immediate assessment of where "
        "the model sits on the trade-off frontier. The cost-benefit table provides "
        "a systematic framework for threshold selection meetings. SHAP integration "
        "connects Phases 4 and 5 into a cohesive analytics platform."
    ),
    (3, 1): (
        "The bank's data science team gets a comprehensive validation dashboard that "
        "speaks both technical and business language. The cost-benefit table enables "
        "structured threshold selection: for any given threshold, the team can see "
        "exactly how many frauds will be caught and at what cost. The SHAP integration "
        "connects model performance to explainability."
    ),
    (3, 2): (
        "The tech team gets a complete set of tools to check the model's performance. "
        "The comparison table makes it easy to pick the right alarm setting by showing "
        "the costs and results of each option side by side."
    ),
}

CH4_MATRIX = {
    (0, 0): (
        "Built Case Study Explorer tab with dropdown selector for 6 cases from Phase 4. "
        "Each case displays: 3 KPI metrics (fraud score, model decision, actual outcome), "
        "feature values in 2-column layout with human-readable formatting ($17.52, 6:00 AM, "
        "Yes/No), SHAP waterfall plot (loaded from PNG), plain-English explanation paragraph, "
        "key risk drivers as bullet list, and improvement recommendations (for FN cases). "
        "Cases: TP clear, TP velocity, FN missed, FP flagged, auto-block, borderline."
    ),
    (0, 1): (
        "Created an interactive case study browser where fraud analysts can select any "
        "of the 6 representative cases and see the complete story: what the transaction "
        "looked like, what the model decided, why it made that decision (in plain English), "
        "and what the actual outcome was. False negative cases include recommendations "
        "for model improvement."
    ),
    (0, 2): (
        "We built a page where investigators can flip through 6 example cases using a "
        "dropdown menu. For each case, they see: the transaction details (amount, time, "
        "etc.), the system's decision, a plain-English explanation of why, and whether "
        "the system was right or wrong. For cases where the system was wrong, we "
        "suggest how to improve."
    ),
    (1, 0): (
        "Case studies demonstrate local explainability \u2014 a regulatory requirement "
        "(GDPR Art. 22, SR 11-7). The 6 cases strategically cover all confusion matrix "
        "quadrants (TP, FP, FN) plus operational scenarios (auto-block, borderline). "
        "The plain-English explanations serve as templates for customer communication "
        "during dispute resolution. Risk drivers distill SHAP values into actionable "
        "bullet points."
    ),
    (1, 1): (
        "Fraud analysts need to understand why specific transactions are flagged. "
        "The case studies teach new analysts how the model 'thinks.' For customer "
        "disputes, the plain-English explanations provide ready-made responses. "
        "For compliance audits, the detailed feature breakdown demonstrates that "
        "individual decisions can be explained and justified."
    ),
    (1, 2): (
        "The investigators need to see real examples of how the system works. "
        "The 6 examples cover every situation: correctly caught fraud, missed fraud, "
        "falsely blocked honest customers, and borderline cases. Each example has "
        "a clear explanation that could be read to a customer on the phone."
    ),
    (2, 0): (
        "st.selectbox() for case selection from dict keys. st.metric() row for score/"
        "decision/actual. Feature display: st.columns(2) with st.markdown(f'**{k}:** {v}'). "
        "SHAP waterfall: PIL.Image.open(FIGURES_PATH/'shap_waterfall_cases.png') loaded "
        "conditionally with fallback st.info(). Case data stored as nested dict with "
        "keys: score, actual, decision, features, explanation, drivers, improvement."
    ),
    (2, 1): (
        "A dropdown menu lets the user select which case to examine. Three metric "
        "cards show the fraud score, model decision, and actual outcome. Feature "
        "values are displayed in a two-column layout. The SHAP waterfall image "
        "from Phase 4 is loaded and displayed. The plain-English explanation and "
        "risk drivers are rendered as text and bullet points."
    ),
    (2, 2): (
        "You pick an example from a dropdown list, and the page shows all the "
        "details: three key numbers at the top, the transaction's characteristics "
        "in a neat layout, a chart showing why the decision was made, a plain-English "
        "explanation, and a list of the main risk factors."
    ),
    (3, 0): (
        "Production-ready explainability interface. The case study template is "
        "extensible: new cases can be added by appending to the dict. The SHAP "
        "waterfall integration demonstrates full traceability from model output "
        "to visual explanation to plain-English summary. The improvement "
        "recommendations for FN cases create a feedback loop for model iteration."
    ),
    (3, 1): (
        "The bank gets an analyst training tool and a dispute resolution resource. "
        "New fraud analysts can learn the model's behavior through the 6 case studies. "
        "Customer service teams can reference the explanations when handling disputes. "
        "The improvement recommendations document known limitations and planned "
        "enhancements."
    ),
    (3, 2): (
        "The bank gets a training tool for new investigators and a reference guide "
        "for customer service. When a customer calls to ask why their purchase was "
        "blocked, the team can look up a similar case and read the explanation "
        "directly."
    ),
}

CH5_MATRIX = {
    (0, 0): (
        "Built Regulatory Compliance tab with: SR 11-7 checklist (8 done + 4 pending, "
        "displayed as disabled checkboxes), fair lending feature risk assessment table "
        "(5 rows with risk levels), model governance framework (identification, risk "
        "classification Tier 2, monitoring schedule \u2014 daily through annual), key "
        "assumptions, right-to-explanation capabilities (3 explanation types + 5-step "
        "dispute workflow), and data lineage (6-step processing pipeline + audit "
        "trail specification)."
    ),
    (0, 1): (
        "Created a one-stop compliance page for regulators and internal audit. It "
        "shows what has been completed (8 items) and what still needs to be done "
        "(4 items). It reviews each feature for potential discrimination risk. It "
        "documents the model's identity, assumptions, and monitoring schedule. It "
        "demonstrates how individual decisions can be explained to customers. And "
        "it traces data from raw source to final prediction."
    ),
    (0, 2): (
        "We built a page for the bank inspectors that shows: a checklist of completed "
        "and pending items, a fairness review of each clue the system uses, the system's "
        "identity card (name, type, version), a schedule for regular check-ups, a "
        "demonstration of how the system explains its decisions, and a map of how "
        "the data flows from start to finish."
    ),
    (1, 0): (
        "Regulatory compliance is a hard requirement for production deployment in "
        "banking. SR 11-7 mandates model documentation, validation, and ongoing "
        "monitoring. ECOA requires fair lending analysis even for fraud models. "
        "GDPR Art. 22 requires right to explanation. Consolidating all compliance "
        "artifacts in a single tab streamlines regulatory review and demonstrates "
        "governance maturity."
    ),
    (1, 1): (
        "No model goes live in a bank without regulatory approval. This tab "
        "pre-packages everything regulators need: documentation, performance evidence, "
        "fairness analysis, and governance framework. Showing 8 completed items and "
        "4 pending items is honest and builds trust \u2014 regulators respect "
        "transparency about what still needs to be done."
    ),
    (1, 2): (
        "Bank inspectors need to see that everything is properly documented and "
        "fair. This page puts all the paperwork in one place. Being honest about "
        "what's done (8 items) and what's still needed (4 items) shows the bank "
        "is serious about doing things right."
    ),
    (2, 0): (
        "st.checkbox(value=True/False, disabled=True) for governance checklist "
        "(non-interactive visual indicator). st.dataframe() for fair lending table "
        "(5 features, risk level, assessment). st.text() for model identification "
        "fields. Monitoring schedule as st.dataframe() with Frequency/Activity columns. "
        "Key assumptions as numbered markdown list. Right-to-explanation as structured "
        "markdown (3 capabilities + 5-step workflow). Data lineage as 6-step pipeline "
        "with audit trail fields."
    ),
    (2, 1): (
        "The checklist uses visual checkboxes (checked for done, unchecked for pending). "
        "Fair lending is presented as a table with risk levels for each feature. Model "
        "governance is displayed as identification fields and a monitoring schedule table. "
        "Right-to-explanation and data lineage are documented as structured text with "
        "numbered steps."
    ),
    (2, 2): (
        "The compliance page uses checkboxes to show what's done and what's not. "
        "A table rates each clue for fairness. The system's identity card shows "
        "its name, type, and version. A schedule shows when different types of "
        "check-ups happen. And a step-by-step description shows how data flows "
        "through the system."
    ),
    (3, 0): (
        "Complete compliance dashboard ready for regulatory review. The disabled "
        "checkboxes provide visual clarity without false interactivity. The fair "
        "lending table documents the proxy discrimination analysis required by ECOA. "
        "The monitoring schedule establishes the operational cadence for ongoing "
        "governance. The data lineage creates an auditable processing chain."
    ),
    (3, 1): (
        "The bank gets a regulatory-ready compliance dashboard that can be shown "
        "directly to auditors. The honest split between completed and pending items "
        "demonstrates governance maturity. The monitoring schedule shows the bank "
        "has a plan for ongoing oversight. The data lineage proves every step is "
        "traceable and reproducible."
    ),
    (3, 2): (
        "The bank can show inspectors a single page with everything they need. "
        "It's organized, honest about what's done and what's not, and shows "
        "the bank has a plan for keeping the system well-maintained."
    ),
}

CH6_MATRIX = {
    (0, 0): (
        "Generated deployment artifacts: dashboard_app.py (~470 lines, self-contained "
        "Streamlit application) and requirements.txt (9 packages with minimum versions). "
        "Local deployment: 'cd notebooks/dashboard && streamlit run dashboard_app.py'. "
        "Cloud deployment: Streamlit Community Cloud via GitHub integration (repo: "
        "JuanCRuizA/Agent-Fraud-Sentinel, branch: main, path: notebooks/dashboard/"
        "dashboard_app.py). Performance: @st.cache_resource and @st.cache_data ensure "
        "model/data loaded once per session."
    ),
    (0, 1): (
        "Prepared the dashboard for both local and cloud deployment. The application "
        "is a single Python file with all dependencies listed in requirements.txt. "
        "Locally, one command launches it. On the cloud, Streamlit Community Cloud "
        "connects directly to the GitHub repository and deploys automatically. "
        "Caching ensures fast performance even with 118,000 transactions."
    ),
    (0, 2): (
        "We prepared everything so the dashboard can run on a laptop or on the "
        "internet. On a laptop, one command starts it. On the internet, it connects "
        "to GitHub and runs automatically. The system is smart about loading data \u2014 "
        "it only loads the model and data once, so everything stays fast."
    ),
    (1, 0): (
        "Production deployment requires: (1) reproducible environment (requirements.txt "
        "with pinned minimum versions), (2) artifact management (model files and data "
        "in the repository), (3) performance optimization (caching decorators prevent "
        "redundant computation), (4) platform selection (Streamlit Cloud provides "
        "free hosting with GitHub integration, 1 GB memory \u2014 sufficient for "
        "118K rows \u00d7 7 features)."
    ),
    (1, 1): (
        "Making the dashboard available to stakeholders requires deployment. "
        "Streamlit Cloud provides free hosting with minimal configuration. The "
        "requirements.txt ensures the same software versions are used everywhere. "
        "Caching means the model loads once and serves all users quickly. "
        "The GitHub integration enables automatic updates when code changes."
    ),
    (1, 2): (
        "A dashboard that only works on one laptop isn't very useful. We set it up "
        "so it can be published on the internet for anyone at the bank to use. "
        "The instructions list tells the computer which software to install. "
        "And smart caching keeps it fast even with lots of data."
    ),
    (2, 0): (
        "%%writefile magic writes dashboard_app.py and requirements.txt from notebook "
        "cells. Local: 'streamlit run dashboard_app.py' (default port 8501). Cloud: "
        "share.streamlit.io with repo/branch/path configuration. @st.cache_resource: "
        "singleton pattern for model (loaded once, shared across sessions/users). "
        "@st.cache_data: hash-based caching for DataFrames and predictions. Paths "
        "resolved relative to __file__ with fallback to Path.cwd()."
    ),
    (2, 1): (
        "The notebook writes the application code to a separate file. Locally, "
        "one terminal command starts the dashboard. On the cloud, Streamlit's "
        "hosting platform reads the code from GitHub and deploys it automatically. "
        "The model is loaded once per server restart, and data predictions are "
        "cached for performance."
    ),
    (2, 2): (
        "The notebook creates the application file and an ingredients list. On a "
        "laptop, one command starts it. On the internet, a cloud service reads "
        "the code from GitHub and runs it automatically. The system is smart "
        "enough to not reload the model every time someone visits."
    ),
    (3, 0): (
        "Deployment-ready application with documented local and cloud paths. "
        "Portfolio-grade artifact: demonstrates full ML lifecycle (data \u2192 "
        "features \u2192 model \u2192 explainability \u2192 interactive dashboard). "
        "Streamlit Cloud deployment enables sharing with recruiters, hiring managers, "
        "and stakeholders via URL. Caching strategy supports multi-user access."
    ),
    (3, 1): (
        "The bank gets a deployable application, not just a prototype. The cloud "
        "deployment option means stakeholders can access it from any browser without "
        "installing software. As a portfolio piece, it demonstrates end-to-end "
        "capability: from raw data through model training, explainability, and "
        "interactive deployment."
    ),
    (3, 2): (
        "The dashboard is ready to go live. Anyone at the bank can access it "
        "through a web link \u2014 no installation needed. As a showcase of the "
        "entire project, it demonstrates everything from understanding the data "
        "to building a model to making it explainable and usable."
    ),
}

ALL_CHAPTERS = [
    {
        "number": 0,
        "title": "Executive Overview",
        "subtitle": "Full-Project Summary",
        "narrative": (
            "This notebook completes the BAFS project by creating an interactive Streamlit "
            "dashboard that consolidates all previous phases (EDA, Feature Engineering, "
            "Model Training, SHAP Explainability) into a single, professional web application. "
            "The dashboard has 4 tabs targeting 4 audiences: Executive Summary (KPIs and cost "
            "analysis), Model Performance (confusion matrix, ROC/PR curves, cost-benefit table), "
            "Case Study Explorer (6 SHAP-explained case studies), and Regulatory Compliance "
            "(SR 11-7, fair lending, audit trail). A sidebar with a threshold slider and sample "
            "size selector enables interactive 'what-if' analysis. The application is deployable "
            "to Streamlit Cloud via GitHub integration."
        ),
        "matrix": CH0_MATRIX,
        "figures": [],
        "callouts": [
            ("insight",
             "The dashboard transforms 5 phases of analysis into an interactive decision-support "
             "tool. The threshold slider is the key innovation: it lets any stakeholder explore "
             "the cost-sensitivity trade-off in real-time, answering 'What happens if we're "
             "more/less strict?'"),
        ],
    },
    {
        "number": 1,
        "title": "Setup & Architecture",
        "subtitle": "Notebook Sections 1-2: Environment Verification, Dashboard Design, File Dependencies",
        "narrative": (
            "The notebook begins by verifying all dependencies (9 Python packages) and confirming "
            "that model artifacts, test data, and SHAP figures exist. The 4-tab architecture is "
            "designed around the four stakeholder groups: executives (KPIs), data scientists "
            "(model validation), fraud analysts (case studies), and compliance officers "
            "(governance). Design principles include professional banking aesthetic (blue/gray "
            "palette), interactive filters, cost-sensitive framing, and regulatory awareness. "
            "The %%writefile magic writes the complete application to dashboard_app.py."
        ),
        "matrix": CH1_MATRIX,
        "figures": [],
        "callouts": [
            ("business",
             "The 4-tab design ensures each stakeholder finds what they need without "
             "information overload. Executives see KPIs, data scientists see curves, "
             "analysts see cases, and compliance sees checklists. One application "
             "serves all four audiences."),
        ],
    },
    {
        "number": 2,
        "title": "Executive Summary Tab",
        "subtitle": "Dashboard Tab 1: KPI Cards, Risk Distribution, Cost Analysis",
        "narrative": (
            "The Executive Summary tab provides at-a-glance KPIs for fraud operations "
            "leadership. Four metric cards show Fraud Detected (with recall), False Positive "
            "Rate, Fraud Prevented (in dollars), and Total Operational Cost. A two-column "
            "layout pairs a performance table (recall, precision, F1, FPR) with a risk score "
            "histogram showing the overlap between fraud and legitimate distributions. "
            "A cost analysis row breaks down missed fraud cost, false alarm cost, and "
            "savings compared to having no model. All metrics update dynamically."
        ),
        "matrix": CH2_MATRIX,
        "figures": [],
        "callouts": [
            ("business",
             "The 'Savings vs No Model' metric quantifies the system's ROI: how much the "
             "fraud detection model saves compared to having no fraud prevention. This is "
             "the single most important metric for justifying the project's budget."),
        ],
    },
    {
        "number": 3,
        "title": "Model Performance Tab",
        "subtitle": "Dashboard Tab 2: Confusion Matrix, ROC/PR Curves, Feature Importance, Cost-Benefit Table",
        "narrative": (
            "The Model Performance tab provides deep technical evaluation for data science "
            "teams and model validators. The confusion matrix heatmap includes dollar-cost "
            "annotations. ROC and Precision-Recall curves display the model's full trade-off "
            "landscape with the current operating point highlighted. SHAP feature importance "
            "from Phase 4 is loaded directly. A cost-benefit table evaluates 8 threshold "
            "values (0.20 to 0.90) side by side, showing recall, precision, true positives, "
            "false positives, missed frauds, and total cost for each."
        ),
        "matrix": CH3_MATRIX,
        "figures": [],
        "callouts": [
            ("insight",
             "The cost-benefit table transforms threshold selection from a technical decision "
             "into a business decision. For each threshold, stakeholders see the exact number "
             "of frauds caught, false alarms generated, and total cost in dollars. This bridges "
             "the gap between data science and business strategy."),
        ],
    },
    {
        "number": 4,
        "title": "Case Study Explorer Tab",
        "subtitle": "Dashboard Tab 3: 6 Case Studies with SHAP Explanations and Plain-English Summaries",
        "narrative": (
            "The Case Study Explorer tab demonstrates individual transaction explainability. "
            "A dropdown selector lets users choose from 6 representative cases from Phase 4: "
            "true positive (clear), true positive (velocity-driven), false negative (missed "
            "fraud), false positive (legitimate flagged), auto-block candidate, and borderline "
            "case. Each case shows metric cards (score, decision, actual), feature values in "
            "readable format, the SHAP waterfall plot, a plain-English explanation, key risk "
            "drivers, and improvement recommendations where applicable."
        ),
        "matrix": CH4_MATRIX,
        "figures": [],
        "callouts": [
            ("business",
             "The Case Study Explorer serves three operational functions: (1) analyst training "
             "tool \u2014 new team members learn how the model 'thinks,' (2) dispute resolution "
             "resource \u2014 ready-made explanations for customer complaints, (3) compliance "
             "evidence \u2014 demonstrates right-to-explanation capability per GDPR Art. 22."),
        ],
    },
    {
        "number": 5,
        "title": "Regulatory Compliance Tab",
        "subtitle": "Dashboard Tab 4: SR 11-7 Checklist, Fair Lending, Model Governance, Audit Trail",
        "narrative": (
            "The Regulatory Compliance tab consolidates all governance artifacts into a single "
            "page for regulatory review. An SR 11-7 checklist displays 8 completed and 4 pending "
            "items as visual checkboxes. A fair lending table assesses 5 feature groups for "
            "proxy discrimination risk (2 flagged as MEDIUM: first-time transactions and "
            "temporal features). Model governance documents identification, risk classification "
            "(Tier 2), and a 5-tier monitoring schedule (daily through annual). "
            "Right-to-explanation and data lineage sections complete the compliance package."
        ),
        "matrix": CH5_MATRIX,
        "figures": [],
        "callouts": [
            ("caution",
             "The 4 pending governance items (disparate impact testing, champion/challenger, "
             "monitoring dashboard, revalidation schedule) require production data or "
             "operational infrastructure. These should be addressed before full production "
             "deployment."),
            ("business",
             "The compliance tab can be shown directly to regulators during SR 11-7 review. "
             "The honest disclosure of 4 pending items demonstrates governance maturity \u2014 "
             "regulators trust organizations that acknowledge what still needs to be done."),
        ],
    },
    {
        "number": 6,
        "title": "Deployment & Production Readiness",
        "subtitle": "Notebook Section 5: Local Deployment, Cloud Deployment, Performance Optimization",
        "narrative": (
            "The deployment section prepares the dashboard for real-world use. Two artifacts "
            "are generated: dashboard_app.py (the complete Streamlit application, ~470 lines) "
            "and requirements.txt (9 Python packages with minimum versions). Local deployment "
            "requires one command: 'streamlit run dashboard_app.py'. Cloud deployment uses "
            "Streamlit Community Cloud with GitHub integration (free hosting, automatic updates). "
            "Performance is optimized via @st.cache_resource (model loaded once) and "
            "@st.cache_data (predictions cached per input hash)."
        ),
        "matrix": CH6_MATRIX,
        "figures": [],
        "callouts": [
            ("insight",
             "The dashboard represents the complete ML lifecycle: data understanding (Phase 1) "
             "\u2192 feature engineering (Phase 2) \u2192 model training (Phase 3) \u2192 "
             "explainability (Phase 4) \u2192 interactive deployment (Phase 5). This end-to-end "
             "pipeline is what differentiates a data scientist from a model builder."),
        ],
    },
]

SUMMARY_TABLE_DATA = [
    ("Dashboard Framework", "Streamlit (Python web framework)"),
    ("Application File", "dashboard_app.py (~470 lines)"),
    ("Number of Tabs", "4 (Executive, Performance, Cases, Compliance)"),
    ("Target Audiences", "Executives, Data Scientists, Analysts, Compliance"),
    ("Interactive Controls", "Threshold slider (0.0-1.0) + Sample size selector"),
    ("KPI Cards", "4 (Fraud Detected, FPR, Fraud Prevented $, Total Cost)"),
    ("Charts", "Confusion matrix, ROC, PR curve, Risk distribution"),
    ("Case Studies", "6 (from Phase 4 SHAP analysis)"),
    ("SHAP Figures Loaded", "2 (feature importance bar, waterfall cases)"),
    ("Cost-Benefit Table", "8 threshold comparisons (0.20 to 0.90)"),
    ("SR 11-7 Checklist", "8 completed + 4 pending"),
    ("Fair Lending Assessment", "5 feature groups (2 MEDIUM risk)"),
    ("Monitoring Schedule", "5 tiers (daily through annual)"),
    ("Deployment Options", "Local (one command) + Streamlit Cloud"),
    ("Caching Strategy", "@st.cache_resource + @st.cache_data"),
    ("Dependencies", "9 Python packages (requirements.txt)"),
    ("Data Loaded", "Test set (118,108 transactions)"),
    ("Model Loaded", "XGBoost (xgboost_final.pkl, scaler.pkl, threshold_config.pkl)"),
]

GLOSSARY = [
    ("@st.cache_data", "Streamlit decorator that caches function results based on input hash, avoiding redundant computation."),
    ("@st.cache_resource", "Streamlit decorator that loads a resource once per server restart, sharing it across all users and sessions."),
    ("Confusion Matrix", "A 2x2 table showing True Negatives, False Positives, False Negatives, and True Positives."),
    ("Cost-Benefit Table", "A comparison table showing recall, precision, and total cost at multiple threshold values."),
    ("Dashboard", "An interactive web application that displays data, charts, and metrics for decision-making."),
    ("Deployment", "Making an application available to users, either locally (on a laptop) or in the cloud (on the internet)."),
    ("KPI", "Key Performance Indicator \u2014 a measurable metric that tracks important business objectives."),
    ("Operating Point", "The specific position on the ROC or PR curve corresponding to the current threshold setting."),
    ("PR Curve", "Precision-Recall curve \u2014 shows the trade-off between precision and recall at all threshold values."),
    ("ROC Curve", "Receiver Operating Characteristic curve \u2014 shows the trade-off between True Positive Rate and False Positive Rate."),
    ("Sidebar", "A collapsible panel on the left side of the dashboard containing navigation and controls."),
    ("Streamlit", "An open-source Python framework for building interactive data dashboards and web applications."),
    ("Streamlit Cloud", "A free hosting platform for Streamlit applications, integrated with GitHub for automatic deployment."),
    ("Threshold Slider", "An interactive control that lets users adjust the fraud detection threshold and see metrics update in real-time."),
]


# ══════════════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_paragraph_spacing(paragraph, before=0, after=0, line=240):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_formatted_text(cell, text, font_name="Calibri", font_size=10,
                       bold=False, color_hex=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=20, after=20, line=240)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def add_callout_box(doc, text, box_type="insight"):
    colors = {
        "insight": (C_INSIGHT_BG, C_INSIGHT_BORDER, "Key Insight"),
        "business": (C_BUSINESS_BG, C_BUSINESS_BORDER, "Business Impact"),
        "caution": (C_CAUTION_BG, C_CAUTION_BORDER, "Caution"),
    }
    bg_color, border_color, label = colors.get(box_type, colors["insight"])

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=120, line=264)

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), bg_color)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)

    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), border_color)
    left.set(qn("w:space"), "4")
    pBdr.append(left)
    pPr.append(pBdr)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:right"), "360")
    pPr.append(ind)

    label_run = p.add_run(f"{label}: ")
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(10)
    label_run.font.bold = True
    label_run.font.color.rgb = RGBColor.from_string(border_color)

    content_run = p.add_run(text)
    content_run.font.name = "Calibri"
    content_run.font.size = Pt(10)
    content_run.font.italic = True
    content_run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)


def add_matrix_table(doc, matrix_data):
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9360")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    headers = [
        "Layer", "Technical\n(BDS Colleague)",
        "Business\n(Manager / Regulator)", "Simple\n(Grandmother)"
    ]
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, layer in enumerate(LAYERS):
        row_idx = i + 1
        layer_cell = table.cell(row_idx, 0)
        set_cell_shading(layer_cell, C_MED_BLUE)
        set_cell_margins(layer_cell, top=60, bottom=60, left=100, right=100)
        add_formatted_text(layer_cell, layer, font_size=9, bold=True, color_hex=C_WHITE)
        layer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for j in range(3):
            cell = table.cell(row_idx, j + 1)
            bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            text = matrix_data.get((i, j), "")
            add_formatted_text(cell, text, font_size=9)

    doc.add_paragraph("")
    return table


def add_section_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)
    return heading


def add_page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
#  UNIQUE ELEMENTS
# ══════════════════════════════════════════════════════════════════════

def add_dashboard_tabs_table(doc):
    """Add the dashboard tab overview table."""
    add_section_heading(doc, "Dashboard Tabs Overview", level=2)

    data = [
        ("Executive Summary", "KPIs, risk distribution, cost analysis", "Executives & Management"),
        ("Model Performance", "Confusion matrix, ROC/PR curves, SHAP, cost-benefit", "Data Science Teams"),
        ("Case Study Explorer", "6 SHAP case studies with plain-English explanations", "Fraud Analysts"),
        ("Regulatory Compliance", "SR 11-7, fair lending, governance, audit trail", "Compliance Officers"),
    ]

    table = doc.add_table(rows=len(data) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Tab", "Content", "Audience"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (tab, content, audience) in enumerate(data):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        for j, text in enumerate([tab, content, audience]):
            cell = table.cell(i + 1, j)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=80, right=80)
            add_formatted_text(cell, text, font_size=10, bold=(j == 0))

    doc.add_paragraph("")


def add_interactive_controls_table(doc):
    """Add the interactive controls table."""
    add_section_heading(doc, "Interactive Controls", level=2)

    data = [
        ("Risk Threshold", "Sidebar slider", "0.00 \u2013 1.00 (default: 0.41)", "Updates all metrics, charts, and tables"),
        ("Sample Size", "Sidebar dropdown", "1K / 5K / 10K / 50K / Full", "Subsamples data for faster exploration"),
        ("Case Study", "Tab 3 dropdown", "6 pre-defined cases", "Shows detailed SHAP explanation per case"),
    ]

    table = doc.add_table(rows=len(data) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Control", "Location", "Range / Options", "Effect"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (control, location, range_opts, effect) in enumerate(data):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        for j, text in enumerate([control, location, range_opts, effect]):
            cell = table.cell(i + 1, j)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=80, right=80)
            add_formatted_text(cell, text, font_size=10, bold=(j == 0))

    doc.add_paragraph("")


def add_ml_lifecycle_table(doc):
    """Add the end-to-end ML lifecycle table."""
    add_section_heading(doc, "End-to-End ML Lifecycle", level=2)

    data = [
        ("Phase 1", "EDA", "01_eda_fraud_patterns.ipynb", "Data understanding, fraud patterns, cost assumptions"),
        ("Phase 2", "Feature Engineering", "02_feature_engineering.ipynb", "7 behavioral features, client identity, temporal split"),
        ("Phase 3", "Model Training", "03_model_training.ipynb", "XGBoost, threshold optimization, production strategy"),
        ("Phase 4", "Explainability", "04_shap_explainability.ipynb", "SHAP analysis, case studies, regulatory compliance"),
        ("Phase 5", "Dashboard", "05_streamlit_dashboard.ipynb", "Interactive deployment, stakeholder communication"),
    ]

    table = doc.add_table(rows=len(data) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Phase", "Topic", "Notebook", "Key Output"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (phase, topic, notebook, output) in enumerate(data):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        for j, text in enumerate([phase, topic, notebook, output]):
            cell = table.cell(i + 1, j)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=80, right=80)
            add_formatted_text(cell, text, font_size=10,
                               bold=(j == 0),
                               font_name="Consolas" if j == 2 else "Calibri")

    doc.add_paragraph("")


# ══════════════════════════════════════════════════════════════════════
#  DOCUMENT CONSTRUCTION
# ══════════════════════════════════════════════════════════════════════

def add_cover_page(doc):
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Agent Fraud Sentinel")
    run.font.name = "Calibri"
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Streamlit Dashboard Deep Dive: 4 Layers \u00d7 3 Perspectives")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("\u2500" * 50)
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)
    run.font.size = Pt(12)

    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ref.add_run("Notebook: 05_streamlit_dashboard.ipynb")
    run.font.name = "Consolas"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        "IEEE-CIS Fraud Detection Dataset\n"
        "Interactive 4-Tab Dashboard | Executive KPIs | Model Validation\n"
        "Case Study Explorer | Regulatory Compliance"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)

    for _ in range(3):
        doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("BAFS Project \u2014 February 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("888888")

    add_page_break(doc)


def add_toc_placeholder(doc):
    add_section_heading(doc, "Table of Contents", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)
    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar_separate)

    placeholder_run = p.add_run(
        "(Right-click here and select 'Update Field' to generate Table of Contents)"
    )
    placeholder_run.font.name = "Calibri"
    placeholder_run.font.size = Pt(10)
    placeholder_run.font.italic = True
    placeholder_run.font.color.rgb = RGBColor.from_string("999999")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    placeholder_run._r.append(fldChar_end)

    add_page_break(doc)


def add_framework_explanation(doc):
    add_section_heading(doc, "The 4-Layer \u00d7 3-Perspective Framework", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "This document examines each component of the Streamlit Dashboard notebook "
        "through two dimensions: four analytical layers and three audience perspectives. "
        "This framework ensures complete understanding \u2014 from raw technical detail "
        "to business impact \u2014 accessible to any reader regardless of their background."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    add_section_heading(doc, "The Four Layers", level=3)
    for title_text, desc in [
        ("Layer 1 \u2014 WHAT did I do?",
         "Describes the concrete actions: tabs built, components designed, deployment prepared."),
        ("Layer 2 \u2014 WHY did I do it?",
         "Explains the motivation: why a dashboard, why these tabs, why interactive controls."),
        ("Layer 3 \u2014 HOW does it work?",
         "Details the mechanics: Streamlit components, caching, CSS, chart generation."),
        ("Layer 4 \u2014 WHAT does the bank gain?",
         "Translates results into value: stakeholder communication, decision support, portfolio piece."),
    ]:
        p = doc.add_paragraph()
        bold_run = p.add_run(title_text + "  ")
        bold_run.font.name = "Calibri"
        bold_run.font.size = Pt(11)
        bold_run.font.bold = True
        bold_run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)
        desc_run = p.add_run(desc)
        desc_run.font.name = "Calibri"
        desc_run.font.size = Pt(11)

    add_section_heading(doc, "The Three Perspectives", level=3)
    for title_text, desc in [
        ("Technical (BDS Colleague)",
         "Uses Streamlit API, caching decorators, layout components, deployment patterns."),
        ("Business (Manager / Regulator)",
         "Focuses on stakeholder communication, decision support, and operational value."),
        ("Simple (Grandmother)",
         "Uses everyday analogies and plain language. No technical background assumed."),
    ]:
        p = doc.add_paragraph()
        bold_run = p.add_run(title_text + "  ")
        bold_run.font.name = "Calibri"
        bold_run.font.size = Pt(11)
        bold_run.font.bold = True
        bold_run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)
        desc_run = p.add_run(desc)
        desc_run.font.name = "Calibri"
        desc_run.font.size = Pt(11)


def add_chapter(doc, chapter_data):
    num = chapter_data["number"]
    title = chapter_data["title"]

    add_page_break(doc)
    add_section_heading(doc, f"Chapter {num}: {title}", level=1)

    if chapter_data.get("subtitle"):
        p = doc.add_paragraph()
        run = p.add_run(chapter_data["subtitle"])
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=200, line=276)
    run = p.add_run(chapter_data["narrative"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    add_section_heading(doc, "Analysis Matrix", level=2)
    add_matrix_table(doc, chapter_data["matrix"])

    for box_type, text in chapter_data.get("callouts", []):
        add_callout_box(doc, text, box_type)


def add_summary_statistics_table(doc):
    add_section_heading(doc, "Dashboard Summary", level=2)

    table = doc.add_table(rows=len(SUMMARY_TABLE_DATA) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Component", "Details"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (metric, value) in enumerate(SUMMARY_TABLE_DATA):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        for j, text in enumerate([metric, value]):
            cell = table.cell(i + 1, j)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=100, right=100)
            add_formatted_text(cell, text, font_size=10, bold=(j == 0))
            if j == 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_appendix_glossary(doc):
    add_page_break(doc)
    add_section_heading(doc, "Appendix: Glossary", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "Key terms used throughout this document, defined for non-technical readers."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    set_paragraph_spacing(p, after=200)

    table = doc.add_table(rows=len(GLOSSARY) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Term", "Definition"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)

    for i, (term, definition) in enumerate(GLOSSARY):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        cell_term = table.cell(i + 1, 0)
        set_cell_shading(cell_term, bg)
        set_cell_margins(cell_term, top=30, bottom=30, left=100, right=100)
        add_formatted_text(cell_term, term, font_size=10, bold=True)

        cell_def = table.cell(i + 1, 1)
        set_cell_shading(cell_def, bg)
        set_cell_margins(cell_def, top=30, bottom=30, left=100, right=100)
        add_formatted_text(cell_def, definition, font_size=10)


def setup_header_footer(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        h_para.text = ""
        run_left = h_para.add_run("Agent Fraud Sentinel \u2014 Streamlit Dashboard")
        run_left.font.name = "Calibri"
        run_left.font.size = Pt(8)
        run_left.font.color.rgb = RGBColor.from_string("999999")
        h_para.add_run("\t\t")
        run_right = h_para.add_run("BAFS")
        run_right.font.name = "Calibri"
        run_right.font.size = Pt(8)
        run_right.font.bold = True
        run_right.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

        footer = section.footer
        footer.is_linked_to_previous = False
        f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        f_para.text = ""
        f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_page = f_para.add_run("Page ")
        run_page.font.name = "Calibri"
        run_page.font.size = Pt(8)
        run_page.font.color.rgb = RGBColor.from_string("999999")
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run_page._r.append(fldChar1)
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run_page._r.append(instrText)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        run_page._r.append(fldChar2)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run_page._r.append(fldChar3)


# ══════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════

def main():
    print("Generating Streamlit Dashboard Analysis Matrix document...")

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i in range(1, 4):
        h_style = doc.styles[f"Heading {i}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

    # ── Build Document ─────────────────────────────────────────────
    add_cover_page(doc)
    add_toc_placeholder(doc)

    # Chapter 0
    ch0 = ALL_CHAPTERS[0]
    add_section_heading(doc, f"Chapter 0: {ch0['title']}", level=1)
    p = doc.add_paragraph()
    run = p.add_run(ch0["subtitle"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=200, line=276)
    run = p.add_run(ch0["narrative"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    add_framework_explanation(doc)
    add_dashboard_tabs_table(doc)
    add_interactive_controls_table(doc)
    add_ml_lifecycle_table(doc)

    add_section_heading(doc, "Full-Project Analysis Matrix", level=2)
    add_matrix_table(doc, ch0["matrix"])

    for box_type, text in ch0.get("callouts", []):
        add_callout_box(doc, text, box_type)

    # Chapters 1-6
    for chapter_data in ALL_CHAPTERS[1:]:
        add_chapter(doc, chapter_data)

    # Summary
    add_summary_statistics_table(doc)

    # Glossary
    add_appendix_glossary(doc)

    # Header/Footer
    setup_header_footer(doc)

    # Save
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"File size: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
