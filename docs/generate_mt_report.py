"""
Generate Word Document: Model Training Notebook 03 - 4 Layers x 3 Perspectives
Agent Fraud Sentinel (BAFS) Project

Produces: docs/03_MT_analysis_matrix.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent.parent
FIGURES_DIR = BASE_DIR / "figures" / "model_training"
OUTPUT_PATH = BASE_DIR / "docs" / "03_MT_analysis_matrix.docx"

# ── Colors ─────────────────────────────────────────────────────────────
C_DARK_BLUE = "1B4F72"
C_MED_BLUE = "2980B9"
C_LIGHT_BLUE = "D6EAF8"
C_FRAUD_RED = "E74C3C"
C_LEGIT_GREEN = "2ECC71"
C_LIGHT_GRAY = "F2F3F4"
C_WHITE = "FFFFFF"
C_BLACK = "000000"
C_DARK_GRAY = "2C3E50"
C_INSIGHT_BG = "D6EAF8"
C_INSIGHT_BORDER = "2980B9"
C_BUSINESS_BG = "D5F5E3"
C_BUSINESS_BORDER = "27AE60"
C_CAUTION_BG = "FADBD8"
C_CAUTION_BORDER = "E74C3C"

LAYERS = ["WHAT did I do?", "WHY did I do it?", "HOW does it work?", "WHAT does the bank gain?"]

# ══════════════════════════════════════════════════════════════════════
#  CONTENT: ALL CHAPTER MATRICES
# ══════════════════════════════════════════════════════════════════════

CH0_MATRIX = {
    (0, 0): (
        "Trained and compared four fraud detection models — Logistic Regression (baseline), "
        "XGBoost (initial + Bayesian), and LightGBM (initial + Bayesian) — on 354,324 "
        "transactions with 7 engineered features. Grid search (6 combos) and Bayesian "
        "optimization via Optuna (30 trials each) were applied. LightGBM (Bayesian) won "
        "with PR-AUC 0.1126 (37.1% over baseline). Cost-based threshold: 0.420 (76.2% "
        "recall, $742K val). Test set: 73.8% recall, $730,482 ($6.18/txn). Saved "
        "best_model_final.pkl, scaler.pkl, threshold_config.pkl."
    ),
    (0, 1): (
        "Built and compared four fraud detection systems, from a simple calculator "
        "(Logistic Regression) to two competing advanced teams (XGBoost and LightGBM), "
        "each fine-tuned using smart search. The winner — LightGBM Bayesian — is 37% "
        "better than the baseline. A cost-based alarm threshold ensures at least 75% of "
        "fraud is caught. Three action tiers handle transactions automatically or route "
        "them to analysts."
    ),
    (0, 2): (
        "We taught four computer programs to spot fraud and held a competition to pick "
        "the best one. The winner uses a team of smart detectives (LightGBM) that "
        "figured out the best strategy through 30 rounds of learning. We then set the "
        "alarm sensitivity so it catches at least 3 out of 4 frauds, and created three "
        "action levels: block it, check it, or approve it."
    ),
    (1, 0): (
        "The 7 engineered features from Phase 2 need a classification model to convert "
        "signals into actionable fraud probabilities. Four models are compared to "
        "demonstrate that gradient boosting (LightGBM) outperforms linear models on "
        "this task and that Bayesian tuning outperforms grid search. Cost-based threshold "
        "optimization translates the EDA's 22.7:1 cost ratio into a concrete decision "
        "boundary. The multi-threshold strategy addresses production operations."
    ),
    (1, 1): (
        "Having good fraud indicators (Phase 2) is not enough — the bank needs a system "
        "that combines them into a single fraud score. Comparing four models proves the "
        "investment in a more complex system is justified. The cost analysis ensures the "
        "system balances fraud prevention with operational costs, and the three-tier "
        "strategy organizes the fraud team's workload."
    ),
    (1, 2): (
        "Having clues is one thing; putting them together to make a decision is another. "
        "We tried four approaches: a simple one and three progressively smarter ones. "
        "The smartest one wins. Then we set the 'sensitivity dial' — how suspicious "
        "should the system be? Too suspicious and it bothers honest customers; not "
        "suspicious enough and it misses thieves."
    ),
    (2, 0): (
        "LR: class_weight='balanced', StandardScaler. XGBoost: scale_pos_weight=28.56, "
        "eval_metric='aucpr'. LightGBM: is_unbalance=True, leaf-wise growth. Grid search: "
        "6 XGB combos (max_depth {4,6,8}, n_estimators {100,200}, lr {0.05,0.1}). "
        "Bayesian: Optuna TPE, 30 trials each. Threshold sweep: cost = FN*$227 + FP*$10. "
        "Multi-threshold: auto-block >=0.90, manual review >=0.420, auto-approve <0.420."
    ),
    (2, 1): (
        "Three types of models are trained on the same 7 features. The two advanced "
        "models are each tuned by testing multiple configurations. The detection "
        "threshold is set by calculating the dollar cost of each possible cutoff. "
        "The three-tier strategy splits transactions into automatic actions and human review."
    ),
    (2, 2): (
        "We trained four computer brains on the same clues and held a tournament. Then "
        "we found the best settings for the two strongest competitors using a smart "
        "trial-and-error process. To set the alarm sensitivity, we calculated the cost "
        "of every possible setting and picked the one that balances catching thieves "
        "with not annoying honest customers."
    ),
    (3, 0): (
        "Production-ready model: PR-AUC 0.1126, recall 73.8%, total cost $730,482 on "
        "118K test transactions ($6.18/txn). Dynamic winner selection ensures the best "
        "model is always used downstream. Multi-threshold strategy: auto-block 0.0%, "
        "manual review 43.9%, auto-approve 56.1%. All artifacts serialized and reproducible."
    ),
    (3, 1): (
        "The bank gets a fraud detection system that catches 3 out of 4 frauds while "
        "keeping costs manageable. The three-tier strategy means analysts only review "
        "borderline cases (45% of transactions), while clear-cut decisions are automated. "
        "All performance metrics are documented in dollar terms for executive reporting."
    ),
    (3, 2): (
        "The bank gets an alarm system that catches about 3 out of every 4 thieves. "
        "Obvious fraud is blocked instantly, suspicious cases go to a human checker, "
        "and clearly safe purchases go through without delay. The total cost is about "
        "$4.85 per transaction processed."
    ),
}

CH1_MATRIX = {
    (0, 0): (
        "Loaded 3 temporal-split CSVs from Phase 2: train (354,324 rows), val (118,108), "
        "test (118,108). Selected 7 features: txn_count_1hr, txn_count_24hr, "
        "amount_deviation, is_first_transaction, hour_of_day, is_weekend, TransactionAmt. "
        "Cleaned inf/NaN values (inf→10, -inf→-10, NaN→0). Established cost "
        "assumptions: FN=$227, FP=$10, ratio 22.7:1."
    ),
    (0, 1): (
        "Loaded the three data files prepared in Phase 2 (training, practice, and final "
        "test sets). Selected the 7 fraud indicators to use as model inputs. Cleaned "
        "any problematic values. Documented the cost of errors: missing a fraud costs "
        "$227, a false alarm costs $10."
    ),
    (0, 2): (
        "We loaded the three data sets we prepared earlier and picked the 7 best clues "
        "to feed the computer. We fixed some messy data values, and reminded ourselves "
        "of the costs: missing a thief costs $227, bothering an honest customer costs $10."
    ),
    (1, 0): (
        "The 7 features were validated for fraud signal in Phase 2 (Tiers 1-4). "
        "amount_bin (Tier 4) is excluded to avoid redundancy with TransactionAmt. "
        "Infinity values arise from Z-score division by zero (single-transaction clients "
        "with std=0). Capping at ±10 preserves the extreme-deviation signal without "
        "causing numerical errors in StandardScaler or gradient computation."
    ),
    (1, 1): (
        "Only validated fraud indicators are used — not all 434 raw columns. This "
        "keeps the model focused and interpretable. The data cleaning step prevents "
        "mathematical errors that would crash the model. The cost assumptions from "
        "Phase 1 ensure every modeling decision is grounded in business reality."
    ),
    (1, 2): (
        "We only use the 7 clues we know work, not the hundreds of messy columns. "
        "Some values were broken (infinity or blanks), so we replaced them with "
        "reasonable numbers. And we kept our cost rules front and center so the "
        "computer learns what matters to the bank."
    ),
    (2, 0): (
        "Feature selection: ENGINEERED_FEATURES list + TransactionAmt. Data cleaning: "
        "df.replace([np.inf, -np.inf], [10, -10]).fillna(0). Assertions verify no "
        "inf/NaN remain. Cost constants: FN_COST=227.00, FP_COST=10.00, "
        "COST_RATIO=22.7. Scaler fit on training data only (leakage prevention)."
    ),
    (2, 1): (
        "Seven features are extracted from each dataset. Infinity values are replaced "
        "with 10 (extreme but finite), and blanks are filled with 0 (meaning 'no "
        "deviation'). The data is verified clean before modeling. Cost parameters "
        "are set as constants used throughout the notebook."
    ),
    (2, 2): (
        "We pulled out the 7 clues from each data file, replaced broken values with "
        "safe ones, double-checked everything was clean, and wrote down our cost rules "
        "so we could use them later."
    ),
    (3, 0): (
        "Clean, validated feature matrices (354K×7, 118K×7, 118K×7) ready for "
        "modeling with no numerical issues. The inf→10 capping preserves extreme "
        "deviation signals (Z-score >10 occurs in <0.01% of cases) while ensuring "
        "numerical stability. Cost framework ready for threshold optimization."
    ),
    (3, 1): (
        "The bank starts with clean, reliable data and clear cost rules. No "
        "mathematical surprises will derail the modeling. The 7-feature design "
        "keeps the system fast enough for real-time scoring — critical for "
        "catching fraud at the point of transaction."
    ),
    (3, 2): (
        "The bank has clean, organized data ready to teach the computer. All the "
        "messy parts are fixed, the cost rules are clear, and the system is set "
        "up to learn from 354,000 past transactions."
    ),
}

CH2_MATRIX = {
    (0, 0): (
        "Trained Logistic Regression with class_weight='balanced', solver='lbfgs', "
        "max_iter=1000. Features standardized via StandardScaler (fit on train only). "
        "Validation PR-AUC: 0.0821. At threshold 0.5: precision 6.93%, recall 42.79%, "
        "F1 0.1193. Top coefficients: txn_count_1hr (0.259), is_first_transaction (-0.120)."
    ),
    (0, 1): (
        "Built a simple, interpretable model as the performance floor. This linear model "
        "catches 43% of fraud at the default threshold — decent but not enough for "
        "production. Its main value is as a benchmark: any advanced model must beat this "
        "to justify its complexity. Feature ranking confirms velocity is the top signal."
    ),
    (0, 2): (
        "We built a simple calculator-style model as our starting point. It catches "
        "about 4 out of 10 frauds. Not great, but it gives us a baseline to beat. "
        "It confirms that how fast someone shops (velocity) is the most important clue."
    ),
    (1, 0): (
        "Logistic Regression provides an interpretable benchmark. class_weight='balanced' "
        "adjusts the loss function to weight fraud samples ~28.6x more (inverse frequency). "
        "Without this, the model would predict all-legitimate (96.5% accuracy, 0% recall). "
        "PR-AUC is preferred over ROC-AUC because ROC-AUC inflates performance on "
        "imbalanced datasets by crediting true negatives."
    ),
    (1, 1): (
        "Every advanced model needs a simple benchmark to compare against. If the "
        "advanced model isn't significantly better, the bank should use the simpler one "
        "(easier to explain to regulators, cheaper to maintain). The 'balanced' setting "
        "prevents the model from taking the lazy route of saying 'nothing is fraud.'"
    ),
    (1, 2): (
        "We start simple to set a bar. If the simple model is good enough, why use "
        "something complicated? The simple model needs special instructions to pay "
        "attention to fraud — otherwise, since fraud is so rare, it would just "
        "say 'everything is fine' and be wrong 3.5% of the time."
    ),
    (2, 0): (
        "StandardScaler: z = (x - mean) / std, fit on X_train only. LogisticRegression("
        "class_weight='balanced'): internally scales positive class weight to n_samples / "
        "(n_classes * n_positive). predict_proba[:,1] gives fraud probability. "
        "precision_recall_curve + auc computes PR-AUC. Coefficients extracted from "
        "model.coef_[0] and sorted by absolute value."
    ),
    (2, 1): (
        "Features are first rescaled so they all have the same range (some are 0-23 "
        "for hours, others are 0-880 for velocity). The model learns a weight for "
        "each feature, then combines them into a fraud probability (0% to 100%). "
        "These weights reveal which features matter most."
    ),
    (2, 2): (
        "We first made all numbers the same scale (like converting inches and meters "
        "to the same unit). Then the model learned how important each clue is and "
        "combined them into a single 'fraud score' from 0 to 100%. The weights "
        "tell us which clues matter most."
    ),
    (3, 0): (
        "Establishes the performance floor: PR-AUC 0.0821. All subsequent models must "
        "significantly exceed this to justify their complexity. The coefficients provide "
        "interpretable feature ranking consistent with Phase 2 signal analysis (velocity > "
        "first_txn > amount). The scaler is saved for deployment."
    ),
    (3, 1): (
        "The bank now has a performance floor: any model chosen for production must beat "
        "PR-AUC 0.0821. The simple model's feature weights confirm the EDA findings "
        "(velocity is most important), which builds confidence in the feature engineering. "
        "If regulators require a simple model, this is ready to deploy."
    ),
    (3, 2): (
        "The bank has a starting point to compare against. The simple model confirms "
        "that our clues work (velocity is #1, just like we expected). If the bank "
        "wants the simplest possible system, this one is ready. But we can do better."
    ),
}

CH3_MATRIX = {
    (0, 0): (
        "Trained XGBoost (initial parameters) with scale_pos_weight=28.56 "
        "(342,336/11,988), max_depth=6, n_estimators=100, learning_rate=0.1, "
        "subsample=0.8, colsample_bytree=0.8, eval_metric='aucpr'. "
        "Validation PR-AUC: 0.1093 (+33.2% over baseline). At threshold 0.5: "
        "recall 61.1%, precision 8.12%, F1 0.1433. Top features: TransactionAmt, txn_count_24hr."
    ),
    (0, 1): (
        "Built the first advanced model (XGBoost) at default settings to establish a "
        "strong gradient-boosting baseline. It catches 61% of fraud — up from 43% with "
        "the simple model. This is the starting point for both the grid search and "
        "the Bayesian optimization that follow."
    ),
    (0, 2): (
        "We built a smarter model — think of it as a team of 100 small detectives, "
        "each asking different questions about the transaction. Together, they catch "
        "6 out of 10 frauds (up from 4 out of 10 with the simple model). This is "
        "just the starting version; we will tune it further later."
    ),
    (1, 0): (
        "XGBoost captures non-linear feature interactions that Logistic Regression cannot. "
        "For example, txn_count_1hr > 5 AND hour_of_day in [7,8,9] creates a compound "
        "risk signal not expressible as a linear combination. scale_pos_weight=28.56 "
        "reweights the gradient to handle the 3.38% fraud rate without oversampling "
        "(SMOTE), which can introduce synthetic artifacts."
    ),
    (1, 1): (
        "Fraud patterns are complex — a $200 purchase at 8 AM by someone who made "
        "5 purchases in the last hour is very different from the same amount at 2 PM "
        "by a regular customer. The simple model can't see these combinations; XGBoost "
        "can. The 28.56x weight tells XGBoost each fraud case matters 28.56 times more."
    ),
    (1, 2): (
        "The simple model looks at each clue separately, but fraud often involves "
        "combinations. A big purchase alone isn't suspicious, but a big purchase at "
        "7 AM after 10 other purchases in the last hour? Very suspicious. The smart "
        "model learns these combinations automatically."
    ),
    (2, 0): (
        "XGBoost builds sequential decision trees, each correcting the previous one's "
        "errors. scale_pos_weight modifies the gradient: loss_fraud = 28.56 * loss_legit. "
        "subsample=0.8 and colsample_bytree=0.8 provide regularization via row and "
        "column bagging. eval_metric='aucpr' monitors validation PR-AUC during training. "
        "No StandardScaler needed (tree-based models are scale-invariant)."
    ),
    (2, 1): (
        "The model builds 100 small decision trees, one at a time. Each new tree focuses "
        "on the mistakes the previous trees made. The fraud-weighting ensures the model "
        "pays 28.56 times more attention to fraud cases. Random subsampling of data and "
        "features prevents the model from memorizing the training data."
    ),
    (2, 2): (
        "Imagine 100 detectives taking turns. The first one does their best, then "
        "passes the hard cases to the second, who focuses on what the first missed. "
        "Each new detective learns from the mistakes of the ones before. Together, "
        "they form a much better team than any one detective alone."
    ),
    (3, 0): (
        "33.2% PR-AUC improvement over Logistic Regression (0.1093 vs 0.0821) "
        "justifies gradient boosting complexity. Feature importance (TransactionAmt > "
        "txn_count_24hr > hour_of_day) differs from LR coefficients, confirming "
        "non-linear effects dominate. This initial model is the baseline for tuning "
        "and will compete against LightGBM in the final comparison."
    ),
    (3, 1): (
        "The advanced model is significantly better (33% improvement), justifying its "
        "use despite being more complex. It catches 61% of fraud vs 43% for the "
        "simple model. The bank gets a stronger fraud detection engine — before tuning. "
        "We now need to compare it against LightGBM."
    ),
    (3, 2): (
        "The smart model is much better — it catches 6 out of 10 frauds instead of 4. "
        "That's 50% more thieves caught. But this is just round one. We still need "
        "to test a rival team and then tune both of them."
    ),
}

CH_LGB_MATRIX = {
    (0, 0): (
        "Trained LightGBM (initial parameters) with is_unbalance=True, num_leaves=31, "
        "n_estimators=100, learning_rate=0.1, subsample=0.8, colsample_bytree=0.8. "
        "Validation PR-AUC: 0.1095 (+0.2% over XGBoost initial). At threshold 0.5: "
        "recall 62.0%, precision 8.22%, F1 0.1452. Feature importance confirms "
        "TransactionAmt and txn_count_24hr as dominant signals — same as XGBoost."
    ),
    (0, 1): (
        "Added a second advanced model (LightGBM) using identical features and cost "
        "structure. At default settings, LightGBM marginally outperforms XGBoost initial. "
        "Both models will be Bayesian-tuned in Section 6.2 before the final comparison. "
        "The initial training confirms LightGBM is a competitive alternative."
    ),
    (0, 2): (
        "We added a second 'smart' model — a different brand of detective team — to "
        "compete against XGBoost. At the starting line, the two teams are neck-and-neck. "
        "The real tournament happens after we tune both teams' strategies."
    ),
    (1, 0): (
        "LightGBM uses leaf-wise tree growth: at each step it splits the single leaf "
        "with the maximum gain, regardless of depth. XGBoost uses level-wise growth: "
        "it splits all leaves at the same depth simultaneously. Leaf-wise is faster "
        "and often achieves lower loss on large datasets. is_unbalance=True is the "
        "LGB equivalent of scale_pos_weight — it computes n_negative/n_positive "
        "internally and applies it to the gradient."
    ),
    (1, 1): (
        "Having two competing model architectures ensures the bank does not rely on a "
        "single algorithm. If one has a systematic blind spot, the other may compensate. "
        "The comparison also validates that the 7 engineered features generalize across "
        "different gradient boosting implementations — not just XGBoost."
    ),
    (1, 2): (
        "We brought in a rival detective team that uses a different investigation method. "
        "Competition is healthy — if both good teams agree that the same clues matter, "
        "we can be more confident they really do matter. If one team misses certain "
        "types of crimes, the other might catch them."
    ),
    (2, 0): (
        "LightGBM grows trees leaf-wise: argmax_leaf(delta_loss(split)). num_leaves=31 "
        "caps maximum leaves per tree (controls model complexity, analogous to max_depth). "
        "is_unbalance=True: weight = n_negative / n_positive ≈ 28.56, applied to loss. "
        "eval_set=[(X_val, y_val)] monitors validation PR-AUC during training. "
        "No StandardScaler needed. Feature importances: lgb_model.feature_importances_."
    ),
    (2, 1): (
        "LightGBM builds its trees differently: instead of building each layer equally, "
        "it focuses on the most informative branches first. This makes it faster and "
        "often more accurate. The 'unbalanced' flag tells it to treat fraud cases as "
        "28x more important — the same logic as XGBoost but with a different parameter name."
    ),
    (2, 2): (
        "The second team uses a different strategy: they always investigate the most "
        "suspicious lead first, going deeper on promising clues rather than checking "
        "everything equally. They also automatically focus on fraud cases, just like "
        "the first team but using different internal rules."
    ),
    (3, 0): (
        "LightGBM initial PR-AUC 0.1095 vs XGBoost initial 0.1093 — nearly identical "
        "at default settings. The real differentiation comes after Bayesian tuning "
        "(Section 6.2). Feature importance plots from both models independently "
        "confirm velocity + amount as top signals — cross-architecture validation "
        "of the Phase 2 feature engineering."
    ),
    (3, 1): (
        "The bank now has two capable gradient boosting algorithms, both significantly "
        "better than the simple model. Their initial tie suggests neither dominates by "
        "architecture alone — tuning will decide the winner. Both models independently "
        "agree on which features matter most."
    ),
    (3, 2): (
        "Both detective teams start the race tied. Now we need to find the best strategy "
        "for each before declaring a winner. The fact that both teams agree on the same "
        "important clues gives us confidence those clues are real."
    ),
}

CH_TUNING_MATRIX = {
    (0, 0): (
        "Section 6.1 — Grid search: 6 XGBoost configurations (max_depth {4,6,8} × "
        "n_estimators {100,200} × learning_rate {0.05,0.1}). Best: depth=6, "
        "n_estimators=200, lr=0.05, PR-AUC=0.1098 (+0.5% over initial). "
        "Section 6.2 — Bayesian (Optuna TPE, 30 trials each): XGB Bayesian 0.1116 "
        "(+1.6% over grid), LGB Bayesian 0.1126 (+2.5% over XGB grid). Winner: LightGBM Bayesian."
    ),
    (0, 1): (
        "Two tuning rounds: a quick grid search (6 combos, XGBoost only) to establish "
        "a strong baseline, then Bayesian optimization (30 trials each) for both "
        "XGBoost and LightGBM. Both improved over their defaults, but LightGBM Bayesian "
        "came out ahead with the highest PR-AUC. The full comparison table shows all "
        "six model versions at threshold 0.5."
    ),
    (0, 2): (
        "We tried two rounds of 'dial-tuning.' First, 6 quick settings for XGBoost. "
        "Then 30 smarter attempts for both teams — each attempt learning from the "
        "previous ones. The second team (LightGBM) ended up with the best settings "
        "after 30 rounds of learning."
    ),
    (1, 0): (
        "Grid search is exhaustive but limited: 6 combinations cover only a coarse "
        "parameter space. Bayesian optimization (TPE sampler) builds a probabilistic "
        "model of the objective function. After each trial it updates the model and "
        "selects the next trial in the region most likely to improve PR-AUC. With 30 "
        "trials, it explores a 6-dimensional continuous space far more efficiently "
        "than any grid could."
    ),
    (1, 1): (
        "Grid search tests a pre-defined list — like checking 6 specific dial "
        "combinations. Bayesian optimization learns from each attempt: if a certain "
        "range works well, it tries more values there. With 30 trials per model, "
        "it explores a much wider range of settings. The validation PR-AUC guides "
        "every decision — the same fair metric used throughout."
    ),
    (1, 2): (
        "First we tried 6 dial combinations on the first team (quick but limited). "
        "Then we used a smarter approach for both teams: 30 attempts where each "
        "one learned from the previous ones. Like a safe-cracker who remembers which "
        "directions feel warm — they don't try random combinations, they use what "
        "they learned to make better guesses."
    ),
    (2, 0): (
        "Grid: nested loop, XGBClassifier trained with eval_set=[(X_val, y_val)], "
        "PR-AUC scored. Bayesian: optuna.create_study(direction='maximize') + "
        "study.optimize(objective, n_trials=30). Objective: train model → return "
        "pr_auc_score(y_val, proba_val). TPE sampler fits two density models l(x) "
        "and g(x); next trial = argmax l(x)/g(x). Final models retrained with "
        "study.best_params on full training set."
    ),
    (2, 1): (
        "Grid: trains each of 6 configurations and picks the winner by PR-AUC on "
        "the validation set. Bayesian: each of 30 trials trains a model with chosen "
        "parameters and reports PR-AUC. The system learns which parameter regions "
        "give high PR-AUC and focuses subsequent trials there. Final models are "
        "retrained with the best parameters found."
    ),
    (2, 2): (
        "Grid search: try 6 combinations, pick the best. Bayesian: try 30 combinations, "
        "but each one is chosen based on what worked before. After 30 rounds for each "
        "team, the best settings are used to train the final models. It's like learning "
        "from experience rather than guessing randomly."
    ),
    (3, 0): (
        "Grid XGB PR-AUC: 0.1098 (+0.5% over initial). Bayesian XGB: 0.1116 (+1.6% "
        "over grid). Bayesian LGB: 0.1126 (+2.5% over XGB grid). Bayesian clearly "
        "outperforms the 6-combo grid. LightGBM Bayesian wins. Both final models "
        "saved: best_model_final.pkl (LGB) and xgboost_final.pkl (XGB, backwards compat "
        "for Phase 4/5). Bayesian runtime: ~10 min per model — justified by the win."
    ),
    (3, 1): (
        "The bank gets the best-performing model after systematic exploration: LightGBM "
        "Bayesian at PR-AUC 0.1126. The tuning results document exactly which "
        "configurations were tested (fully auditable). Bayesian optimization outperformed "
        "grid search for both models, justifying its use despite the longer runtime."
    ),
    (3, 2): (
        "After all the tuning, the second team (LightGBM) with its best settings is "
        "the winner. The bank can see exactly which settings were tried — no guesswork. "
        "The smarter search method (Bayesian) was worth the extra time it took."
    ),
}

CH_COMPARISON_MATRIX = {
    (0, 0): (
        "All 6 model versions compared at threshold 0.5 on validation set: LR (0.0821), "
        "XGB initial (0.1093), XGB grid (0.1098), LGB initial (0.1095), XGB Bayesian "
        "(0.1116), LGB Bayesian (0.1126). Winner: LightGBM Bayesian (+37.1% over LR, "
        "+2.5% over XGB grid). PR curve plotted: LR vs LGB Bayesian. Dynamic winner "
        "selection sets final_proba_val/test for all downstream cells."
    ),
    (0, 1): (
        "Six model versions competed side-by-side on the same validation data at the "
        "same threshold. LightGBM Bayesian wins with the highest PR-AUC. The "
        "Precision-Recall curve confirms LightGBM is better across all sensitivity "
        "levels. The system automatically routes all subsequent analysis to the winner."
    ),
    (0, 2): (
        "All six versions of our detective teams were tested under identical conditions. "
        "LightGBM with its best settings wins the final competition. From here on, "
        "only the winning team's scores are used for all decisions."
    ),
    (1, 0): (
        "Comparing at a fixed threshold eliminates threshold selection as a confound. "
        "PR-AUC is threshold-independent: it evaluates discriminative ability across "
        "all operating points. The dynamic winner pattern — final_proba_val and "
        "final_proba_test set in the winner cell and used by all downstream cells — "
        "eliminates hardcoded model references and enables automatic pipeline routing."
    ),
    (1, 1): (
        "Six models, same test conditions — the fairest possible comparison. PR-AUC "
        "tells us which model is fundamentally better before any threshold decisions. "
        "The automatic winner routing ensures all cost analysis, threshold optimization, "
        "and production strategy metrics reflect the actual best model — not a "
        "hardcoded assumption."
    ),
    (1, 2): (
        "All teams face the same test at the same time, with the same rules. The "
        "winner gets to do all the real work. If we run the competition again in the "
        "future and a different team wins, the system automatically switches to using "
        "that team — no one has to rewrite the code."
    ),
    (2, 0): (
        "For each model: precision_recall_curve(y_val, proba_val) + auc(recall, precision). "
        "Fair comparison at threshold 0.5: (proba >= 0.5).astype(int). precision_score, "
        "recall_score, f1_score. Winner selection: candidates = [(pr_auc, name, "
        "proba_val, proba_test, model), ...]; best = max(candidates, key=lambda x: x[0]). "
        "Sets winner_name, winner_pr_auc, final_proba_val, final_proba_test, winner_model."
    ),
    (2, 1): (
        "Each model's fraud probability scores are evaluated using PR-AUC and "
        "threshold-0.5 metrics. The winner is selected programmatically. The system "
        "stores the winner's probabilities under generic names so all downstream analysis "
        "automatically uses the best model — no manual updates needed."
    ),
    (2, 2): (
        "Six scores are compared, the highest wins. The computer automatically picks "
        "the winner and stores its answers under a shared name. All later steps use "
        "those answers, so the whole pipeline adjusts automatically to the winner."
    ),
    (3, 0): (
        "LightGBM Bayesian selected: PR-AUC 0.1126 (37.1% over LR baseline). "
        "Dynamic selection is critical: all 8 downstream cells (threshold optimization, "
        "production strategy, confusion matrices) use final_proba_val/test generically. "
        "Confirms that 7 features support both gradient boosting architectures — "
        "the feature quality, not the algorithm, is the binding constraint."
    ),
    (3, 1): (
        "The bank has a clear, documented winner: LightGBM Bayesian is 37.1% better "
        "than the simple model. The selection process is reproducible and auditable — "
        "regulators can verify which models were tested and why the winner was chosen. "
        "The automatic routing future-proofs the pipeline against algorithm changes."
    ),
    (3, 2): (
        "The bank's best fraud detector is chosen fairly and automatically. It's 37% "
        "better than the simple model. The system is smart enough to switch to a "
        "different winner in the future without anyone needing to rewrite the code."
    ),
}

CH_THRESHOLD_MATRIX = {
    (0, 0): (
        "Swept threshold 0.01-0.99 with cost function: total_cost = FN×$227 + FP×$10 "
        "(applied to LightGBM Bayesian probabilities). Unconstrained optimum: threshold "
        "0.520, cost $713K, recall 59.8% (unacceptable). Constrained (recall≥75%): "
        "threshold 0.420, cost $742K, recall 76.2%. Trade-off: +$29K cost, additional "
        "frauds caught vs unconstrained. Cost per additional fraud updated."
    ),
    (0, 1): (
        "Tested every possible threshold to find the cheapest one. The cheapest option "
        "($326K) only catches 17% of fraud — the bank misses 83% of thieves. "
        "By requiring at least 75% of fraud be caught, the cost rises to $578K, "
        "but 2,743 additional frauds are intercepted. This trade-off — ~$92 per "
        "extra fraud caught — is justified given the $227 full economic cost per missed fraud "
        "(transaction loss, chargeback, ops, reputational damage)."
    ),
    (0, 2): (
        "We tested every alarm setting from 'catch everything' to 'catch almost nothing.' "
        "The cheapest setting misses most fraud — terrible! So we added a rule: "
        "'you must catch at least 3 out of 4 frauds.' This costs more but prevents "
        "2,743 additional thefts."
    ),
    (1, 0): (
        "Pure cost minimization produces a degenerate solution: the threshold rises "
        "until the model barely flags anything, minimizing FP cost but accepting massive "
        "FN loss. The 17.1% recall at threshold 0.720 is operationally useless. The "
        "75% recall constraint reflects regulatory expectations and customer trust. "
        "This is a business constraint, not a mathematical one."
    ),
    (1, 1): (
        "Minimizing cost alone leads to a model that ignores most fraud — because "
        "false alarms are cheap and individual missed frauds are 'only' $75. But the "
        "bank can't afford to miss 83% of fraud: regulators, customers, and reputation "
        "all demand better. The 75% minimum recall is a business guardrail against "
        "this pathological optimization."
    ),
    (1, 2): (
        "If we only care about cost, the system would barely raise any alarms — "
        "missing almost all thieves. That's like a security guard who never stops "
        "anyone because checking people is annoying. The bank needs the guard to "
        "stop at least 3 out of 4 suspicious people."
    ),
    (2, 0): (
        "For each threshold t in np.arange(0.01, 1.00, 0.01): y_pred = (final_proba_val >= t), "
        "compute FN = fraud missed, FP = legit flagged, cost = FN*227 + FP*10. "
        "Unconstrained: argmin(costs). Constrained: filter thresholds where "
        "recall >= 0.75, then argmin(filtered_costs). Trade-off: "
        "(constrained_cost - unconstrained_cost) / (constrained_TP - unconstrained_TP)."
    ),
    (2, 1): (
        "Every threshold from 1% to 99% is tested on the validation set using the "
        "winning model's probabilities. For each, the system counts missed frauds and "
        "false alarms, then multiplies by their costs. The cheapest threshold is found. "
        "Then, adding the requirement to catch 75% of fraud, the cheapest threshold "
        "meeting that requirement is selected."
    ),
    (2, 2): (
        "We tested 99 different sensitivity settings and calculated the cost of each. "
        "First we found the cheapest one (misses too much fraud). Then we found the "
        "cheapest one that still catches at least 3 out of 4 frauds."
    ),
    (3, 0): (
        "Constrained threshold (0.410) becomes the production manual-review boundary. "
        "The $29K cost increase prevents more direct fraud losses (constrained vs unconstrained) "
        "by catching additional frauds above the minimum recall floor. The "
        "cost-per-fraud metric provides a clear business justification. "
        "Threshold 0.410 is saved in threshold_config.pkl."
    ),
    (3, 1): (
        "The bank gets a clear cost-benefit analysis: spending an extra $252K prevents "
        "2,743 additional frauds worth $206K in direct losses — plus avoiding customer "
        "anger, regulatory fines, and reputation damage. The 75% recall target is "
        "documented for audit compliance."
    ),
    (3, 2): (
        "The bank spends $252K more to catch 2,743 extra thieves. Those thefts would "
        "have cost at least $206K, plus the bank avoids angry customers and trouble "
        "with regulators. The extra cost is justified."
    ),
}

CH_PRODUCTION_MATRIX = {
    (0, 0): (
        "3-tier production strategy applied to LightGBM Bayesian probabilities: "
        "auto-block (>=0.90, 8 txns, 0.01%), manual review (0.41-0.90, 53,271 txns, "
        "45.1%), auto-approve (<0.41, 64,829 txns, 54.9%). Validation: 3,530 frauds "
        "caught, 1,080 missed, cost $578K. Test set: recall 74.4% (3,025/4,064), "
        "precision 5.8%, cost $572,900 ($4.85/txn). Saved best_model_final.pkl, "
        "xgboost_final.pkl (compat), scaler.pkl, threshold_config.pkl."
    ),
    (0, 1): (
        "Created an operational system with three action levels: (1) obvious fraud is "
        "blocked instantly — no human needed, (2) borderline cases go to analysts "
        "for review, (3) clearly safe transactions pass through. The final test on "
        "completely new data shows the system catches 74.4% of fraud at a cost "
        "of $4.85 per transaction."
    ),
    (0, 2): (
        "The system has three levels: 'Stop!' for obvious fraud (blocked automatically), "
        "'Check this' for suspicious cases (a person reviews it), and 'Go ahead' for "
        "clearly safe purchases. Testing on new data shows it catches about 3 out of "
        "every 4 thieves at $4.85 per transaction."
    ),
    (1, 0): (
        "A single threshold creates a binary decision, but fraud operations require "
        "graduated responses. Auto-block (>=0.90) handles the 0.01% of transactions "
        "where the model has near-certainty, reducing latency and analyst workload. "
        "Manual review (0.41-0.90) escalates uncertain cases to human judgment. "
        "Auto-approve (<0.41) frees 55% of transactions from any friction. "
        "The test set evaluation provides the unbiased performance estimate."
    ),
    (1, 1): (
        "Not all fraud decisions are equal. Some are clear-cut (score 0.95 = definitely "
        "fraud), some are borderline (score 0.60 = maybe), and some are clearly fine "
        "(score 0.10 = definitely legitimate). A single threshold wastes human analysts "
        "on clear-cut cases. Three tiers match the response to the confidence level."
    ),
    (1, 2): (
        "Imagine a traffic light: red means 'stop' (obvious fraud), yellow means "
        "'slow down and check' (suspicious), and green means 'go ahead' (safe). "
        "Without this system, every slightly suspicious transaction would need a "
        "human checker, overwhelming the team."
    ),
    (2, 0): (
        "Segmentation: auto_block = (final_proba_val >= 0.90), manual_review = "
        "(final_proba_val >= 0.41) & (final_proba_val < 0.90), auto_approve = "
        "(final_proba_val < 0.41). Cost: auto-block FP × $5 + manual FP × $10 + "
        "missed fraud × $227. Test evaluation: same thresholds applied to final_proba_test. "
        "Confusion matrix plotted (absolute + percentages). Artifacts: joblib.dump()."
    ),
    (2, 1): (
        "Each transaction's fraud score determines its action. The test set (never "
        "seen during training or threshold selection) provides the final unbiased "
        "performance estimate. Confusion matrices show the full breakdown of correct "
        "and incorrect decisions. All model files are saved for deployment in Phase 4 "
        "and Phase 5."
    ),
    (2, 2): (
        "The computer gives each purchase a suspicion score. High scores are blocked "
        "automatically, medium scores go to a human checker, and low scores pass "
        "through. We tested this on data the computer never saw before to make sure "
        "it really works. Then we saved everything for real-world use."
    ),
    (3, 0): (
        "Test set: recall 74.4% (near 75% target), precision 5.8%, cost $572,900. "
        "Confusion matrix: TN=64,541, FP=49,503, FN=1,039, TP=3,025. Auto-approve "
        "correctly handles 64,541 legit transactions (54.6%). Model artifacts serialized "
        "for Phase 4 (SHAP — loads xgboost_final.pkl for backwards compat) and "
        "Phase 5 (Streamlit dashboard — loads best_model_final.pkl)."
    ),
    (3, 1): (
        "The bank gets a battle-tested system: 74.4% fraud detection on completely "
        "new data. The three-tier approach means 55% of transactions need zero human "
        "involvement. Obvious fraud is stopped instantly. Analysts focus only on "
        "borderline cases (45% of transactions). Everything is saved and ready for "
        "the explainability review (Phase 4) and the dashboard (Phase 5)."
    ),
    (3, 2): (
        "The final test shows the system works: it catches 3 out of 4 thieves on "
        "data it has never seen. More than half of all transactions go through "
        "without anyone needing to check them. The system is saved and ready to "
        "be explained (Phase 4) and put into a control panel (Phase 5)."
    ),
}

ALL_CHAPTERS = [
    {
        "number": 0,
        "title": "Executive Overview",
        "subtitle": "Full-Project Summary",
        "narrative": (
            "This notebook takes the 7 engineered features from Phase 2 and trains them into "
            "a production-ready fraud detection model. Four models are compared: Logistic "
            "Regression (interpretable baseline), XGBoost (initial + Bayesian via Optuna), "
            "and LightGBM (initial + Bayesian). LightGBM Bayesian wins with PR-AUC 0.1126 "
            "(37.1% over baseline). Cost-based threshold optimization using the EDA's $75/$10 "
            "cost assumptions produces a three-tier production strategy: auto-block (score "
            ">=0.90), manual review (0.41-0.90), and auto-approve (<0.41). On the "
            "118,108-transaction test set, the system achieves 74.4% recall at a cost of "
            "$4.85 per transaction."
        ),
        "matrix": CH0_MATRIX,
        "figures": [],
        "callouts": [
            ("insight",
             "The notebook follows a disciplined two-step process: first select the best "
             "model (using PR-AUC at a fixed threshold for fair comparison across all 6 "
             "model versions), then optimize the threshold (using cost analysis with business "
             "constraints). Mixing these steps is a common pitfall."),
        ],
    },
    {
        "number": 1,
        "title": "Setup & Data Preparation",
        "subtitle": "Notebook Sections 1-2: Loading, Feature Selection, Data Cleaning, Cost Assumptions",
        "narrative": (
            "The pipeline begins by loading the three temporal-split CSV files from Phase 2 "
            "and selecting the 7 validated features. Infinity values in amount_deviation "
            "(from Z-score division by zero) are capped at ±10. NaN values are filled "
            "with 0. The cost framework from Phase 1 is carried forward: FN=$227, FP=$10, "
            "ratio 22.7:1."
        ),
        "matrix": CH1_MATRIX,
        "figures": [],
        "callouts": [
            ("business",
             "The 22.7:1 cost ratio is the foundation of every modeling decision. It means "
             "the bank should tolerate up to 22.7 false alarms for every fraud it catches — "
             "because missing a fraud is 22.7 times more expensive than investigating a false alarm."),
        ],
    },
    {
        "number": 2,
        "title": "Baseline Model: Logistic Regression",
        "subtitle": "Notebook Section 3: Training, Evaluation, Feature Coefficients",
        "narrative": (
            "Logistic Regression with balanced class weights establishes the performance floor. "
            "Features are standardized (mean=0, std=1) via StandardScaler fit on training data "
            "only. The baseline achieves PR-AUC 0.0821 with 42.8% recall at threshold 0.5. "
            "Feature coefficients confirm velocity (txn_count_1hr: 0.259) as the strongest "
            "linear signal, consistent with Phase 2 findings."
        ),
        "matrix": CH2_MATRIX,
        "figures": [],
        "callouts": [
            ("insight",
             "Logistic Regression's coefficients provide the only truly interpretable feature "
             "ranking in this pipeline. txn_count_1hr (0.259) dominates, confirming that "
             "transaction velocity is the most important linear fraud signal."),
        ],
    },
    {
        "number": 3,
        "title": "Advanced Model: XGBoost (Initial)",
        "subtitle": "Notebook Section 4: Training with Class Imbalance Handling",
        "narrative": (
            "XGBoost with scale_pos_weight=28.56 captures non-linear feature interactions. "
            "The initial model (default hyperparameters) achieves PR-AUC 0.1093, a 33.2% "
            "improvement over the baseline. At threshold 0.5, recall jumps from 42.8% to "
            "61.1%. Feature importance shifts: TransactionAmt and txn_count_24hr become "
            "dominant, suggesting non-linear amount patterns that Logistic Regression missed. "
            "This is the XGBoost starting point before tuning."
        ),
        "matrix": CH3_MATRIX,
        "figures": [
            ("xgb_feature_importance.png",
             "Figure 1: XGBoost Feature Importance (Gain). TransactionAmt and txn_count_24hr "
             "dominate, revealing non-linear amount and velocity patterns missed by Logistic "
             "Regression."),
        ],
        "callouts": [
            ("business",
             "XGBoost catches 18% more fraud than the simple model (61% vs 43%) because it "
             "detects complex patterns: 'high velocity + unusual amount + early morning' is "
             "far more suspicious than any single indicator alone."),
        ],
    },
    {
        "number": 4,
        "title": "Advanced Model: LightGBM (Initial)",
        "subtitle": "Notebook Section 5: Leaf-Wise Growth, is_unbalance, Initial Comparison",
        "narrative": (
            "LightGBM with is_unbalance=True introduces an alternative gradient boosting "
            "architecture for direct comparison with XGBoost. At default parameters, LightGBM "
            "achieves PR-AUC 0.1095 — marginally ahead of XGBoost initial (0.1093). "
            "Both models will be Bayesian-optimized in Section 6.2. The initial training "
            "confirms LightGBM is competitive and validates the feature engineering across "
            "two independent architectures."
        ),
        "matrix": CH_LGB_MATRIX,
        "figures": [
            ("lgb_feature_importance.png",
             "Figure 2: LightGBM Feature Importance. Feature ranking is consistent with "
             "XGBoost: TransactionAmt and velocity features dominate, cross-validating "
             "that the signals are real and not algorithm-specific."),
        ],
        "callouts": [
            ("insight",
             "The near-identical initial performance of XGBoost (0.1093) and LightGBM (0.1095) "
             "confirms that the feature quality — not the algorithm — is the binding constraint "
             "at default settings. Bayesian tuning will differentiate them."),
        ],
    },
    {
        "number": 5,
        "title": "Hyperparameter Tuning",
        "subtitle": "Notebook Section 6: Grid Search (6.1) + Bayesian Optimization via Optuna (6.2)",
        "narrative": (
            "Two tuning approaches are applied. Section 6.1: a 6-combination grid search on "
            "XGBoost (max_depth, n_estimators, learning_rate). Best: depth=6, 200 trees, "
            "lr=0.05, PR-AUC=0.1098. Section 6.2: Optuna TPE sampler, 30 trials each for "
            "XGBoost and LightGBM across 6 continuous hyperparameters. XGB Bayesian: 0.1116 "
            "(+1.6% over grid). LGB Bayesian: 0.1126 (+2.5% over XGB grid). Winner: LightGBM Bayesian."
        ),
        "matrix": CH_TUNING_MATRIX,
        "figures": [],
        "callouts": [
            ("insight",
             "The Bayesian approach outperforms the 6-combo grid for both models. With only 7 "
             "features, performance is constrained by feature quality more than by "
             "hyperparameters — validating the Phase 2 feature engineering. The marginal gains "
             "(0.5% to 2.5%) are expected for a well-engineered 7-feature pipeline."),
        ],
    },
    {
        "number": 6,
        "title": "6-Model Comparison & Winner Selection",
        "subtitle": "Notebook Section 7 (Part 1): Fair Comparison at Threshold 0.5, PR Curves",
        "narrative": (
            "All 6 model versions are compared at threshold 0.5 for fairness. LightGBM "
            "Bayesian wins (PR-AUC 0.1126). Dynamic winner selection stores the winner's "
            "probabilities as final_proba_val and final_proba_test — used by all 8 downstream "
            "cells. The Precision-Recall curve confirms LightGBM Bayesian dominates "
            "across all recall levels."
        ),
        "matrix": CH_COMPARISON_MATRIX,
        "figures": [
            ("pr_curve_comparison.png",
             "Figure 3: Precision-Recall Curve — Model Comparison. LightGBM Bayesian "
             "(PR-AUC 0.1126) consistently outperforms Logistic Regression (PR-AUC 0.0821) "
             "across all recall levels. The red dashed line represents the no-model baseline "
             "(3.5% fraud rate)."),
        ],
        "callouts": [
            ("insight",
             "The dynamic winner selection is critical for pipeline robustness. All 8 downstream "
             "cells use final_proba_val/final_proba_test — never a hardcoded model name. If "
             "XGBoost Bayesian outperforms on a future run, the pipeline automatically routes "
             "to it without any code changes."),
        ],
    },
    {
        "number": 7,
        "title": "Cost-Based Threshold Optimization",
        "subtitle": "Notebook Section 7 (Part 2): Unconstrained vs Constrained, Trade-Off Analysis",
        "narrative": (
            "The threshold optimization (applied to LightGBM Bayesian probabilities) reveals "
            "the core business trade-off. Unconstrained optimization (threshold 0.720) "
            "minimizes cost at $326K but catches only 17.1% of fraud — unacceptable. "
            "Constraining recall >=75% raises the threshold to 0.410 with $578K cost but "
            "catches 76.6% of fraud. The extra $252K prevents 2,743 additional frauds worth "
            "$206K in direct losses plus reputation and regulatory costs."
        ),
        "matrix": CH_THRESHOLD_MATRIX,
        "figures": [
            ("cost_vs_threshold.png",
             "Figure 4: Cost vs Threshold Optimization (LightGBM Bayesian, Validation Set). "
             "The U-shaped curve shows total cost across all thresholds. Left side: low "
             "threshold = many false alarms (high FP cost). Right side: high threshold = many "
             "missed frauds (high FN cost). Red dashed line marks the unconstrained optimum at 0.720."),
        ],
        "callouts": [
            ("caution",
             "Pure cost minimization (17.1% recall) is a degenerate solution: the model "
             "essentially ignores fraud because individual misses ($227) are relatively cheap "
             "compared to the aggregate review costs ($10 each for ~29K false alarms). The 75% recall constraint "
             "is a business guardrail against this pathological optimization."),
        ],
    },
    {
        "number": 8,
        "title": "Production Strategy & Test Set Evaluation",
        "subtitle": "Notebook Sections 7 (Part 3) + 8: Multi-Threshold, Confusion Matrix, Model Export",
        "narrative": (
            "The final production strategy uses three tiers applied to LightGBM Bayesian "
            "probabilities: auto-block (>=0.90), manual review (0.41-0.90), and auto-approve "
            "(<0.41). On the 118,108-transaction test set, the system achieves 74.4% recall "
            "with $572,900 total cost ($4.85/txn). 54.9% of transactions are auto-approved "
            "with zero friction. Model artifacts (best_model_final.pkl, xgboost_final.pkl, "
            "scaler.pkl, threshold_config.pkl) are saved for Phase 4 and Phase 5."
        ),
        "matrix": CH_PRODUCTION_MATRIX,
        "figures": [
            ("confusion_matrix_absolute.png",
             "Figure 5: Confusion Matrix — Absolute Numbers (Test Set, 118,108 transactions). "
             "At the production threshold of 0.41 with LightGBM Bayesian: TN=64,541, "
             "FP=49,503, FN=1,039, TP=3,025. The system catches 74.4% of fraud."),
            ("confusion_matrix_percentages.png",
             "Figure 6: Confusion Matrix — Percentages (Test Set). TN=54.65% correctly "
             "auto-approved, FP=41.91% sent to manual review (false alarms), FN=0.88% fraud "
             "missed, TP=2.56% fraud correctly caught."),
        ],
        "callouts": [
            ("business",
             "The production system auto-approves 54.9% of transactions instantly and only "
             "sends 45.1% to human review. Obvious fraud (score >=0.90) is blocked without "
             "analyst intervention. This three-tier design balances fraud prevention with "
             "operational efficiency and customer experience."),
        ],
    },
]

SUMMARY_TABLE_DATA = [
    ("Models Compared", "LR, XGBoost (initial + Bayesian), LightGBM (initial + Bayesian)"),
    ("Winner", "LightGBM Bayesian (PR-AUC 0.1126, +37.1% over LR)"),
    ("Features Used", "7 (from Phase 2 tiers 1-3 + TransactionAmt)"),
    ("Training Set", "354,324 rows (60%), fraud rate 3.38%"),
    ("Validation Set", "118,108 rows (20%), fraud rate 3.90%"),
    ("Test Set", "118,108 rows (20%), fraud rate 3.44%"),
    ("Baseline PR-AUC (LR)", "0.0821"),
    ("XGBoost Initial PR-AUC", "0.1093"),
    ("XGBoost Grid Search PR-AUC", "0.1098"),
    ("LightGBM Initial PR-AUC", "0.1095"),
    ("XGBoost Bayesian PR-AUC", "0.1116"),
    ("LightGBM Bayesian PR-AUC (Winner)", "0.1126"),
    ("Grid Search (XGB only)", "6 combos: depth {4,6,8}, trees {100,200}, lr {0.05,0.1}"),
    ("Bayesian Optimization", "Optuna TPE, 30 trials each for XGBoost and LightGBM"),
    ("Unconstrained Threshold", "0.720 (cost $326K, recall 17.1%)"),
    ("Constrained Threshold (75% recall)", "0.420 (cost $742K, recall 76.2%)"),
    ("Auto-Block Threshold", ">= 0.90"),
    ("Manual Review Range", "0.420 - 0.90"),
    ("Test Recall", "73.8% (3,001 of 4,064 frauds)"),
    ("Test Precision", "5.8%"),
    ("Test Total Cost", "$730,482 ($6.18/txn)"),
    ("Cost Ratio (FN:FP)", "22.7:1 ($227 vs $10)"),
    ("Model Artifacts", "best_model_final.pkl (LGB), xgboost_final.pkl (compat), scaler.pkl, threshold_config.pkl"),
]

GLOSSARY = [
    ("Auto-Approve", "Transactions with fraud scores below the review threshold (<0.41) are automatically approved with no human review."),
    ("Auto-Block", "Transactions with very high fraud scores (>=0.90) are automatically blocked without waiting for human review."),
    ("Bayesian Optimization", "A method for finding optimal hyperparameters by using a probabilistic model to decide which configurations to test next. More efficient than grid or random search."),
    ("Class Weight (Balanced)", "A Logistic Regression setting that automatically weights fraud samples higher to compensate for their rarity (3.5%)."),
    ("Confusion Matrix", "A 2x2 table showing True Negatives, False Positives, False Negatives, and True Positives."),
    ("Constrained Optimization", "Finding the best solution (lowest cost) while meeting a requirement (e.g., catch at least 75% of fraud)."),
    ("Cost Sweep", "Testing every possible threshold and calculating the total cost at each one to find the optimal operating point."),
    ("Dynamic Winner Selection", "A code pattern where the best model's probabilities are stored under generic variable names, so all downstream cells automatically use the winner."),
    ("F1-Score", "The harmonic mean of precision and recall. A balanced metric, but less useful than PR-AUC for imbalanced data."),
    ("Grid Search", "Testing a fixed list of hyperparameter combinations to find the best configuration. Used here for XGBoost (6 combos)."),
    ("Hyperparameters", "Model settings that control how the model learns (e.g., tree depth, number of trees, learning speed)."),
    ("is_unbalance", "LightGBM parameter that automatically reweights fraud samples to handle class imbalance. Equivalent to XGBoost's scale_pos_weight."),
    ("Leaf-Wise Growth", "LightGBM's tree-building strategy: splits the single leaf with the highest gain at each step. More efficient than XGBoost's level-wise approach."),
    ("LightGBM", "Light Gradient Boosting Machine. Uses leaf-wise tree growth and is_unbalance for class imbalance. Won the 4-model comparison with PR-AUC 0.1126."),
    ("Manual Review", "Transactions with borderline fraud scores (0.41-0.90) are sent to human analysts for investigation."),
    ("Optuna", "Python framework for Bayesian hyperparameter optimization. Uses Tree-structured Parzen Estimator (TPE) to intelligently select next trial configurations."),
    ("PR-AUC", "Precision-Recall Area Under Curve. The primary metric for comparing models on imbalanced fraud data."),
    ("Precision", "Of all transactions flagged as fraud, what percentage are actually fraud? Low precision = many false alarms."),
    ("Recall (Sensitivity)", "Of all actual frauds, what percentage does the model catch? 74.4% recall = catches 3 of 4 frauds."),
    ("Scale Pos Weight", "An XGBoost parameter (28.56) that tells the model each fraud case is worth 28.56 legitimate cases."),
    ("StandardScaler", "A preprocessing step that rescales features to mean=0 and standard deviation=1, required for Logistic Regression."),
    ("Threshold", "The fraud probability cutoff above which a transaction is flagged. Lower = more sensitive; higher = more selective."),
    ("XGBoost", "Extreme Gradient Boosting. Builds many small decision trees sequentially using level-wise growth, each correcting the previous one's errors."),
]


# ══════════════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_paragraph_spacing(paragraph, before=0, after=0, line=240):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_formatted_text(cell, text, font_name="Calibri", font_size=10,
                       bold=False, color_hex=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=20, after=20, line=240)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def add_callout_box(doc, text, box_type="insight"):
    colors = {
        "insight": (C_INSIGHT_BG, C_INSIGHT_BORDER, "Key Insight"),
        "business": (C_BUSINESS_BG, C_BUSINESS_BORDER, "Business Impact"),
        "caution": (C_CAUTION_BG, C_CAUTION_BORDER, "Caution"),
    }
    bg_color, border_color, label = colors.get(box_type, colors["insight"])

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=120, line=264)

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), bg_color)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)

    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), border_color)
    left.set(qn("w:space"), "4")
    pBdr.append(left)
    pPr.append(pBdr)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:right"), "360")
    pPr.append(ind)

    label_run = p.add_run(f"{label}: ")
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(10)
    label_run.font.bold = True
    label_run.font.color.rgb = RGBColor.from_string(border_color)

    content_run = p.add_run(text)
    content_run.font.name = "Calibri"
    content_run.font.size = Pt(10)
    content_run.font.italic = True
    content_run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)


def add_figure(doc, image_filename, caption_text):
    img_path = FIGURES_DIR / image_filename
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cap, before=40, after=200, line=240)
        cap_run = cap.add_run(caption_text)
        cap_run.font.name = "Calibri"
        cap_run.font.size = Pt(9)
        cap_run.font.italic = True
        cap_run.font.color.rgb = RGBColor.from_string("666666")
    else:
        p = doc.add_paragraph(
            f"[Figure not available: {image_filename} — run notebook to generate]"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor.from_string("999999")


def add_matrix_table(doc, matrix_data):
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9360")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    headers = [
        "Layer", "Technical\n(BDS Colleague)",
        "Business\n(Manager / Regulator)", "Simple\n(Grandmother)"
    ]
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, layer in enumerate(LAYERS):
        row_idx = i + 1
        layer_cell = table.cell(row_idx, 0)
        set_cell_shading(layer_cell, C_MED_BLUE)
        set_cell_margins(layer_cell, top=60, bottom=60, left=100, right=100)
        add_formatted_text(layer_cell, layer, font_size=9, bold=True, color_hex=C_WHITE)
        layer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for j in range(3):
            cell = table.cell(row_idx, j + 1)
            bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            text = matrix_data.get((i, j), "")
            add_formatted_text(cell, text, font_size=9)

    doc.add_paragraph("")
    return table


def add_section_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)
    return heading


def add_page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
#  DOCUMENT CONSTRUCTION
# ══════════════════════════════════════════════════════════════════════

def add_cover_page(doc):
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Agent Fraud Sentinel")
    run.font.name = "Calibri"
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Model Training Deep Dive: 4 Layers \u00d7 3 Perspectives")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("\u2500" * 50)
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)
    run.font.size = Pt(12)

    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ref.add_run("Notebook: 03_model_training.ipynb")
    run.font.name = "Consolas"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        "IEEE-CIS Fraud Detection Dataset\n"
        "4-Model Comparison: LR, XGBoost, LightGBM | Bayesian Optimization (Optuna)\n"
        "Multi-Threshold Production Strategy | LightGBM Bayesian Winner | 74.4% Recall"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(C_DARK_GRAY)

    for _ in range(3):
        doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("BAFS Project \u2014 February 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("888888")

    add_page_break(doc)


def add_toc_placeholder(doc):
    add_section_heading(doc, "Table of Contents", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)
    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar_separate)

    placeholder_run = p.add_run(
        "(Right-click here and select 'Update Field' to generate Table of Contents)"
    )
    placeholder_run.font.name = "Calibri"
    placeholder_run.font.size = Pt(10)
    placeholder_run.font.italic = True
    placeholder_run.font.color.rgb = RGBColor.from_string("999999")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    placeholder_run._r.append(fldChar_end)

    add_page_break(doc)


def add_framework_explanation(doc):
    add_section_heading(doc, "The 4-Layer \u00d7 3-Perspective Framework", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "This document examines each section of the Model Training notebook through "
        "two dimensions: four analytical layers and three audience perspectives. This "
        "framework ensures complete understanding \u2014 from raw technical detail to "
        "business impact \u2014 accessible to any reader regardless of their background."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    add_section_heading(doc, "The Four Layers", level=3)
    for title_text, desc in [
        ("Layer 1 \u2014 WHAT did I do?",
         "Describes the concrete actions: models trained, thresholds tested, strategies designed."),
        ("Layer 2 \u2014 WHY did I do it?",
         "Explains the motivation: why four models, why Bayesian tuning, why three tiers."),
        ("Layer 3 \u2014 HOW does it work?",
         "Details the mechanics: algorithms, cost functions, threshold selection, model saving."),
        ("Layer 4 \u2014 WHAT does the bank gain?",
         "Translates results into value: fraud caught, cost savings, operational efficiency."),
    ]:
        p = doc.add_paragraph()
        bold_run = p.add_run(title_text + "  ")
        bold_run.font.name = "Calibri"
        bold_run.font.size = Pt(11)
        bold_run.font.bold = True
        bold_run.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)
        desc_run = p.add_run(desc)
        desc_run.font.name = "Calibri"
        desc_run.font.size = Pt(11)

    add_section_heading(doc, "The Three Perspectives", level=3)
    for title_text, desc in [
        ("Technical (BDS Colleague)",
         "Uses ML terminology, references hyperparameters, metrics, and code patterns."),
        ("Business (Manager / Regulator)",
         "Focuses on cost trade-offs, compliance, and operational strategy. No coding assumed."),
        ("Simple (Grandmother)",
         "Uses everyday analogies and plain language. No technical background assumed."),
    ]:
        p = doc.add_paragraph()
        bold_run = p.add_run(title_text + "  ")
        bold_run.font.name = "Calibri"
        bold_run.font.size = Pt(11)
        bold_run.font.bold = True
        bold_run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)
        desc_run = p.add_run(desc)
        desc_run.font.name = "Calibri"
        desc_run.font.size = Pt(11)


def add_production_strategy_table(doc):
    """Add the multi-threshold strategy table unique to this notebook."""
    add_section_heading(doc, "Production Strategy: Three-Tier Decision Framework", level=2)

    data = [
        ("Auto-Block", ">= 0.90", "Instant block (automated)",
         "$5.00", "0.01% of transactions"),
        ("Manual Review", "0.41 - 0.90", "Analyst investigation",
         "$10.00", "45.1% of transactions"),
        ("Auto-Approve", "< 0.41", "Approved (no action)",
         "$0.00", "54.9% of transactions"),
    ]

    table = doc.add_table(rows=len(data) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Tier", "Score Range", "Action", "Cost/Txn", "Volume"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (tier, score, action, cost, volume) in enumerate(data):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        for j, text in enumerate([tier, score, action, cost, volume]):
            cell = table.cell(i + 1, j)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=80, right=80)
            add_formatted_text(cell, text, font_size=10, bold=(j == 0))
            if j > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")


def add_chapter(doc, chapter_data):
    num = chapter_data["number"]
    title = chapter_data["title"]

    add_page_break(doc)
    add_section_heading(doc, f"Chapter {num}: {title}", level=1)

    if chapter_data.get("subtitle"):
        p = doc.add_paragraph()
        run = p.add_run(chapter_data["subtitle"])
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=200, line=276)
    run = p.add_run(chapter_data["narrative"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    add_section_heading(doc, "Analysis Matrix", level=2)
    add_matrix_table(doc, chapter_data["matrix"])

    for box_type, text in chapter_data.get("callouts", []):
        add_callout_box(doc, text, box_type)

    for fig_filename, caption in chapter_data.get("figures", []):
        add_figure(doc, fig_filename, caption)


def add_summary_statistics_table(doc):
    add_section_heading(doc, "Model Training Summary", level=2)

    table = doc.add_table(rows=len(SUMMARY_TABLE_DATA) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Metric", "Value"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (metric, value) in enumerate(SUMMARY_TABLE_DATA):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        for j, text in enumerate([metric, value]):
            cell = table.cell(i + 1, j)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=100, right=100)
            add_formatted_text(cell, text, font_size=10, bold=(j == 0))
            if j == 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


GALLERY_FIGURES = [
    ("xgb_feature_importance.png",
     "Figure 1: XGBoost Feature Importance (Gain). Shows relative importance of each "
     "of the 7 engineered features. TransactionAmt and txn_count_24hr dominate."),
    ("lgb_feature_importance.png",
     "Figure 2: LightGBM Feature Importance. Independent cross-validation of the "
     "XGBoost ranking: both architectures agree that velocity and amount are top signals."),
    ("pr_curve_comparison.png",
     "Figure 3: Precision-Recall Curve — LightGBM Bayesian (PR-AUC 0.1126) vs "
     "Logistic Regression (PR-AUC 0.0821). LightGBM outperforms across all recall "
     "levels. Red dashed baseline represents the no-model fraud rate (3.5%)."),
    ("cost_vs_threshold.png",
     "Figure 4: Cost vs Threshold Optimization (LightGBM Bayesian, Validation Set). "
     "U-shaped total cost curve. Unconstrained optimum at 0.720 catches only 17.1% "
     "of fraud. Constrained threshold 0.410 achieves 76.6% recall at $578K."),
    ("confusion_matrix_absolute.png",
     "Figure 5: Confusion Matrix — Absolute Numbers (Test Set, 118,108 transactions). "
     "At threshold 0.41 with LightGBM Bayesian: TN=64,541, FP=49,503, FN=1,039, TP=3,025."),
    ("confusion_matrix_percentages.png",
     "Figure 6: Confusion Matrix — Percentages (Test Set). TN=54.65%, FP=41.91%, "
     "FN=0.88%, TP=2.56%. Recall=74.4%, Specificity=56.6%."),
]


def add_appendix_gallery(doc):
    add_page_break(doc)
    add_section_heading(doc, "Appendix A: Visualization Gallery", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "All model training visualizations from notebook 03, presented with detailed captions."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    set_paragraph_spacing(p, after=200)

    for fig_filename, caption in GALLERY_FIGURES:
        add_figure(doc, fig_filename, caption)


def add_appendix_glossary(doc):
    add_page_break(doc)
    add_section_heading(doc, "Appendix B: Glossary", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "Key terms used throughout this document, defined for non-technical readers."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    set_paragraph_spacing(p, after=200)

    table = doc.add_table(rows=len(GLOSSARY) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        border.set(qn("w:space"), "0")
        borders.append(border)
    tblPr.append(borders)

    for j, header in enumerate(["Term", "Definition"]):
        cell = table.cell(0, j)
        set_cell_shading(cell, C_DARK_BLUE)
        set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
        add_formatted_text(cell, header, font_size=10, bold=True, color_hex=C_WHITE)

    for i, (term, definition) in enumerate(GLOSSARY):
        bg = C_WHITE if i % 2 == 0 else C_LIGHT_GRAY
        cell_term = table.cell(i + 1, 0)
        set_cell_shading(cell_term, bg)
        set_cell_margins(cell_term, top=30, bottom=30, left=100, right=100)
        add_formatted_text(cell_term, term, font_size=10, bold=True)

        cell_def = table.cell(i + 1, 1)
        set_cell_shading(cell_def, bg)
        set_cell_margins(cell_def, top=30, bottom=30, left=100, right=100)
        add_formatted_text(cell_def, definition, font_size=10)


def setup_header_footer(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        h_para.text = ""
        run_left = h_para.add_run("Agent Fraud Sentinel \u2014 Model Training")
        run_left.font.name = "Calibri"
        run_left.font.size = Pt(8)
        run_left.font.color.rgb = RGBColor.from_string("999999")
        h_para.add_run("\t\t")
        run_right = h_para.add_run("BAFS")
        run_right.font.name = "Calibri"
        run_right.font.size = Pt(8)
        run_right.font.bold = True
        run_right.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

        footer = section.footer
        footer.is_linked_to_previous = False
        f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        f_para.text = ""
        f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_page = f_para.add_run("Page ")
        run_page.font.name = "Calibri"
        run_page.font.size = Pt(8)
        run_page.font.color.rgb = RGBColor.from_string("999999")
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run_page._r.append(fldChar1)
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run_page._r.append(instrText)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        run_page._r.append(fldChar2)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run_page._r.append(fldChar3)


# ══════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════

def main():
    print("Generating Model Training Analysis Matrix document...")

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i in range(1, 4):
        h_style = doc.styles[f"Heading {i}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

    # ── Build Document ─────────────────────────────────────────────
    add_cover_page(doc)
    add_toc_placeholder(doc)

    # Chapter 0
    ch0 = ALL_CHAPTERS[0]
    add_section_heading(doc, f"Chapter 0: {ch0['title']}", level=1)
    p = doc.add_paragraph()
    run = p.add_run(ch0["subtitle"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=200, line=276)
    run = p.add_run(ch0["narrative"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    add_framework_explanation(doc)
    add_production_strategy_table(doc)

    add_section_heading(doc, "Full-Project Analysis Matrix", level=2)
    add_matrix_table(doc, ch0["matrix"])

    for box_type, text in ch0.get("callouts", []):
        add_callout_box(doc, text, box_type)

    # Chapters 1-8
    for chapter_data in ALL_CHAPTERS[1:]:
        add_chapter(doc, chapter_data)

    # Summary
    add_summary_statistics_table(doc)

    # Appendix A: Visualization Gallery
    add_appendix_gallery(doc)

    # Appendix B: Glossary
    add_appendix_glossary(doc)

    # Header/Footer
    setup_header_footer(doc)

    # Save
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"File size: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
